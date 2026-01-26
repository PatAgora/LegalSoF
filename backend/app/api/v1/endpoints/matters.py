"""
Matters API Endpoints
"""
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from datetime import datetime
from pydantic import BaseModel
import json
import os

from app.db.session import get_db
from app.models import Matter, MatterStatus

router = APIRouter()


# Request/Response Models
class StatusUpdateRequest(BaseModel):
    new_status: str
    reason: Optional[str] = None
    auto_transition: bool = False


class StatusUpdateResponse(BaseModel):
    success: bool
    previous_status: str
    new_status: str
    completion_percentage: int
    message: str
    auto_transitions_applied: List[str] = []


@router.get("/matters")
async def list_matters(
    skip: int = 0,
    limit: int = 100,
    status: Optional[str] = None,
    risk_rating: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """
    List all matters with optional filtering
    """
    # Use sync DB helper for blocking operations
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    sync_db = SessionLocal()
    
    try:
        query = sync_db.query(Matter)
        
        # Apply filters if provided
        if status:
            query = query.filter(Matter.status == status)
        if risk_rating:
            query = query.filter(Matter.risk_rating == risk_rating)
        
        # Get matters with pagination
        matters = query.offset(skip).limit(limit).all()
        
        # Load SoF assessment storage for completion percentages
        import os
        storage_file = "/tmp/sof_assessment_storage.json"
        storage = {}
        
        if os.path.exists(storage_file):
            with open(storage_file, 'r') as f:
                storage = json.load(f)
        
        # Convert to dict format
        result = []
        for matter in matters:
            # Get SoF data for this matter
            sof_data = storage.get(str(matter.id))
            
            # Calculate completion percentage
            completion_percentage = calculate_completion_percentage(matter, sof_data)
            
            result.append({
                "id": matter.id,
                "reference_number": matter.reference_number,
                "client_name": matter.client_name,
                "client_entity_name": matter.client_entity_name,
                "transaction_type": matter.transaction_type.value if matter.transaction_type else None,
                "target_business_name": matter.target_business_name,
                "target_amount": float(matter.target_amount) if matter.target_amount else None,
                "target_currency": "GBP",  # Default currency
                "transaction_date": matter.transaction_date.isoformat() if matter.transaction_date else None,
                "status": matter.status.value if matter.status else "draft",
                "risk_rating": matter.risk_rating.value if matter.risk_rating else "medium",
                "description": matter.description,
                "created_at": matter.created_at.isoformat() if matter.created_at else None,
                "updated_at": matter.updated_at.isoformat() if matter.updated_at else None,
                "completion_percentage": completion_percentage,
            })
        
        return result
    
    finally:
        sync_db.close()


@router.get("/matters/{matter_id}")
async def get_matter(
    matter_id: int,
    db: Session = Depends(get_db)
):
    """
    Get a single matter by ID
    """
    # Use sync DB helper for blocking operations
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    sync_db = SessionLocal()
    
    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data for completion percentage calculation
        import os
        import json
        storage_file = "/tmp/sof_assessment_storage.json"
        sof_data = None
        
        if os.path.exists(storage_file):
            with open(storage_file, 'r') as f:
                storage = json.load(f)
                matter_key = str(matter_id)
                if matter_key in storage:
                    sof_data = storage[matter_key]
        
        # Calculate completion percentage (this function will be available after workflow implementation)
        completion_percentage = calculate_completion_percentage(matter, sof_data)
        
        return {
            "id": matter.id,
            "reference_number": matter.reference_number,
            "client_name": matter.client_name,
            "client_entity_name": matter.client_entity_name,
            "transaction_type": matter.transaction_type.value if matter.transaction_type else None,
            "target_business_name": matter.target_business_name,
            "target_amount": float(matter.target_amount) if matter.target_amount else None,
            "target_currency": "GBP",
            "transaction_date": matter.transaction_date.isoformat() if matter.transaction_date else None,
            "status": matter.status.value if matter.status else "draft",
            "risk_rating": matter.risk_rating.value if matter.risk_rating else "medium",
            "risk_rating_auto": matter.risk_rating_auto.value if matter.risk_rating_auto else None,
            "risk_rating_override": matter.risk_rating_override.value if matter.risk_rating_override else None,
            "risk_notes": matter.risk_notes,
            "description": matter.description,
            "created_at": matter.created_at.isoformat() if matter.created_at else None,
            "updated_at": matter.updated_at.isoformat() if matter.updated_at else None,
            "completion_percentage": completion_percentage,
        }
    
    finally:
        sync_db.close()


@router.get("/matters/{matter_id}/report")
async def generate_matter_report(
    matter_id: int,
    db: Session = Depends(get_db)
):
    """
    Generate comprehensive Matter Summary Report as Word document (.docx)
    
    Includes:
    - Matter overview and details
    - Status and progress summary
    - SoF Assessment results
    - Transaction Review alerts
    - Document checklist
    - Next actions
    """
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from fastapi.responses import StreamingResponse
    import json
    
    # Get sync DB session
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    sync_db = SessionLocal()
    
    try:
        # Get matter details
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment storage
        import os
        storage_file = "/tmp/sof_assessment_storage.json"
        sof_data = None
        transaction_data = None
        
        if os.path.exists(storage_file):
            with open(storage_file, 'r') as f:
                storage = json.load(f)
                matter_key = str(matter_id)
                if matter_key in storage:
                    sof_data = storage[matter_key]
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ==================== TITLE ====================
        title = doc.add_heading('Matter Summary Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ==================== MATTER DETAILS ====================
        header = doc.add_paragraph()
        header.add_run(f'Matter Reference: {matter.reference_number}').bold = True
        header.add_run(f'\nClient: {matter.client_name}')
        header.add_run(f'\nReport Generated: {datetime.now().strftime("%d %B %Y at %H:%M")}')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # ==================== SECTION 1: MATTER OVERVIEW ====================
        doc.add_heading('1. Matter Overview', level=1)
        
        overview_table = doc.add_table(rows=8, cols=2)
        overview_table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Reference Number', matter.reference_number or 'N/A'),
            ('Client Name', matter.client_name or 'N/A'),
            ('Transaction Type', matter.transaction_type.value.replace('_', ' ').title() if matter.transaction_type else 'N/A'),
            ('Target Amount', f"£{matter.target_amount:,.2f} {matter.target_currency}" if matter.target_amount else 'N/A'),
            ('Status', matter.status.value if matter.status else 'DRAFT'),
            ('Risk Rating', matter.risk_rating.value if matter.risk_rating else 'MEDIUM'),
            ('Created Date', matter.created_at.strftime("%d %B %Y") if matter.created_at else 'N/A'),
            ('Last Updated', matter.updated_at.strftime("%d %B %Y") if matter.updated_at else 'N/A'),
        ]
        
        for idx, (label, value) in enumerate(rows_data):
            row = overview_table.rows[idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # ==================== SECTION 2: SoF ASSESSMENT ====================
        doc.add_heading('2. Source of Funds Assessment', level=1)
        
        if sof_data and 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            
            # Assessment Status
            para = doc.add_paragraph()
            para.add_run('Assessment Status: ').bold = True
            
            outcome = assessment.get('outcome', {})
            status = outcome.get('status', 'NOT COMPLETED').upper()
            confidence = outcome.get('confidence', 0)
            
            para.add_run(f"{status} ({confidence}% confidence)")
            
            # Claims Summary
            claims = assessment.get('claims', [])
            if claims:
                doc.add_paragraph()
                claims_heading = doc.add_paragraph()
                claims_heading.add_run('Claims Verified:').bold = True
                
                for idx, claim in enumerate(claims, 1):
                    claim_para = doc.add_paragraph(style='List Bullet')
                    source_type = claim.get('source_type', 'Unknown')
                    amount = claim.get('expected_amount', 0)
                    claim_para.add_run(f"Claim {idx}: {source_type} - £{amount:,.2f}")
                
                # Evidence Matches
                evidence_matches = assessment.get('evidence_matches', [])
                if evidence_matches:
                    doc.add_paragraph()
                    evidence_heading = doc.add_paragraph()
                    evidence_heading.add_run('Verification Results:').bold = True
                    
                    for idx, evidence in enumerate(evidence_matches, 1):
                        verified = evidence.get('verified', False)
                        doc_verified = evidence.get('document_verified', False)
                        
                        status_icon = '✅' if (verified and doc_verified) else '⚠️' if verified else '❌'
                        status_text = 'FULLY VERIFIED' if (verified and doc_verified) else 'BANK ONLY' if verified else 'NOT VERIFIED'
                        
                        ev_para = doc.add_paragraph(style='List Bullet')
                        ev_para.add_run(f"{status_icon} Claim {idx}: {status_text}")
            
            # Next Actions
            next_actions = assessment.get('next_actions', {})
            questions = next_actions.get('questions', [])
            documents = next_actions.get('documents', [])
            
            if questions or documents:
                doc.add_paragraph()
                doc.add_paragraph().add_run('Outstanding Actions:').bold = True
                
                if questions:
                    doc.add_paragraph().add_run('Questions for Client:').bold = True
                    for q in questions[:5]:
                        doc.add_paragraph(q, style='List Bullet')
                
                if documents:
                    doc.add_paragraph().add_run('Documents Required:').bold = True
                    for d in documents[:5]:
                        doc.add_paragraph(d, style='List Bullet')
        else:
            para = doc.add_paragraph()
            para.add_run('Status: ').bold = True
            para.add_run('Assessment not yet completed')
        
        doc.add_paragraph()
        
        # ==================== SECTION 3: TRANSACTION REVIEW ====================
        doc.add_heading('3. Transaction Review & AML Alerts', level=1)
        
        if sof_data and 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            tr_summary = assessment.get('transaction_review_summary', {})
            
            total_alerts = tr_summary.get('total_alerts', 0)
            critical_alerts = tr_summary.get('critical_alerts', 0)
            high_alerts = tr_summary.get('high_alerts', 0)
            medium_alerts = tr_summary.get('medium_alerts', 0)
            
            # Alert Summary
            alert_para = doc.add_paragraph()
            alert_para.add_run('Alert Summary: ').bold = True
            
            if total_alerts > 0:
                alert_para.add_run(f"{total_alerts} alert(s) identified")
                
                alert_breakdown = doc.add_paragraph(style='List Bullet')
                alert_breakdown.add_run(f"🔴 CRITICAL: {critical_alerts}")
                
                alert_breakdown = doc.add_paragraph(style='List Bullet')
                alert_breakdown.add_run(f"🟠 HIGH: {high_alerts}")
                
                alert_breakdown = doc.add_paragraph(style='List Bullet')
                alert_breakdown.add_run(f"🟡 MEDIUM: {medium_alerts}")
                
                # Key Concerns
                key_concerns = tr_summary.get('key_concerns', [])
                if key_concerns:
                    doc.add_paragraph()
                    doc.add_paragraph().add_run('Key Concerns:').bold = True
                    for concern in key_concerns[:5]:
                        doc.add_paragraph(concern, style='List Bullet')
            else:
                alert_para.add_run('No alerts identified ✅')
        else:
            para = doc.add_paragraph()
            para.add_run('Status: ').bold = True
            para.add_run('Transaction review not yet completed')
        
        doc.add_paragraph()
        
        # ==================== SECTION 4: DOCUMENT CHECKLIST ====================
        doc.add_heading('4. Document Checklist', level=1)
        
        if sof_data and 'uploaded_files' in sof_data:
            uploaded_files = sof_data.get('uploaded_files', [])
            
            if uploaded_files:
                doc_para = doc.add_paragraph()
                doc_para.add_run(f'Total Documents Uploaded: {len(uploaded_files)}').bold = True
                doc.add_paragraph()
                
                # Group by category
                categories = {}
                for file_info in uploaded_files:
                    category = file_info.get('category', 'unknown')
                    if category not in categories:
                        categories[category] = []
                    categories[category].append(file_info)
                
                for category, files in categories.items():
                    cat_para = doc.add_paragraph()
                    cat_para.add_run(f"{category.replace('_', ' ').title()}: {len(files)} file(s)").bold = True
                    
                    for file_info in files[:3]:  # Limit to first 3 per category
                        file_para = doc.add_paragraph(style='List Bullet')
                        file_para.add_run(f"✓ {file_info.get('filename', 'Unknown')}")
                    
                    if len(files) > 3:
                        doc.add_paragraph(f"... and {len(files) - 3} more", style='List Bullet')
            else:
                doc.add_paragraph('No documents uploaded yet')
        else:
            doc.add_paragraph('Document upload not yet started')
        
        doc.add_paragraph()
        
        # ==================== SECTION 5: RECOMMENDATIONS ====================
        doc.add_heading('5. Recommendations & Next Steps', level=1)
        
        # Determine recommendations based on status
        recommendations = []
        
        if sof_data and 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            outcome = assessment.get('outcome', {})
            status = outcome.get('status', '').lower()
            
            if status == 'sufficient':
                recommendations.append('✅ Matter has sufficient source of funds documentation')
                recommendations.append('✅ All material claims have been verified')
                recommendations.append('➡️ Matter can proceed to completion subject to standard monitoring')
            elif status == 'borderline':
                recommendations.append('⚠️ Additional verification recommended before proceeding')
                recommendations.append('➡️ Review outstanding questions and documents')
                recommendations.append('➡️ Consider enhanced due diligence for high-risk elements')
            else:
                recommendations.append('❌ Insufficient source of funds documentation')
                recommendations.append('➡️ Request missing documents before proceeding')
                recommendations.append('➡️ Obtain client responses to outstanding questions')
            
            # Check for critical alerts
            tr_summary = assessment.get('transaction_review_summary', {})
            if tr_summary.get('critical_alerts', 0) > 0:
                recommendations.append('🔴 CRITICAL: Address transaction review alerts before proceeding')
        else:
            recommendations.append('➡️ Complete SoF Assessment to generate recommendations')
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
        
        # ==================== FOOTER ====================
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('_' * 80)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = doc.add_paragraph()
        footer_text.add_run('\nEnd of Report')
        footer_text.add_run(f'\nGenerated: {datetime.now().strftime("%d %B %Y at %H:%M")}')
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Return as downloadable Word document
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=Matter_Summary_{matter.reference_number}.docx"
            }
        )
    
    finally:
        sync_db.close()


# ==================== WORKFLOW ENGINE ====================

def calculate_completion_percentage(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> int:
    """
    Calculate matter completion percentage based on workflow state and data
    """
    percentage = 0
    
    # Base percentage by status
    status_percentages = {
        MatterStatus.DRAFT: 10,
        MatterStatus.AWAITING_CLIENT: 20,
        MatterStatus.CLIENT_UPLOADING: 40,
        MatterStatus.UNDER_REVIEW: 60,
        MatterStatus.QUERIES_RAISED: 50,  # Step back
        MatterStatus.APPROVED: 90,
        MatterStatus.REJECTED: 100,  # Terminal state
        MatterStatus.COMPLETED: 100,
    }
    
    percentage = status_percentages.get(matter.status, 0)
    
    # Add bonus for SoF assessment completion
    if sof_data:
        if 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            outcome = assessment.get('outcome', {})
            
            # Add 10% if assessment run
            percentage = min(percentage + 10, 100)
            
            # Add 10% more if sufficient
            if outcome.get('status', '').lower() == 'sufficient':
                percentage = min(percentage + 10, 100)
            
            # Add 5% if all claims verified
            evidence_matches = assessment.get('evidence_matches', [])
            if evidence_matches:
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                if all_verified:
                    percentage = min(percentage + 5, 100)
    
    return percentage


def get_next_status_suggestions(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """
    Get smart status suggestions based on current state and data
    """
    current_status = matter.status
    suggestions = []
    
    # Status transition rules with reasons
    transitions = {
        MatterStatus.DRAFT: [
            {'status': MatterStatus.AWAITING_CLIENT, 'reason': 'Ready to request documents from client', 'auto': False},
            {'status': MatterStatus.CLIENT_UPLOADING, 'reason': 'Client portal access granted', 'auto': False},
        ],
        MatterStatus.AWAITING_CLIENT: [
            {'status': MatterStatus.CLIENT_UPLOADING, 'reason': 'Client started uploading documents', 'auto': True},
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'Skip to review (documents received externally)', 'auto': False},
        ],
        MatterStatus.CLIENT_UPLOADING: [
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'All required documents uploaded', 'auto': True},
        ],
        MatterStatus.UNDER_REVIEW: [
            {'status': MatterStatus.QUERIES_RAISED, 'reason': 'Additional information required', 'auto': False},
            {'status': MatterStatus.APPROVED, 'reason': 'Review complete - all requirements met', 'auto': True},
            {'status': MatterStatus.REJECTED, 'reason': 'Cannot proceed with this matter', 'auto': False},
        ],
        MatterStatus.QUERIES_RAISED: [
            {'status': MatterStatus.AWAITING_CLIENT, 'reason': 'Waiting for client response', 'auto': False},
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'Client responded - resume review', 'auto': False},
        ],
        MatterStatus.APPROVED: [
            {'status': MatterStatus.COMPLETED, 'reason': 'Matter completed successfully', 'auto': False},
            {'status': MatterStatus.QUERIES_RAISED, 'reason': 'New issue identified', 'auto': False},
        ],
        MatterStatus.REJECTED: [
            # Terminal state - no transitions
        ],
        MatterStatus.COMPLETED: [
            # Terminal state - no transitions
        ],
    }
    
    base_suggestions = transitions.get(current_status, [])
    
    # Check if auto-transitions should be suggested based on data
    if sof_data and 'assessment_result' in sof_data:
        assessment = sof_data['assessment_result']
        outcome = assessment.get('outcome', {})
        
        # If UNDER_REVIEW and assessment is SUFFICIENT, suggest APPROVED
        if current_status == MatterStatus.UNDER_REVIEW:
            if outcome.get('status', '').lower() == 'sufficient':
                # Check if all claims fully verified
                evidence_matches = assessment.get('evidence_matches', [])
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                
                # Check no critical alerts
                tr_summary = assessment.get('transaction_review_summary', {})
                critical_alerts = tr_summary.get('critical_alerts', 0)
                
                if all_verified and critical_alerts == 0:
                    # Mark APPROVED as auto-ready
                    for sug in base_suggestions:
                        if sug['status'] == MatterStatus.APPROVED:
                            sug['auto'] = True
                            sug['reason'] = '✅ AUTO: All claims verified, no critical alerts'
        
        # If UNDER_REVIEW and there are outstanding questions, suggest QUERIES_RAISED
        if current_status == MatterStatus.UNDER_REVIEW:
            next_actions = assessment.get('next_actions', {})
            questions = next_actions.get('questions', [])
            if len(questions) > 0:
                for sug in base_suggestions:
                    if sug['status'] == MatterStatus.QUERIES_RAISED:
                        sug['auto'] = True
                        sug['reason'] = f'⚠️ AUTO: {len(questions)} outstanding question(s)'
    
    return base_suggestions


def apply_auto_transitions(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> List[str]:
    """
    Apply automatic status transitions based on workflow rules
    """
    transitions_applied = []
    
    # Check if documents have been uploaded (transition to UNDER_REVIEW)
    if matter.status == MatterStatus.CLIENT_UPLOADING and sof_data:
        if 'uploaded_files' in sof_data and len(sof_data.get('uploaded_files', [])) > 0:
            # Check if client_info and bank_statements uploaded
            has_client_info = any(f.get('category') == 'client_info' for f in sof_data.get('uploaded_files', []))
            has_bank = any(f.get('category') == 'bank_statement' for f in sof_data.get('uploaded_files', []))
            
            if has_client_info and has_bank:
                matter.status = MatterStatus.UNDER_REVIEW
                transitions_applied.append(f"CLIENT_UPLOADING → UNDER_REVIEW (documents uploaded)")
    
    # Check if assessment complete and sufficient (transition to APPROVED)
    if matter.status == MatterStatus.UNDER_REVIEW and sof_data:
        if 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            outcome = assessment.get('outcome', {})
            
            if outcome.get('status', '').lower() == 'sufficient':
                # Check all claims verified
                evidence_matches = assessment.get('evidence_matches', [])
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                
                # Check no critical alerts
                tr_summary = assessment.get('transaction_review_summary', {})
                critical_alerts = tr_summary.get('critical_alerts', 0)
                
                if all_verified and critical_alerts == 0:
                    matter.status = MatterStatus.APPROVED
                    transitions_applied.append(f"UNDER_REVIEW → APPROVED (all verified, no critical alerts)")
    
    return transitions_applied


@router.patch("/matters/{matter_id}/status")
async def update_matter_status(
    matter_id: int,
    request: StatusUpdateRequest,
    db: Session = Depends(get_db)
):
    """
    Update matter status with workflow validation and smart transitions
    
    Features:
    - Workflow validation (prevent invalid transitions)
    - Automatic status transitions based on data
    - Completion percentage calculation
    - Status change audit trail
    """
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    import os
    
    # Get sync DB session
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    sync_db = SessionLocal()
    
    try:
        # Get matter
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data
        storage_file = "/tmp/sof_assessment_storage.json"
        sof_data = None
        
        if os.path.exists(storage_file):
            import json
            with open(storage_file, 'r') as f:
                storage = json.load(f)
                matter_key = str(matter_id)
                if matter_key in storage:
                    sof_data = storage[matter_key]
        
        # Store previous status
        previous_status = matter.status
        
        # Validate new status
        try:
            new_status = MatterStatus(request.new_status)
        except ValueError:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid status: {request.new_status}. Valid statuses: {[s.value for s in MatterStatus]}"
            )
        
        # Check if transition is allowed
        suggestions = get_next_status_suggestions(matter, sof_data)
        allowed_statuses = [s['status'] for s in suggestions]
        
        # Allow same status (for refresh) or any suggested status
        if new_status != previous_status and new_status not in allowed_statuses:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid transition from {previous_status.value} to {new_status.value}. Allowed: {[s.value for s in allowed_statuses]}"
            )
        
        # Apply status change
        matter.status = new_status
        matter.updated_at = datetime.now()
        
        # Apply automatic transitions if requested
        auto_transitions = []
        if request.auto_transition:
            auto_transitions = apply_auto_transitions(matter, sof_data)
        
        # Calculate new completion percentage
        completion_percentage = calculate_completion_percentage(matter, sof_data)
        
        # Commit changes
        sync_db.commit()
        sync_db.refresh(matter)
        
        # Build response
        message = f"Status updated from {previous_status.value} to {matter.status.value}"
        if auto_transitions:
            message += f". Auto-transitions: {', '.join(auto_transitions)}"
        
        return StatusUpdateResponse(
            success=True,
            previous_status=previous_status.value,
            new_status=matter.status.value,
            completion_percentage=completion_percentage,
            message=message,
            auto_transitions_applied=auto_transitions
        )
    
    finally:
        sync_db.close()


@router.get("/matters/{matter_id}/status/suggestions")
async def get_status_suggestions(
    matter_id: int,
    db: Session = Depends(get_db)
):
    """
    Get smart status transition suggestions for a matter
    
    Returns suggested next statuses with reasons and auto-transition flags
    """
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    import os
    
    # Get sync DB session
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    sync_db = SessionLocal()
    
    try:
        # Get matter
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data
        storage_file = "/tmp/sof_assessment_storage.json"
        sof_data = None
        
        if os.path.exists(storage_file):
            import json
            with open(storage_file, 'r') as f:
                storage = json.load(f)
                matter_key = str(matter_id)
                if matter_key in storage:
                    sof_data = storage[matter_key]
        
        # Get suggestions
        suggestions = get_next_status_suggestions(matter, sof_data)
        
        # Calculate current completion percentage
        completion_percentage = calculate_completion_percentage(matter, sof_data)
        
        # Check if auto-transitions available
        auto_transitions = apply_auto_transitions(matter, sof_data)
        
        return {
            'current_status': matter.status.value,
            'completion_percentage': completion_percentage,
            'suggestions': [
                {
                    'status': s['status'].value,
                    'reason': s['reason'],
                    'auto_recommended': s.get('auto', False)
                }
                for s in suggestions
            ],
            'auto_transitions_available': len(auto_transitions) > 0,
            'auto_transition_preview': auto_transitions
        }
    
    finally:
        sync_db.close()
