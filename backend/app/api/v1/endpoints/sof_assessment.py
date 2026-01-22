"""
Source of Funds Assessment API Endpoints

100% LOCAL - No external API calls
Handles file uploads, processing, and SoF assessment results
"""
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, Form
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from datetime import datetime
import json

from app.db.session import get_db
from app.services.file_processor import file_processor
from app.services.sof_assessment_engine import SoFAssessmentEngine
from app.models import Matter

router = APIRouter()


# Sync DB helper (for blocking file operations)
def get_sync_db():
    """Get synchronous database session for blocking operations"""
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.core.config import settings
    
    db_url = str(settings.DATABASE_URL).replace("sqlite+aiosqlite", "sqlite")
    engine = create_engine(db_url)
    SessionLocal = sessionmaker(bind=engine)
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


import json
import os
from pathlib import Path

# Persistent storage using JSON file (survives backend restarts)
STORAGE_FILE = Path("/tmp/sof_assessment_storage.json")

def load_storage() -> Dict[int, Dict[str, Any]]:
    """Load storage from file"""
    if STORAGE_FILE.exists():
        try:
            with open(STORAGE_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_storage(storage: Dict[int, Dict[str, Any]]):
    """Save storage to file"""
    # Convert matter_id keys to strings for JSON
    storage_str_keys = {str(k): v for k, v in storage.items()}
    with open(STORAGE_FILE, 'w') as f:
        json.dump(storage_str_keys, f)

# Load storage on module import
assessment_storage = load_storage()
# Convert string keys back to ints
assessment_storage = {int(k): v for k, v in assessment_storage.items()}


@router.post("/matters/{matter_id}/sof-assessment/upload")
async def upload_sof_files(
    matter_id: int,
    file: UploadFile = File(...),
    file_category: str = Form(...),  # 'client_info', 'bank_statement', 'supporting_doc'
    db: Session = Depends(get_sync_db)
):
    """
    Upload and process a file for SoF assessment
    
    File categories:
    - client_info: JSON with client details, purchase info, SoF explanation
    - bank_statement: CSV or PDF bank statement
    - supporting_doc: PDF supporting document
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Initialize storage for this matter if needed
    if matter_id not in assessment_storage:
        assessment_storage[matter_id] = {
            "client_info": None,
            "bank_statements": [],
            "supporting_docs": [],
            "uploaded_files": [],
            "status": "pending",
            "last_updated": None
        }
    
    # Determine file type from extension and MIME
    file_ext = file.filename.split('.')[-1].lower() if file.filename else ''
    
    if file_category == 'client_info':
        if file_ext != 'json':
            raise HTTPException(
                status_code=400, 
                detail="Client info must be JSON file"
            )
        file_type = 'json'
    
    elif file_category == 'bank_statement':
        if file_ext == 'csv':
            file_type = 'csv'
        elif file_ext == 'pdf':
            file_type = 'pdf'
        else:
            raise HTTPException(
                status_code=400,
                detail="Bank statement must be CSV or PDF"
            )
    
    elif file_category == 'supporting_doc':
        if file_ext != 'pdf':
            raise HTTPException(
                status_code=400,
                detail="Supporting documents must be PDF"
            )
        file_type = 'pdf'
    
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file category: {file_category}"
        )
    
    # Process the file
    result = await file_processor.process_upload(file, file_type)
    
    if not result['success']:
        raise HTTPException(
            status_code=400,
            detail=f"File processing failed: {result['error']}"
        )
    
    # Validate result based on file category
    if file_category == 'bank_statement':
        if 'bank_statements' not in result.get('data', {}):
            raise HTTPException(
                status_code=400,
                detail="PDF file does not contain valid bank statement data. No transactions could be extracted."
            )
    
    # Store the processed data
    storage = assessment_storage[matter_id]
    
    if file_category == 'client_info':
        storage['client_info'] = result['data']
    
    elif file_category == 'bank_statement':
        # Merge bank statements
        new_transactions = result['data']['bank_statements']
        storage['bank_statements'].extend(new_transactions)
    
    elif file_category == 'supporting_doc':
        # Add filename and upload timestamp to the document data for audit trail
        doc_data = result['data'].copy()
        doc_data['filename'] = file.filename
        doc_data['uploaded_at'] = datetime.utcnow().isoformat()
        storage['supporting_docs'].append(doc_data)
    
    # Track uploaded file
    storage['uploaded_files'].append({
        "filename": file.filename,
        "category": file_category,
        "file_type": result['file_type'],
        "uploaded_at": datetime.utcnow().isoformat(),
        "records_count": result['data'].get('transaction_count') or 1
    })
    
    storage['last_updated'] = datetime.utcnow().isoformat()
    storage['status'] = 'files_uploaded'
    
    # Persist storage to file
    save_storage(assessment_storage)
    
    return {
        "success": True,
        "matter_id": matter_id,
        "file_category": file_category,
        "filename": file.filename,
        "records_processed": result['data'].get('transaction_count') or 1,
        "message": f"File processed successfully"
    }


@router.get("/matters/{matter_id}/sof-assessment/status")
async def get_sof_assessment_status(
    matter_id: int,
    db: Session = Depends(get_sync_db)
):
    """
    Get current status of SoF assessment for a matter
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Get assessment data
    if matter_id not in assessment_storage:
        return {
            "matter_id": matter_id,
            "status": "no_data",
            "uploaded_files": [],
            "ready_for_assessment": False
        }
    
    storage = assessment_storage[matter_id]
    
    # Check if ready for assessment
    has_client_info = storage['client_info'] is not None
    has_bank_statements = len(storage['bank_statements']) > 0
    ready = has_client_info and has_bank_statements
    
    return {
        "matter_id": matter_id,
        "status": storage['status'],
        "uploaded_files": storage['uploaded_files'],
        "files_summary": {
            "client_info": "uploaded" if has_client_info else "missing",
            "bank_statements_count": len(storage['bank_statements']),
            "supporting_docs_count": len(storage['supporting_docs'])
        },
        "ready_for_assessment": ready,
        "last_updated": storage['last_updated']
    }


@router.post("/matters/{matter_id}/sof-assessment/run")
async def run_sof_assessment(
    matter_id: int,
    db: Session = Depends(get_sync_db)
):
    """
    Run SoF assessment engine on uploaded data
    Integrates with Transaction Review automatically
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Get assessment data
    if matter_id not in assessment_storage:
        raise HTTPException(
            status_code=400,
            detail="No files uploaded for this matter"
        )
    
    storage = assessment_storage[matter_id]
    
    # Validate required data
    if not storage['client_info']:
        raise HTTPException(
            status_code=400,
            detail="Client info file required (JSON with client details, purchase, SoF explanation)"
        )
    
    if not storage['bank_statements']:
        raise HTTPException(
            status_code=400,
            detail="At least one bank statement file required (CSV or PDF)"
        )
    
    # Extract data
    client_info = storage['client_info']['client_info']
    purchase = storage['client_info']['purchase']
    sof_explanation = storage['client_info']['sof_explanation']
    bank_statements = storage['bank_statements']
    
    # Check if client_info JSON has explicit claims array
    # If so, convert to structured format for the assessment engine
    print(f"\n=== CLAIMS CHECK ===")
    print(f"Has 'claims' key: {'claims' in storage['client_info']}")
    if 'claims' in storage['client_info']:
        print(f"Claims array: {storage['client_info']['claims']}")
    
    if 'claims' in storage['client_info'] and storage['client_info']['claims']:
        # Convert claims array to structured format
        print(f"✅ Converting claims array to structured format")
        sof_explanation = {
            'sources': storage['client_info']['claims']
        }
        print(f"sof_explanation type: {type(sof_explanation)}")
    print(f"====================\n")
    
    # IMPORTANT: Build fresh known_documents list from current uploads only
    # Do NOT accumulate from previous assessments
    known_documents = []
    supporting_docs_data = storage['supporting_docs']  # Full document data with extracted info
    
    # Add uploaded supporting docs to known documents
    for doc in storage['supporting_docs']:
        doc_type = doc.get('document_type', 'unknown')
        if doc_type != 'unknown':
            known_documents.append(doc_type)
    
    flags = storage['client_info'].get('flags', {})
    constraints = storage['client_info'].get('constraints', {})
    
    # DEBUG LOGGING
    print(f"\n=== SoF ASSESSMENT DEBUG ===")
    print(f"Matter ID: {matter_id}")
    print(f"Supporting docs uploaded: {len(supporting_docs_data)}")
    for idx, doc in enumerate(supporting_docs_data):
        print(f"  Doc {idx}: Type={doc.get('document_type')}, Has extracted_data={bool(doc.get('extracted_data'))}")
        if doc.get('extracted_data'):
            print(f"    Extracted fields: {list(doc.get('extracted_data', {}).keys())}")
    print(f"Known documents: {known_documents}")
    print(f"===========================\n")
    
    # Run assessment engine
    engine = SoFAssessmentEngine(matter_id=matter_id, db=db)
    
    try:
        assessment_result = engine.assess(
            client_info=client_info,
            purchase=purchase,
            sof_explanation=sof_explanation,
            bank_statements=bank_statements,
            known_documents=known_documents,
            supporting_docs_data=supporting_docs_data,  # Pass full document data
            constraints=constraints,
            flags=flags
        )
        
        # Store result
        storage['assessment_result'] = assessment_result
        storage['status'] = 'completed'
        storage['last_updated'] = datetime.utcnow().isoformat()
        
        # Persist storage to file
        save_storage(assessment_storage)
        
        return {
            "success": True,
            "matter_id": matter_id,
            "assessment": assessment_result
        }
    
    except Exception as e:
        storage['status'] = 'error'
        storage['error'] = str(e)
        raise HTTPException(
            status_code=500,
            detail=f"Assessment engine error: {str(e)}"
        )


@router.get("/matters/{matter_id}/sof-assessment/results")
async def get_sof_assessment_results(
    matter_id: int,
    db: Session = Depends(get_sync_db)
):
    """
    Get full SoF assessment results
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Reload storage from file to get latest data
    assessment_storage_reloaded = load_storage()
    assessment_storage_reloaded = {int(k): v for k, v in assessment_storage_reloaded.items()}
    
    # Get assessment data
    if matter_id not in assessment_storage_reloaded:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )
    
    storage = assessment_storage_reloaded[matter_id]
    
    if storage['status'] != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage['status']})"
        )
    
    return {
        "matter_id": matter_id,
        "assessment": storage['assessment_result'],
        "metadata": {
            "uploaded_files": storage['uploaded_files'],
            "bank_statements_count": len(storage['bank_statements']),
            "supporting_docs_count": len(storage['supporting_docs']),
            "completed_at": storage['last_updated']
        }
    }


@router.get("/matters/{matter_id}/sof-assessment/file-note")
async def download_file_note(
    matter_id: int,
    db: Session = Depends(get_sync_db)
):
    """
    Download audit-ready file note as Word document (.docx)
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Get assessment data
    if matter_id not in assessment_storage:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )
    
    storage = assessment_storage[matter_id]
    
    if storage['status'] != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage['status']})"
        )
    
    # Get file note text
    file_note = storage['assessment_result']['file_note_summary']
    
    # Generate Word document
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from fastapi.responses import StreamingResponse
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Source of Funds Assessment File Note', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add matter reference
    matter_ref = doc.add_paragraph()
    matter_ref.add_run(f'Matter Reference: {matter.reference_number or f"MAT-{matter.id}"}').bold = True
    matter_ref.add_run(f'\nClient: {matter.client_name}')
    matter_ref.add_run(f'\nDate: {datetime.now().strftime("%d %B %Y")}')
    matter_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse and format the file note sections
    sections_text = file_note.split('===')
    
    for section in sections_text:
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        section_title = lines[0].strip()
        section_content = '\n'.join(lines[1:]).strip()
        
        # Add section heading
        if section_title:
            heading = doc.add_heading(section_title, level=1)
            # Style the heading
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add section content
        if section_content:
            # Split content into paragraphs
            paragraphs = section_content.split('\n\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                    
                para = doc.add_paragraph()
                
                # Handle bullet points and formatting
                if para_text.strip().startswith('•') or para_text.strip().startswith('-'):
                    para.style = 'List Bullet'
                    para_text = para_text.strip()[1:].strip()
                
                # Add text with basic formatting
                for line in para_text.split('\n'):
                    if not line.strip():
                        continue
                    
                    # Bold text between ** **
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            para.add_run(part)
                        else:
                            para.add_run(part).bold = True
                    
                    para.add_run('\n')
                
                # Set paragraph spacing
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('_' * 80)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph()
    footer_text.add_run('\nDocument generated on: ' + datetime.now().strftime("%d %B %Y at %H:%M"))
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Return as downloadable Word document
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=SoF_File_Note_Matter_{matter.reference_number or matter.id}.docx"
        }
    )


@router.post("/matters/{matter_id}/sof-assessment/accept-differences")
async def accept_claim_differences(
    matter_id: int,
    request: dict,
    db: Session = Depends(get_sync_db)
):
    """
    Manually accept differences for a specific claim after manual review
    
    Body:
        {
            "claim_index": int,
            "accepted_by": str,
            "reason": str (optional)
        }
    """
    from datetime import datetime
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Get assessment data
    if matter_id not in assessment_storage:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )
    
    storage = assessment_storage[matter_id]
    
    if storage['status'] != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage['status']})"
        )
    
    claim_index = request.get('claim_index')
    accepted_by = request.get('accepted_by', 'User')
    reason = request.get('reason', 'Manual review completed - differences accepted')
    
    if claim_index is None:
        raise HTTPException(status_code=400, detail="claim_index is required")
    
    # Get the evidence match for this claim
    evidence_matches = storage['assessment_result']['evidence_matches']
    
    if claim_index < 0 or claim_index >= len(evidence_matches):
        raise HTTPException(status_code=400, detail="Invalid claim_index")
    
    evidence = evidence_matches[claim_index]
    
    # Check if this claim requires review
    if not evidence.get('document_verification', {}).get('requires_review', False):
        raise HTTPException(
            status_code=400,
            detail="This claim does not require review (already at 100% confidence)"
        )
    
    # Update the manual review status
    doc_verification = evidence['document_verification']
    doc_verification['manual_review_status'] = 'accepted'
    doc_verification['manually_accepted_by'] = accepted_by
    doc_verification['manually_accepted_at'] = datetime.utcnow().isoformat()
    doc_verification['acceptance_reason'] = reason
    
    # Mark all differences as accepted
    if 'differences' in doc_verification:
        for diff in doc_verification['differences']:
            diff['accepted'] = True
            diff['accepted_by'] = accepted_by
            diff['accepted_at'] = datetime.utcnow().isoformat()
    
    # Persist the updated storage
    save_storage(assessment_storage)
    
    return {
        "success": True,
        "message": f"Differences accepted for claim {claim_index + 1}",
        "claim_index": claim_index,
        "accepted_by": accepted_by,
        "accepted_at": doc_verification['manually_accepted_at'],
        "reason": reason,
        "updated_status": {
            "requires_review": doc_verification.get('requires_review', False),
            "manual_review_status": doc_verification.get('manual_review_status'),
            "confidence": doc_verification.get('confidence', 0)
        }
    }


@router.delete("/matters/{matter_id}/sof-assessment/reset")
async def reset_sof_assessment(
    matter_id: int,
    db: Session = Depends(get_sync_db)
):
    """
    Reset/clear SoF assessment data for a matter
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Clear storage
    if matter_id in assessment_storage:
        del assessment_storage[matter_id]
    
    # Persist storage to file
    save_storage(assessment_storage)
    
    return {
        "success": True,
        "matter_id": matter_id,
        "message": "SoF assessment data cleared"
    }
