"""
Matters API Endpoints
"""
import json
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone
from pydantic import BaseModel


def _parse_json(raw):
    """Parse a JSON text column into a Python object; None/blank → None."""
    if not raw:
        return None
    if isinstance(raw, (dict, list)):
        return raw
    try:
        return json.loads(raw)
    except (TypeError, ValueError):
        return None

from app.db.session import get_db, get_sync_db, get_sync_session
from app.api.dependencies.auth import get_current_active_user, require_analyst, require_admin
from app.models import Matter, MatterStatus, RiskRating, TransactionType
from app.models.user import User, UserRole
from app.models.assessment_storage import AssessmentStorage
from app.models.audit import AuditLog, AuditLogAction
from app.models.status_history import MatterStatusHistory
from app.models.document_verification import DocumentVerification
from app.models.statement_validation import StatementValidation
from app.models.notification import Notification

router = APIRouter()


def _safe_load_storage_db(db) -> Dict:
    """Load all assessment storage data from PostgreSQL (keyed by str(matter_id))."""
    try:
        rows = db.query(AssessmentStorage).all()
        return {str(row.matter_id): row.data for row in rows if row.data}
    except Exception as e:
        print(f"Warning: Could not load assessment storage from DB: {e}")
        return {}


def _safe_load_storage_single(db, matter_id: int) -> Optional[Dict]:
    """Load assessment data for a single matter from PostgreSQL."""
    try:
        row = db.query(AssessmentStorage).filter(AssessmentStorage.matter_id == matter_id).first()
        if row and row.data:
            return row.data
        return None
    except Exception as e:
        print(f"Warning: Could not load assessment storage for matter {matter_id}: {e}")
        return None


# Request/Response Models
class StatusUpdateRequest(BaseModel):
    new_status: str
    reason: Optional[str] = None
    auto_transition: bool = False


class StatusUpdateResponse(BaseModel):
    success: bool
    previous_status: str
    new_status: str
    completion_percentage: int
    message: str
    auto_transitions_applied: List[str] = []


class CreateMatterRequest(BaseModel):
    client_name: str
    reference: Optional[str] = None
    transaction_value: Optional[float] = None
    risk_level: Optional[str] = "medium"
    description: Optional[str] = None
    status: Optional[str] = "draft"


@router.post("/matters")
async def create_matter(
    request: CreateMatterRequest,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """Create a new matter"""
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    
    try:
        # Generate reference if not provided
        reference = request.reference
        if not reference:
            # Get highest matter ID and create new reference
            existing = sync_db.query(Matter).order_by(Matter.id.desc()).first()
            next_num = (existing.id + 1) if existing else 1
            reference = f"MAT-{datetime.now().year}-{str(next_num).zfill(3)}"
        
        # Map risk level string to enum
        risk_map = {
            'low': RiskRating.LOW,
            'medium': RiskRating.MEDIUM,
            'high': RiskRating.HIGH,
            'critical': RiskRating.CRITICAL
        }
        risk_rating = risk_map.get((request.risk_level or 'medium').lower(), RiskRating.MEDIUM)
        
        # Map status string to enum
        status_map = {
            'draft': MatterStatus.DRAFT,
            'under_review': MatterStatus.UNDER_REVIEW,
            'approved': MatterStatus.APPROVED
        }
        status = status_map.get(request.status.lower(), MatterStatus.DRAFT)
        
        # Create the new matter. Transaction value is optional — when
        # not supplied it defaults to 0 (the column is NOT NULL).
        new_matter = Matter(
            reference_number=reference,
            client_name=request.client_name,
            target_amount=request.transaction_value if request.transaction_value is not None else 0,
            target_currency="GBP",
            transaction_type=TransactionType.PROPERTY_PURCHASE,
            status=status,
            risk_rating=risk_rating,
            description=request.description or "Property purchase",
            created_by_id=current_user.id
        )
        
        sync_db.add(new_matter)
        sync_db.commit()
        sync_db.refresh(new_matter)

        # Audit log for matter creation
        audit = AuditLog(
            matter_id=new_matter.id,
            user_id=current_user.id,
            action=AuditLogAction.MATTER_CREATED,
            entity_type="matter",
            entity_id=new_matter.id,
            description=f"Matter {new_matter.reference_number} created for client {new_matter.client_name}",
            details={
                "reference": new_matter.reference_number,
                "client": new_matter.client_name,
                "amount": float(new_matter.target_amount),
                "risk_rating": new_matter.risk_rating.value,
            },
        )
        sync_db.add(audit)
        sync_db.commit()

        return {
            "id": new_matter.id,
            "reference_number": new_matter.reference_number,
            "client_name": new_matter.client_name,
            "target_amount": float(new_matter.target_amount),
            "status": new_matter.status.value,
            "risk_rating": new_matter.risk_rating.value,
            "created_at": new_matter.created_at.isoformat() if new_matter.created_at else None
        }
    
    except Exception as e:
        sync_db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to create matter: {str(e)}")
    
    finally:
        sync_db.close()


@router.get("/matters")
async def list_matters(
    skip: int = 0,
    limit: int = 100,
    status: Optional[str] = None,
    risk_rating: Optional[str] = None,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """
    List all matters with optional filtering
    """
    # Use shared sync DB session for blocking operations
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()

    try:
        query = sync_db.query(Matter)

        # Exclude archived matters (database-driven, replaces hardcoded exclusion list)
        query = query.filter(Matter.is_archived == False)

        # Apply filters if provided
        if status:
            query = query.filter(Matter.status == status)
        if risk_rating:
            query = query.filter(Matter.risk_rating == risk_rating)

        # Get matters with pagination
        matters = query.offset(skip).limit(limit).all()
        
        # Load SoF assessment storage for completion percentages
        storage = _safe_load_storage_db(sync_db)
        
        # Convert to dict format
        result = []
        for matter in matters:
            # Get SoF data for this matter
            sof_data = storage.get(str(matter.id))
            
            # Calculate completion percentage
            completion_percentage = calculate_completion_percentage(matter, sof_data)
            
            result.append({
                "id": matter.id,
                "reference_number": matter.reference_number,
                "client_name": matter.client_name,
                "client_entity_name": matter.client_entity_name,
                "transaction_type": matter.transaction_type.value if matter.transaction_type else None,
                "target_business_name": matter.target_business_name,
                "target_amount": float(matter.target_amount) if matter.target_amount else None,
                "target_currency": "GBP",  # Default currency
                "transaction_date": matter.transaction_date.isoformat() if matter.transaction_date else None,
                "status": matter.status.value if matter.status else "draft",
                "risk_rating": matter.risk_rating.value if matter.risk_rating else "medium",
                "description": matter.description,
                "created_at": matter.created_at.isoformat() if matter.created_at else None,
                "updated_at": matter.updated_at.isoformat() if matter.updated_at else None,
                "completion_percentage": completion_percentage,
            })
        
        return result
    
    finally:
        sync_db.close()


@router.get("/matters/{matter_id}")
async def get_matter(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """
    Get a single matter by ID
    """
    # Use shared sync DB session for blocking operations
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    
    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data for completion percentage calculation
        sof_data = _safe_load_storage_single(sync_db, matter_id)
        
        # Calculate completion percentage (this function will be available after workflow implementation)
        completion_percentage = calculate_completion_percentage(matter, sof_data)
        
        return {
            "id": matter.id,
            "reference_number": matter.reference_number,
            "client_name": matter.client_name,
            "client_entity_name": matter.client_entity_name,
            "transaction_type": matter.transaction_type.value if matter.transaction_type else None,
            "target_business_name": matter.target_business_name,
            "target_amount": float(matter.target_amount) if matter.target_amount else None,
            "target_currency": "GBP",
            "transaction_date": matter.transaction_date.isoformat() if matter.transaction_date else None,
            "status": matter.status.value if matter.status else "draft",
            "risk_rating": matter.risk_rating.value if matter.risk_rating else "medium",
            "risk_rating_auto": matter.risk_rating_auto.value if matter.risk_rating_auto else None,
            "risk_rating_override": bool(matter.risk_rating_override),
            "risk_notes": matter.risk_notes,
            "risk_factors": _parse_json(getattr(matter, 'risk_factors', None)),
            "risk_assessed_at": (matter.risk_assessed_at.isoformat()
                                 if getattr(matter, 'risk_assessed_at', None) else None),
            "risk_assessed_by": getattr(matter, 'risk_assessed_by', None),
            "source_of_wealth": getattr(matter, 'source_of_wealth', None),
            "source_of_wealth_evidence": getattr(matter, 'source_of_wealth_evidence', None),
            "compliance_submitted_at": (matter.compliance_submitted_at.isoformat()
                                        if getattr(matter, 'compliance_submitted_at', None) else None),
            "compliance_submitted_by": getattr(matter, 'compliance_submitted_by', None),
            "description": matter.description,
            "created_at": matter.created_at.isoformat() if matter.created_at else None,
            "updated_at": matter.updated_at.isoformat() if matter.updated_at else None,
            "completion_percentage": completion_percentage,
        }

    finally:
        sync_db.close()


# ── Request models for the risk assessment / compliance endpoints ──
class RiskAssessmentRequest(BaseModel):
    risk_rating: str                                  # low | medium | high | critical
    risk_factors: Optional[Dict[str, Any]] = None     # {category: [indicators]}
    risk_notes: Optional[str] = None
    source_of_wealth: Optional[str] = None
    source_of_wealth_evidence: Optional[str] = None


@router.put("/matters/{matter_id}/risk-assessment")
async def save_risk_assessment(
    matter_id: int,
    request: RiskAssessmentRequest,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db),
):
    """
    Record the matter-level risk assessment and Source of Wealth.

    Captures the documented matter risk assessment required by MLR 2017
    Reg 18 / LSAG §4.3-4.4 (rating + higher-risk indicators by
    category), the reviewer's risk reasoning, and the client's Source
    of Wealth narrative + supporting evidence. The saved rating drives
    the per-risk-tier rule configuration used by every assessment.
    """
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        risk_map = {
            'low': RiskRating.LOW, 'medium': RiskRating.MEDIUM,
            'high': RiskRating.HIGH, 'critical': RiskRating.CRITICAL,
        }
        new_rating = risk_map.get((request.risk_rating or 'medium').lower(), RiskRating.MEDIUM)

        matter.risk_rating = new_rating
        matter.risk_rating_override = True
        matter.risk_notes = request.risk_notes
        matter.risk_factors = json.dumps(request.risk_factors) if request.risk_factors else None
        matter.risk_assessed_at = datetime.now(timezone.utc)
        matter.risk_assessed_by = current_user.full_name or current_user.email
        matter.source_of_wealth = request.source_of_wealth
        matter.source_of_wealth_evidence = request.source_of_wealth_evidence

        sync_db.add(AuditLog(
            matter_id=matter.id,
            user_id=current_user.id,
            action=AuditLogAction.UPDATED,
            entity_type="matter",
            entity_id=matter.id,
            description=(
                f"Matter risk assessment saved by {matter.risk_assessed_by}: "
                f"rating {new_rating.value}."
            ),
            details={
                "risk_rating": new_rating.value,
                "risk_factors": request.risk_factors or {},
                "source_of_wealth_provided": bool(request.source_of_wealth),
            },
        ))
        sync_db.commit()

        return {
            "success": True,
            "matter_id": matter_id,
            "risk_rating": new_rating.value,
            "risk_assessed_at": matter.risk_assessed_at.isoformat(),
            "risk_assessed_by": matter.risk_assessed_by,
        }
    finally:
        sync_db.close()


@router.post("/matters/{matter_id}/send-to-compliance")
async def send_to_compliance(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db),
):
    """
    Submit the matter to the firm's compliance function for review.

    Records the submission on the matter, writes an audit-trail entry,
    and notifies every admin (compliance) user so the matter and its
    file appear in their queue.
    """
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        submitted_by = current_user.full_name or current_user.email
        now = datetime.now(timezone.utc)
        matter.compliance_submitted_at = now
        matter.compliance_submitted_by = submitted_by

        sync_db.add(AuditLog(
            matter_id=matter.id,
            user_id=current_user.id,
            action=AuditLogAction.UPDATED,
            entity_type="matter",
            entity_id=matter.id,
            description=(
                f"Matter {matter.reference_number} sent to the compliance "
                f"team for review by {submitted_by}."
            ),
            details={"reference": matter.reference_number, "submitted_by": submitted_by},
        ))

        # Notify every admin user — they act as the compliance function.
        recipients = sync_db.query(User).filter(User.role == UserRole.ADMIN).all()
        for admin in recipients:
            try:
                sync_db.add(Notification(
                    user_id=admin.id,
                    matter_id=matter.id,
                    type="compliance_review",
                    title="Matter sent for compliance review",
                    message=(
                        f"{submitted_by} has sent matter {matter.reference_number} "
                        f"({matter.client_name}) to the compliance team for review."
                    ),
                ))
            except Exception:
                pass  # notification is best-effort

        sync_db.commit()
        return {
            "success": True,
            "matter_id": matter_id,
            "compliance_submitted_at": now.isoformat(),
            "compliance_submitted_by": submitted_by,
            "notified": len(recipients),
        }
    finally:
        sync_db.close()


@router.delete("/matters/{matter_id}")
async def delete_matter(
    matter_id: int,
    current_user: User = Depends(require_admin),
    db: Session = Depends(get_db),
):
    """
    Hard-delete a matter and every record associated with it: documents,
    verifications, transactions, assessments, uploaded files on disk.
    Admin-only. Irreversible.
    """
    import os, shutil

    SessionLocal = get_sync_session()
    sync_db = SessionLocal()

    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        reference_number = matter.reference_number

        # Tables that DO NOT cascade via Matter's relationships — wipe explicitly.
        # DocumentVerification and StatementValidation cascade their own
        # flags/transactions children via ondelete=CASCADE on the FK.
        sync_db.query(AssessmentStorage).filter(AssessmentStorage.matter_id == matter_id).delete(synchronize_session=False)
        sync_db.query(DocumentVerification).filter(DocumentVerification.matter_id == matter_id).delete(synchronize_session=False)
        sync_db.query(StatementValidation).filter(StatementValidation.matter_id == matter_id).delete(synchronize_session=False)
        sync_db.query(Notification).filter(Notification.matter_id == matter_id).delete(synchronize_session=False)

        # Matter delete cascades through questionnaire_responses, documents,
        # entities, funds_events, checks, notes, audit_logs, approvals,
        # status_history, transactions, transaction_alerts, kyc_profile.
        sync_db.delete(matter)
        sync_db.commit()

        # Remove any persisted upload files for this matter.
        upload_dir = f"/app/uploads/{matter_id}"
        if os.path.isdir(upload_dir):
            shutil.rmtree(upload_dir, ignore_errors=True)

        return {
            "success": True,
            "matter_id": matter_id,
            "reference_number": reference_number,
            "message": "Matter and all related records deleted",
        }

    finally:
        sync_db.close()


@router.get("/matters/{matter_id}/report")
async def generate_matter_report(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """
    Generate an SRA-grade Source of Funds / CDD Memorandum as a Word
    document (.docx).

    Structured to evidence compliance with the Money Laundering,
    Terrorist Financing and Transfer of Funds (Information on the Payer)
    Regulations 2017 ("MLR 2017") and the SRA's AML Sectoral Guidance
    for the Legal Sector. Sections cover: regulatory framework, risk
    assessment, CDD measures, source of funds and source of wealth
    analysis, document forensics, bank statement validation, funds
    lineage, outstanding matters, audit trail, reviewer attestation,
    and retention notice.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from io import BytesIO
    from fastapi.responses import StreamingResponse

    SessionLocal = get_sync_session()
    sync_db = SessionLocal()

    try:
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        sof_data = _safe_load_storage_single(sync_db, matter_id) or {}
        assessment = sof_data.get('assessment_result') or {}
        outcome = assessment.get('outcome', {}) or {}
        claims = assessment.get('claims', []) or []
        evidence_matches = assessment.get('evidence_matches', []) or []
        tr_summary = assessment.get('transaction_review_summary', {}) or {}
        funds_lineage = sof_data.get('funds_lineage') or {}
        uploaded_files = sof_data.get('uploaded_files', []) or []
        bank_statements = sof_data.get('bank_statements', []) or []
        supporting_docs = sof_data.get('supporting_docs', []) or []
        client_info = sof_data.get('client_info') or {}
        next_actions = assessment.get('next_actions', {}) or {}

        dv_rows = sync_db.query(DocumentVerification).filter(
            DocumentVerification.matter_id == matter_id
        ).order_by(DocumentVerification.id.asc()).all()

        stmt_rows = sync_db.query(StatementValidation).filter(
            StatementValidation.matter_id == matter_id
        ).order_by(StatementValidation.id.asc()).all()

        audit_rows = sync_db.query(AuditLog).filter(
            AuditLog.matter_id == matter_id
        ).order_by(AuditLog.created_at.asc()).all()

        # -----------------------------------------------------------------
        # Render helpers — keep the document construction tidy and the
        # styling consistent across every section.
        # -----------------------------------------------------------------
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10.5)

        def p(text: str = '', bold: bool = False, italic: bool = False, size: float = None, align=None):
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            if align is not None:
                para.alignment = align
            return para

        def kv_table(rows, col_widths=(Cm(5.0), Cm(11.5))):
            t = doc.add_table(rows=len(rows), cols=2)
            t.style = 'Light Grid Accent 1'
            for i, (label, value) in enumerate(rows):
                a, b = t.rows[i].cells
                a.text = str(label)
                a.paragraphs[0].runs[0].bold = True
                a.width = col_widths[0]
                b.text = '' if value is None else str(value)
                b.width = col_widths[1]
            return t

        def header_table(headers, rows):
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = 'Light Grid Accent 1'
            for j, h in enumerate(headers):
                c = t.rows[0].cells[j]
                c.text = str(h)
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].bold = True
            for i, row in enumerate(rows, start=1):
                for j, val in enumerate(row):
                    t.rows[i].cells[j].text = '' if val is None else str(val)
            return t

        def fmt_date(value, fmt: str = '%d/%m/%Y') -> str:
            if not value:
                return ''
            if isinstance(value, str):
                # Accept ISO strings; fall back to raw text.
                try:
                    return datetime.fromisoformat(value.replace('Z', '+00:00')).strftime(fmt)
                except Exception:
                    return value
            try:
                return value.strftime(fmt)
            except Exception:
                return str(value)

        def fmt_money(amount, currency: str = 'GBP') -> str:
            if amount is None:
                return 'Not specified'
            try:
                symbol = '£' if currency == 'GBP' else f"{currency} "
                return f"{symbol}{float(amount):,.2f}"
            except Exception:
                return str(amount)

        # -----------------------------------------------------------------
        # 0. COVER
        # -----------------------------------------------------------------
        title = doc.add_heading('Source of Funds Compliance Memorandum', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p('Customer Due Diligence record prepared under the Money Laundering,',
          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p('Terrorist Financing and Transfer of Funds (Information on the Payer) Regulations 2017',
          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p('and the SRA AML Sectoral Guidance for the Legal Sector.',
          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p('', align=WD_ALIGN_PARAGRAPH.CENTER)
        p('CONFIDENTIAL — for the firm\'s AML records.', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        verdict_status = (outcome.get('status') or '').lower()
        verdict_label = {
            'sufficient':  'PASS — Source of funds adequately evidenced',
            'borderline':  'REVIEW REQUIRED — Manual sign-off needed',
            'insufficient':'FAIL — Source of funds not adequately evidenced',
        }.get(verdict_status, 'INCOMPLETE — Assessment not yet finalised')

        kv_table([
            ('Matter reference',  matter.reference_number or 'Not assigned'),
            ('Client',            matter.client_name or ''),
            ('Transaction type',  (matter.transaction_type.value.replace('_', ' ').title()
                                   if matter.transaction_type else 'Not specified')),
            ('Target amount',     fmt_money(matter.target_amount, 'GBP')),
            ('Transaction date',  fmt_date(matter.transaction_date)),
            ('Matter status',     (matter.status.value if matter.status else 'draft').replace('_', ' ').title()),
            ('CDD outcome',       verdict_label),
            ('Reviewer',          current_user.full_name or current_user.email),
            ('Reviewer role',     (current_user.role.value if getattr(current_user, 'role', None) else 'Analyst').title()),
            ('Report prepared',   datetime.utcnow().strftime('%d %B %Y at %H:%M UTC')),
        ])

        doc.add_paragraph()

        # -----------------------------------------------------------------
        # 1. EXECUTIVE SUMMARY
        # -----------------------------------------------------------------
        doc.add_heading('1. Executive summary', level=1)

        total_claims    = len(claims)
        passed_claims   = sum(
            1 for ev in evidence_matches
            if (ev.get('document_verification', {}) or {}).get('manual_review_status') == 'accepted'
               or (ev.get('document_verified') is True
                   and (ev.get('document_verification', {}) or {}).get('confidence', 0) >= 0.999)
        )
        review_claims   = max(0, total_claims - passed_claims)
        docs_reviewed   = len(dv_rows)
        docs_tampered   = sum(1 for d in dv_rows if d.verdict and d.verdict.value == 'LikelyTampered')
        docs_suspicious = sum(1 for d in dv_rows if d.verdict and d.verdict.value == 'Suspicious')

        kv_table([
            ('Overall verdict',                 verdict_label),
            ('Claims declared by client',       total_claims),
            ('Claims fully evidenced',          passed_claims),
            ('Claims requiring further review', review_claims),
            ('Documents subjected to forensic checks', docs_reviewed),
            ('Documents flagged "Suspicious"',  docs_suspicious),
            ('Documents flagged "Likely tampered"', docs_tampered),
            ('Transaction-review alerts raised', tr_summary.get('total_alerts', 0)),
        ])

        if outcome.get('rationale'):
            doc.add_paragraph()
            p('Reviewer narrative', bold=True)
            doc.add_paragraph(str(outcome.get('rationale')).strip()[:2000])

        # -----------------------------------------------------------------
        # 2. REGULATORY FRAMEWORK
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('2. Regulatory framework', level=1)
        doc.add_paragraph(
            'This memorandum records the customer due diligence (CDD) and source-of-funds '
            'verification undertaken in respect of the above matter. The work has been performed '
            'with reference to:'
        )
        for ref in [
            'The Money Laundering, Terrorist Financing and Transfer of Funds (Information on the Payer) '
            'Regulations 2017 — in particular Regulations 18 (risk assessment), 27–28 (CDD), 33 (EDD), '
            '35 (PEPs), 40 (record-keeping), and 47 (training).',
            'The Proceeds of Crime Act 2002 — in particular Sections 327–330 (principal money '
            'laundering offences and reporting obligations).',
            'The Terrorism Act 2000 — Sections 15–18 and 19.',
            'The SRA Standards and Regulations 2019 — including the SRA Code of Conduct for '
            'Solicitors, RELs and RFLs (paragraphs 7.1, 7.4) and the Code of Conduct for Firms '
            '(paragraph 3).',
            'The SRA AML Sectoral Guidance for the Legal Sector (2023) and the Legal Sector '
            'Affinity Group Guidance.',
        ]:
            doc.add_paragraph(ref, style='List Bullet')

        # -----------------------------------------------------------------
        # 3. MATTER & CLIENT DETAILS
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('3. Matter and client details', level=1)
        kv_table([
            ('Matter reference',          matter.reference_number or ''),
            ('Client name',               matter.client_name or ''),
            ('Client entity (if any)',    matter.client_entity_name or 'Not applicable'),
            ('Target / counterparty',     matter.target_business_name or 'Not specified'),
            ('Transaction type',          (matter.transaction_type.value.replace('_', ' ').title()
                                            if matter.transaction_type else 'Not specified')),
            ('Target amount',             fmt_money(matter.target_amount, 'GBP')),
            ('Transaction date',          fmt_date(matter.transaction_date) or 'Not yet set'),
            ('Matter opened',             fmt_date(matter.created_at)),
            ('Matter last updated',       fmt_date(matter.updated_at)),
            ('Matter description',        (matter.description or '').strip() or 'Not provided'),
        ])

        if isinstance(client_info, dict) and client_info:
            doc.add_paragraph()
            p('Client-supplied source-of-funds explanation', bold=True)
            ci_inner = client_info.get('client_info') if isinstance(client_info.get('client_info'), dict) else client_info
            sof_explanation = (
                ci_inner.get('sof_explanation') or ci_inner.get('source_of_funds')
                or client_info.get('sof_explanation') or ''
            )
            if isinstance(sof_explanation, dict):
                sof_explanation = sof_explanation.get('explanation') or sof_explanation.get('summary') or ''
            if sof_explanation:
                doc.add_paragraph(str(sof_explanation).strip()[:2000])
            else:
                doc.add_paragraph('No free-text explanation supplied by the client.')

        # -----------------------------------------------------------------
        # 4. RISK ASSESSMENT (MLR Reg 18 / 28)
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('4. Matter risk assessment', level=1)
        kv_table([
            ('Effective risk rating',   (matter.risk_rating.value if matter.risk_rating else 'Medium').title()),
            ('Reviewer override applied', 'Yes' if matter.risk_rating_override else 'No'),
            ('Assessment recorded by',  getattr(matter, 'risk_assessed_by', None) or 'Not recorded'),
            ('Assessment recorded on',  fmt_date(getattr(matter, 'risk_assessed_at', None))),
            ('Risk reasoning',          (matter.risk_notes or '').strip() or 'No reasoning recorded.'),
        ])

        # Selected higher-risk indicators (LSAG §4.4), by category.
        _rf = _parse_json(getattr(matter, 'risk_factors', None)) or {}
        _rf_rows = []
        for _cat, _label in (
            ('client', 'Client'), ('geographic', 'Geographic'),
            ('service', 'Service / product'), ('transaction', 'Transaction'),
            ('delivery_channel', 'Delivery channel'),
        ):
            items = _rf.get(_cat) or []
            if items:
                _rf_rows.append((_label, '; '.join(str(i) for i in items)))
        if _rf_rows:
            p('Higher-risk indicators identified', bold=True)
            kv_table(_rf_rows)
        else:
            doc.add_paragraph(
                'No specific higher-risk indicators were recorded against the LSAG risk-factor '
                'categories for this matter.'
            )

        doc.add_paragraph(
            'In line with Regulation 18 MLR 2017 the matter risk has been assessed against the firm-wide '
            'risk assessment and the specific characteristics of the client, the parties, the geography, '
            'the products/services involved, the channels through which the matter is being conducted, '
            'and the transactions concerned. Enhanced Due Diligence (Regulation 33) is applied '
            'where the rating is High or Critical or where any factor in Schedule 3 / Schedule 4 to '
            'MLR 2017 was identified.'
        )

        # Source of Wealth (EDD — LSAG §6.8, §7.2).
        doc.add_paragraph()
        p('Source of Wealth', bold=True)
        _sow = (getattr(matter, 'source_of_wealth', None) or '').strip()
        _sow_ev = (getattr(matter, 'source_of_wealth_evidence', None) or '').strip()
        _rating_l = (matter.risk_rating.value if matter.risk_rating else 'medium').lower()
        _sow_required = _rating_l in ('high', 'critical')
        kv_table([
            ('Source of Wealth required',
             'Yes — Enhanced Due Diligence (high-risk matter)' if _sow_required
             else 'Risk-based — not mandatory at this rating'),
            ('Wealth narrative', _sow or 'Not recorded.'),
            ('Evidence reviewed', _sow_ev or 'Not recorded.'),
        ])
        if _sow_required and not _sow:
            p('RESERVATION: Source of Wealth is required for this high-risk matter but has '
              'not been recorded. This must be completed before sign-off.', bold=True)

        # -----------------------------------------------------------------
        # 5. CDD MEASURES APPLIED
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('5. Customer due diligence measures applied', level=1)
        doc.add_paragraph(
            'The following CDD measures were applied to satisfy Regulation 28 MLR 2017. Each control '
            'is recorded against the corresponding regulatory requirement. Documentary evidence is '
            'retained on the matter file and the assessment platform audit trail.'
        )
        cdd_rows = [
            ('Reg 28(2)(a) — Identification of the client',
             'Client name on file: %s.' % (matter.client_name or 'Not recorded')),
            ('Reg 28(2)(b) — Verification of identity',
             '%d identity / supporting document(s) uploaded and run through the document forensics '
             'pipeline (see Section 7).' % docs_reviewed),
            ('Reg 28(4) — Beneficial ownership',
             'Recorded against the matter where applicable: %s.' % (matter.client_entity_name or 'Individual client; no entity ownership chain.')),
            ('Reg 28(2)(c) — Purpose and intended nature of the business relationship',
             ((matter.description or '').strip() or 'See matter description / SoF explanation above.')),
            ('Reg 28(11) — Source of funds',
             '%d claim(s) of source-of-funds declared; %d fully evidenced, %d require further review '
             '(see Section 6).' % (total_claims, passed_claims, review_claims)),
            ('Reg 28(13) — Ongoing monitoring',
             'Bank statement validation and transaction review applied; %d alert(s) raised '
             '(see Section 8).' % (tr_summary.get('total_alerts', 0))),
            ('Reg 35 — PEP screening',
             'Not yet integrated with an external screening provider; reviewer to confirm '
             'separately. (PLATFORM LIMITATION — to be addressed before sign-off.)'),
            ('Reg 39 — Reliance on third parties',
             'No reliance placed on third parties; CDD performed by the firm directly.'),
        ]
        kv_table(cdd_rows)

        # -----------------------------------------------------------------
        # 6. SOURCE OF FUNDS — CLAIM-BY-CLAIM ANALYSIS
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('6. Source of funds analysis', level=1)

        if not claims:
            doc.add_paragraph('No source-of-funds claims have been recorded against this matter.')
        else:
            for idx, claim in enumerate(claims, 1):
                ev = evidence_matches[idx - 1] if idx - 1 < len(evidence_matches) else {}
                dv = ev.get('document_verification') or {}
                confidence = dv.get('confidence') or 0
                accepted = dv.get('manual_review_status') == 'accepted'
                fully_verified = (
                    accepted
                    or (ev.get('document_verified') is True and confidence >= 0.999)
                )
                claim_status = 'ACCEPTED (manual review)' if accepted else (
                    'VERIFIED' if fully_verified else 'REVIEW REQUIRED'
                )
                source_label = str(claim.get('source_type', 'Unknown')).replace('_', ' ').title()

                p(f'Claim {idx}: {source_label}', bold=True, size=12)
                kv_table([
                    ('Declared source',         source_label),
                    ('Declared amount',         fmt_money(claim.get('expected_amount'), 'GBP')),
                    ('Verification outcome',    claim_status),
                    ('Document-match confidence', f"{confidence * 100:.1f}%" if isinstance(confidence, (int, float)) else 'n/a'),
                    ('Bank transactions matched', len(ev.get('transactions', []) or [])),
                ])

                diffs = dv.get('differences') or []
                if diffs:
                    p('Discrepancies identified', bold=True)
                    diff_rows = []
                    for d in diffs[:8]:
                        field = (d.get('field') or '').replace('_', ' ').title() if isinstance(d, dict) else ''
                        expected = (d.get('expected') if isinstance(d, dict) else '') or ''
                        found    = (d.get('found') if isinstance(d, dict) else '') or ''
                        desc     = (d.get('description') or d.get('issue') if isinstance(d, dict) else '') or ''
                        diff_rows.append((field or 'Discrepancy',
                                          str(expected)[:120],
                                          str(found)[:120],
                                          str(desc)[:200]))
                    header_table(('Field', 'Expected', 'Found', 'Reviewer notes'), diff_rows)

                if accepted:
                    p('Reviewer disposition', bold=True)
                    accepted_by = dv.get('accepted_by') or dv.get('manually_accepted_by') or 'Reviewer'
                    accepted_at = dv.get('accepted_at') or dv.get('manually_accepted_at') or ''
                    accepted_reason = (
                        dv.get('accepted_reason') or dv.get('acceptance_reason') or
                        'Differences accepted on review.'
                    )
                    kv_table([
                        ('Accepted by',     accepted_by),
                        ('Accepted on',     fmt_date(accepted_at, '%d/%m/%Y %H:%M') or 'Not recorded'),
                        ('Rationale',       accepted_reason),
                    ])
                doc.add_paragraph()

        # -----------------------------------------------------------------
        # 7. DOCUMENT VERIFICATION (forensics)
        # -----------------------------------------------------------------
        doc.add_heading('7. Document verification', level=1)
        if not dv_rows:
            doc.add_paragraph('No documents were submitted to the forensics pipeline.')
        else:
            doc.add_paragraph(
                f'{len(dv_rows)} document(s) were subjected to automated forensic checks including '
                'text-layer / OCR consistency, image forensics (ELA, JPEG quantisation), template '
                'fingerprint matching against known bank statement layouts, PDF signature validation '
                'and statement mathematical reconciliation.'
            )
            rows = []
            for d in dv_rows:
                v = d.verdict.value if d.verdict else 'Unverified'
                phase = getattr(d, 'verification_phase', '') or ''
                score = (
                    f"{d.authenticity_score}/100" if d.authenticity_score is not None else 'n/a'
                )
                flag_count = len(d.flags or []) if hasattr(d, 'flags') else 0
                actionable_flags = (
                    [f for f in (d.flags or []) if f.severity not in (None, 'info') and not (f.code or '').endswith('_OK')]
                    if hasattr(d, 'flags') else []
                )
                rows.append((
                    d.filename,
                    v,
                    score,
                    len(actionable_flags),
                    flag_count,
                ))
            header_table(
                ('Filename', 'Verdict', 'Authenticity score', 'Actionable flags', 'Total flags'),
                rows,
            )

            # Drill into each non-Verified document so the reviewer has the
            # full forensic narrative for sign-off.
            for d in dv_rows:
                if d.verdict and d.verdict.value == 'Verified':
                    continue
                doc.add_paragraph()
                p(f'Forensic detail — {d.filename}', bold=True)
                p('Method: ' + (getattr(d, 'method_summary', '') or 'Standard pipeline'))
                if getattr(d, 'flags', None):
                    flag_rows = []
                    for f in d.flags[:25]:
                        if (f.code or '').endswith('_OK'):
                            continue
                        if f.severity == 'info':
                            continue
                        flag_rows.append((
                            (f.severity or '').upper(),
                            f.code or '',
                            (f.message or '')[:300],
                        ))
                    if flag_rows:
                        header_table(('Severity', 'Code', 'Message'), flag_rows)

        # -----------------------------------------------------------------
        # 8. BANK STATEMENT VALIDATION & TRANSACTION REVIEW
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('8. Bank statement validation and transaction review', level=1)
        kv_table([
            ('Bank statements ingested',       sum(1 for f in uploaded_files if f.get('category') == 'bank_statement')),
            ('Total transactions processed',   len(bank_statements)),
            ('Statement validations on file',  len(stmt_rows)),
            ('Transaction-review alerts',      tr_summary.get('total_alerts', 0)),
            ('  Critical',                     tr_summary.get('critical_alerts', 0)),
            ('  High',                         tr_summary.get('high_alerts', 0)),
            ('  Medium',                       tr_summary.get('medium_alerts', 0)),
            ('  Low / informational',          tr_summary.get('low_alerts', 0)),
        ])
        key_concerns = tr_summary.get('key_concerns') or []
        if key_concerns:
            p('Key concerns flagged by the transaction-review engine', bold=True)
            for c in key_concerns[:10]:
                doc.add_paragraph(str(c), style='List Bullet')

        # -----------------------------------------------------------------
        # 9. FUNDS LINEAGE
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('9. Funds lineage', level=1)
        fl_summary = (funds_lineage or {}).get('summary') if isinstance(funds_lineage, dict) else None
        if not fl_summary:
            doc.add_paragraph(
                'A funds lineage trace has not been recorded for this matter. Where the source of '
                'funds includes accumulated savings or inter-account transfers, a lineage trace '
                'should be performed to evidence accumulation back to a verified origin.'
            )
        else:
            total_amt    = float(fl_summary.get('totalAmount') or 0)
            traced_amt   = float(fl_summary.get('tracedAmount') or 0)
            untraced_amt = float(fl_summary.get('untracedAmount') or max(0, total_amt - traced_amt))
            traced_pct   = (traced_amt / total_amt * 100) if total_amt > 0 else 0
            untraced_pct = max(0.0, 100.0 - traced_pct)
            period_days  = int(fl_summary.get('accumulationPeriodDays') or 0)
            kv_table([
                ('Target credit traced',     fmt_money(total_amt, 'GBP')),
                ('Traced to a verified origin', f"{fmt_money(traced_amt, 'GBP')} ({traced_pct:.1f}%)"),
                ('Untraced / requires evidence', f"{fmt_money(untraced_amt, 'GBP')} ({untraced_pct:.1f}%)"),
                ('Matched internal transfers', fl_summary.get('matchedTransfers') or 0),
                ('External-origin entries',    fl_summary.get('externalOrigins') or 0),
                ('Accumulation period (days)', period_days),
            ])
            unresolved = funds_lineage.get('unresolved_items') or []
            if unresolved:
                p('Unresolved items requiring further evidence', bold=True)
                u_rows = []
                for u in unresolved[:15]:
                    u_rows.append((
                        fmt_money(u.get('amount'), 'GBP'),
                        fmt_date(u.get('date')),
                        u.get('account') or '',
                        (u.get('description') or '')[:120],
                    ))
                header_table(('Amount', 'Date', 'Account', 'Description'), u_rows)

        # -----------------------------------------------------------------
        # 10. OUTSTANDING MATTERS / RECOMMENDATIONS
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('10. Outstanding matters and recommendations', level=1)
        questions = next_actions.get('questions') or []
        documents = next_actions.get('documents') or []
        recommendations = []
        if verdict_status == 'sufficient':
            recommendations.append(
                'Source of funds is adequately evidenced. The matter may proceed subject to '
                'standard ongoing monitoring under Regulation 28(13).'
            )
        elif verdict_status == 'borderline':
            recommendations.append(
                'Manual review required before completion. Resolve the outstanding items listed '
                'below and revisit the assessment.'
            )
        elif verdict_status == 'insufficient':
            recommendations.append(
                'Source of funds is NOT adequately evidenced. The matter must not proceed to '
                'completion until the items below are resolved. Consider whether a Suspicious '
                'Activity Report (SAR) to the National Crime Agency is required under sections '
                '330 / 331 of the Proceeds of Crime Act 2002.'
            )
        else:
            recommendations.append('Complete the SoF assessment before drawing conclusions.')
        if tr_summary.get('critical_alerts', 0):
            recommendations.append(
                'Critical transaction-review alerts have been raised — these must be resolved or '
                'escalated to the firm\'s MLRO before sign-off.'
            )
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')

        if documents:
            p('Documents required from the client', bold=True)
            for d_ in documents[:25]:
                doc.add_paragraph(str(d_), style='List Bullet')
        if questions:
            p('Outstanding questions for the client', bold=True)
            for q in questions[:25]:
                doc.add_paragraph(str(q), style='List Bullet')

        # -----------------------------------------------------------------
        # 11. AUDIT TRAIL
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('11. Audit trail', level=1)
        if not audit_rows:
            doc.add_paragraph('No audit events have been recorded for this matter.')
        else:
            audit_data = []
            for a in audit_rows[-60:]:
                actor = '-'
                try:
                    if a.user_id:
                        u = sync_db.query(User).filter(User.id == a.user_id).first()
                        if u:
                            actor = u.full_name or u.email
                except Exception:
                    actor = '-'
                audit_data.append((
                    fmt_date(a.created_at, '%d/%m/%Y %H:%M'),
                    (a.action.value if a.action else '').replace('_', ' ').title(),
                    actor,
                    (a.description or '')[:200],
                ))
            header_table(('Timestamp', 'Action', 'Actor', 'Description'), audit_data)

        # -----------------------------------------------------------------
        # 12. REVIEWER ATTESTATION
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('12. Reviewer attestation', level=1)
        doc.add_paragraph(
            'I confirm that the customer due diligence measures recorded in this memorandum are, '
            'in my professional judgement, commensurate with the assessed risk of the matter and '
            'have been performed in accordance with the firm\'s AML policies, procedures and '
            'controls. I am satisfied (subject to the items in Section 10) that the requirements '
            'of Regulations 27–28 MLR 2017 have been met. Where any reservation is recorded, the '
            'matter has not been signed off and will not proceed to completion until that '
            'reservation is resolved.'
        )
        kv_table([
            ('Reviewer name',  current_user.full_name or current_user.email),
            ('Role',           (current_user.role.value if getattr(current_user, 'role', None) else 'Analyst').title()),
            ('Date of review', datetime.utcnow().strftime('%d %B %Y')),
            ('Signature',      '_____________________________________________'),
        ])

        # -----------------------------------------------------------------
        # 13. RETENTION
        # -----------------------------------------------------------------
        doc.add_paragraph()
        doc.add_heading('13. Record retention', level=1)
        doc.add_paragraph(
            'Pursuant to Regulation 40 MLR 2017 this memorandum and the underlying CDD evidence '
            '(client identification, source-of-funds documents, transaction records and the '
            'platform audit trail) will be retained by the firm for a period of five years from '
            'the date the business relationship ends or the occasional transaction is completed. '
            'On the expiry of that period, and unless retention is otherwise required by law or '
            'court order, the records will be securely deleted.'
        )

        # -----------------------------------------------------------------
        # FOOTER
        # -----------------------------------------------------------------
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('— End of memorandum —').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer2 = doc.add_paragraph()
        footer2.add_run(
            f'Generated by Agora Consulting AI · Matter {matter.reference_number} · '
            f'{datetime.utcnow().strftime("%d %B %Y at %H:%M UTC")}'
        ).italic = True
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # -----------------------------------------------------------------
        # Audit + return
        # -----------------------------------------------------------------
        report_audit = AuditLog(
            matter_id=matter.id,
            user_id=current_user.id,
            action=AuditLogAction.REPORT_GENERATED,
            entity_type="matter",
            entity_id=matter.id,
            description=f"SoF compliance memorandum generated for matter {matter.reference_number}",
            details={
                "reference": matter.reference_number,
                "client": matter.client_name,
                "verdict": verdict_status or 'incomplete',
                "claims": total_claims,
                "passed_claims": passed_claims,
                "documents_reviewed": docs_reviewed,
            },
        )
        sync_db.add(report_audit)
        sync_db.commit()

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        safe_ref = (matter.reference_number or f'matter-{matter.id}').replace('/', '-').replace(' ', '_')
        filename = f"SoF_Compliance_Memorandum_{safe_ref}.docx"
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    finally:
        sync_db.close()


# ==================== WORKFLOW ENGINE ====================

def calculate_completion_percentage(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> int:
    """
    Calculate matter completion percentage based on workflow state and data
    """
    percentage = 0
    
    # Base percentage by status
    status_percentages = {
        MatterStatus.DRAFT: 10,
        MatterStatus.AWAITING_CLIENT: 20,
        MatterStatus.CLIENT_UPLOADING: 40,
        MatterStatus.UNDER_REVIEW: 60,
        MatterStatus.QUERIES_RAISED: 50,  # Step back
        MatterStatus.APPROVED: 90,
        MatterStatus.REJECTED: 100,  # Terminal state
        MatterStatus.COMPLETED: 100,
    }
    
    percentage = status_percentages.get(matter.status, 0)
    
    # Add bonus for SoF assessment completion
    if sof_data:
        if 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            outcome = assessment.get('outcome', {})
            
            # Add 10% if assessment run
            percentage = min(percentage + 10, 100)
            
            # Add 10% more if sufficient
            if outcome.get('status', '').lower() == 'sufficient':
                percentage = min(percentage + 10, 100)
            
            # Add 5% if all claims verified
            evidence_matches = assessment.get('evidence_matches', [])
            if evidence_matches:
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                if all_verified:
                    percentage = min(percentage + 5, 100)
    
    return percentage


def get_next_status_suggestions(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """
    Get smart status suggestions based on current state and data
    """
    current_status = matter.status
    suggestions = []
    
    # Status transition rules with reasons
    transitions = {
        MatterStatus.DRAFT: [
            {'status': MatterStatus.AWAITING_CLIENT, 'reason': 'Ready to request documents from client', 'auto': False},
            {'status': MatterStatus.CLIENT_UPLOADING, 'reason': 'Client portal access granted', 'auto': False},
        ],
        MatterStatus.AWAITING_CLIENT: [
            {'status': MatterStatus.CLIENT_UPLOADING, 'reason': 'Client started uploading documents', 'auto': True},
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'Skip to review (documents received externally)', 'auto': False},
        ],
        MatterStatus.CLIENT_UPLOADING: [
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'All required documents uploaded', 'auto': True},
        ],
        MatterStatus.UNDER_REVIEW: [
            {'status': MatterStatus.QUERIES_RAISED, 'reason': 'Additional information required', 'auto': False},
            {'status': MatterStatus.APPROVED, 'reason': 'Review complete - all requirements met', 'auto': True},
            {'status': MatterStatus.REJECTED, 'reason': 'Cannot proceed with this matter', 'auto': False},
        ],
        MatterStatus.QUERIES_RAISED: [
            {'status': MatterStatus.AWAITING_CLIENT, 'reason': 'Waiting for client response', 'auto': False},
            {'status': MatterStatus.UNDER_REVIEW, 'reason': 'Client responded - resume review', 'auto': False},
        ],
        MatterStatus.APPROVED: [
            {'status': MatterStatus.COMPLETED, 'reason': 'Matter completed successfully', 'auto': False},
            {'status': MatterStatus.QUERIES_RAISED, 'reason': 'New issue identified', 'auto': False},
        ],
        MatterStatus.REJECTED: [
            # Terminal state - no transitions
        ],
        MatterStatus.COMPLETED: [
            # Terminal state - no transitions
        ],
    }
    
    base_suggestions = transitions.get(current_status, [])
    
    # Check if auto-transitions should be suggested based on data
    if sof_data and 'assessment_result' in sof_data:
        assessment = sof_data['assessment_result']
        outcome = assessment.get('outcome', {})
        
        # If UNDER_REVIEW and assessment is SUFFICIENT, suggest APPROVED
        if current_status == MatterStatus.UNDER_REVIEW:
            if outcome.get('status', '').lower() == 'sufficient':
                # Check if all claims fully verified
                evidence_matches = assessment.get('evidence_matches', [])
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                
                # Check no critical alerts
                tr_summary = assessment.get('transaction_review_summary', {})
                critical_alerts = tr_summary.get('critical_alerts', 0)
                
                if all_verified and critical_alerts == 0:
                    # Mark APPROVED as auto-ready
                    for sug in base_suggestions:
                        if sug['status'] == MatterStatus.APPROVED:
                            sug['auto'] = True
                            sug['reason'] = '✅ AUTO: All claims verified, no critical alerts'
        
        # If UNDER_REVIEW and there are outstanding questions, suggest QUERIES_RAISED
        if current_status == MatterStatus.UNDER_REVIEW:
            next_actions = assessment.get('next_actions', {})
            questions = next_actions.get('questions', [])
            if len(questions) > 0:
                for sug in base_suggestions:
                    if sug['status'] == MatterStatus.QUERIES_RAISED:
                        sug['auto'] = True
                        sug['reason'] = f'⚠️ AUTO: {len(questions)} outstanding question(s)'
    
    return base_suggestions


def apply_auto_transitions(matter: Matter, sof_data: Optional[Dict[str, Any]] = None) -> List[str]:
    """
    Apply automatic status transitions based on workflow rules
    """
    transitions_applied = []
    
    # Check if documents have been uploaded (transition to UNDER_REVIEW)
    if matter.status == MatterStatus.CLIENT_UPLOADING and sof_data:
        if 'uploaded_files' in sof_data and len(sof_data.get('uploaded_files', [])) > 0:
            # Check if client_info and bank_statements uploaded
            has_client_info = any(f.get('category') == 'client_info' for f in sof_data.get('uploaded_files', []))
            has_bank = any(f.get('category') == 'bank_statement' for f in sof_data.get('uploaded_files', []))
            
            if has_client_info and has_bank:
                matter.status = MatterStatus.UNDER_REVIEW
                transitions_applied.append(f"CLIENT_UPLOADING → UNDER_REVIEW (documents uploaded)")
    
    # Check if assessment complete and sufficient (transition to APPROVED)
    if matter.status == MatterStatus.UNDER_REVIEW and sof_data:
        if 'assessment_result' in sof_data:
            assessment = sof_data['assessment_result']
            outcome = assessment.get('outcome', {})
            
            if outcome.get('status', '').lower() == 'sufficient':
                # Check all claims verified
                evidence_matches = assessment.get('evidence_matches', [])
                all_verified = all(
                    e.get('verified', False) and e.get('document_verified', False) 
                    for e in evidence_matches
                )
                
                # Check no critical alerts
                tr_summary = assessment.get('transaction_review_summary', {})
                critical_alerts = tr_summary.get('critical_alerts', 0)
                
                if all_verified and critical_alerts == 0:
                    matter.status = MatterStatus.APPROVED
                    transitions_applied.append(f"UNDER_REVIEW → APPROVED (all verified, no critical alerts)")
    
    return transitions_applied


@router.patch("/matters/{matter_id}/status")
async def update_matter_status(
    matter_id: int,
    request: StatusUpdateRequest,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """
    Update matter status with workflow validation and smart transitions
    
    Features:
    - Workflow validation (prevent invalid transitions)
    - Automatic status transitions based on data
    - Completion percentage calculation
    - Status change audit trail
    """
    import os

    # Get shared sync DB session
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    
    try:
        # Get matter
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data from DB
        sof_data = _safe_load_storage_single(sync_db, matter_id)
        
        # Store previous status
        previous_status = matter.status
        
        # Validate new status
        try:
            new_status = MatterStatus(request.new_status)
        except ValueError:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid status: {request.new_status}. Valid statuses: {[s.value for s in MatterStatus]}"
            )
        
        # Check if transition is allowed
        suggestions = get_next_status_suggestions(matter, sof_data)
        allowed_statuses = [s['status'] for s in suggestions]
        
        # Allow same status (for refresh) or any suggested status
        if new_status != previous_status and new_status not in allowed_statuses:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid transition from {previous_status.value} to {new_status.value}. Allowed: {[s.value for s in allowed_statuses]}"
            )

        # Require reason for REJECTED status
        if new_status == MatterStatus.REJECTED and not request.reason:
            raise HTTPException(
                status_code=400,
                detail="A reason is required when rejecting a matter."
            )

        # Apply status change
        matter.status = new_status
        matter.updated_at = datetime.now()

        # Apply automatic transitions if requested
        auto_transitions = []
        if request.auto_transition:
            auto_transitions = apply_auto_transitions(matter, sof_data)

        # Record status history
        if new_status != previous_status:
            history = MatterStatusHistory(
                matter_id=matter.id,
                old_status=previous_status.value,
                new_status=matter.status.value,
                changed_by=current_user.id,
                reason=request.reason,
            )
            sync_db.add(history)

            # Audit log for status change
            audit = AuditLog(
                matter_id=matter.id,
                user_id=current_user.id,
                action=AuditLogAction.STATUS_CHANGED,
                entity_type="matter",
                entity_id=matter.id,
                description=f"Status changed from {previous_status.value} to {matter.status.value}",
                details={
                    "previous_status": previous_status.value,
                    "new_status": matter.status.value,
                    "reason": request.reason,
                },
            )
            sync_db.add(audit)

        # Record any auto-transitions as well
        for transition_desc in auto_transitions:
            auto_history = MatterStatusHistory(
                matter_id=matter.id,
                old_status=previous_status.value,
                new_status=matter.status.value,
                changed_by=current_user.id,
                reason=f"Auto-transition: {transition_desc}",
            )
            sync_db.add(auto_history)

        # Calculate new completion percentage
        completion_percentage = calculate_completion_percentage(matter, sof_data)

        # Commit changes
        sync_db.commit()
        sync_db.refresh(matter)
        
        # Build response
        message = f"Status updated from {previous_status.value} to {matter.status.value}"
        if auto_transitions:
            message += f". Auto-transitions: {', '.join(auto_transitions)}"
        
        return StatusUpdateResponse(
            success=True,
            previous_status=previous_status.value,
            new_status=matter.status.value,
            completion_percentage=completion_percentage,
            message=message,
            auto_transitions_applied=auto_transitions
        )
    
    finally:
        sync_db.close()


@router.get("/matters/{matter_id}/status/suggestions")
async def get_status_suggestions(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_db)
):
    """
    Get smart status transition suggestions for a matter
    
    Returns suggested next statuses with reasons and auto-transition flags
    """
    import os

    # Get shared sync DB session
    SessionLocal = get_sync_session()
    sync_db = SessionLocal()
    
    try:
        # Get matter
        matter = sync_db.query(Matter).filter(Matter.id == matter_id).first()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")
        
        # Load SoF assessment data from DB
        sof_data = _safe_load_storage_single(sync_db, matter_id)
        
        # Get suggestions
        suggestions = get_next_status_suggestions(matter, sof_data)
        
        # Calculate current completion percentage
        completion_percentage = calculate_completion_percentage(matter, sof_data)
        
        # Check if auto-transitions available
        auto_transitions = apply_auto_transitions(matter, sof_data)
        
        return {
            'current_status': matter.status.value,
            'completion_percentage': completion_percentage,
            'suggestions': [
                {
                    'status': s['status'].value,
                    'reason': s['reason'],
                    'auto_recommended': s.get('auto', False)
                }
                for s in suggestions
            ],
            'auto_transitions_available': len(auto_transitions) > 0,
            'auto_transition_preview': auto_transitions
        }
    
    finally:
        sync_db.close()
