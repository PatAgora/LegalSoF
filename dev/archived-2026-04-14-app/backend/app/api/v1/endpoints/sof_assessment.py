"""
Source of Funds Assessment API Endpoints

100% LOCAL - No external API calls
Handles file uploads, processing, and SoF assessment results
"""
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, Form
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone
import json

from app.db.session import get_db, get_sync_db, get_sync_engine, get_sync_session
from app.api.dependencies.auth import require_analyst, require_admin
from app.models.user import User
from app.services.file_processor import file_processor
from app.services.sof_assessment_engine import SoFAssessmentEngine
from app.services.statement_validation_pipeline import StatementValidationPipeline, ValidationResult as PipelineResult
from app.models import Matter
from app.models.statement_validation import (
    StatementValidation, StatementValidationFlag, StatementValidationTransaction,
    ValidationStatus, FlagSeverity,
)
from app.services.document_verification_pipeline import document_verification_pipeline, VerificationResult as DocVerResult
from app.services.statement_validation_pipeline import StatementValidationPipeline, ValidationResult as StmtValResult
from app.services.cross_document_corroborator import corroborate as corroborate_verification
from app.models.document_verification import (
    DocumentVerification, DocumentVerificationFlag, DocumentVerificationTransaction,
    VerificationVerdict,
)
from app.models.transaction import TransactionConfig
from app.models.audit import AuditLog, AuditLogAction
from app.models.notification import Notification
from app.models.user import UserRole

router = APIRouter()


import os
from datetime import datetime as dt

# Helper to parse dates in multiple formats - used throughout the module
def parse_date_flexible(date_str: str):
    """Parse date string in multiple formats, returns datetime object or None"""
    if not date_str:
        return None
    # Try UK format first (dd/mm/yyyy) since that's our standard
    for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%d %b %Y', '%d %B %Y']:
        try:
            return dt.strptime(str(date_str).strip(), fmt)
        except:
            continue
    return None

# Format date to UK format dd/mm/yyyy
def format_date_uk(date_obj) -> str:
    """Format datetime object to UK date string"""
    if date_obj:
        return date_obj.strftime('%d/%m/%Y')
    return 'Unknown'

# ---------------------------------------------------------------------------
# PostgreSQL-backed assessment storage helpers
# (Replaces the old /tmp JSON file + in-memory dict approach)
# ---------------------------------------------------------------------------
from app.models.assessment_storage import AssessmentStorage


def _db_load_storage(db: Session, matter_id: int) -> Optional[Dict[str, Any]]:
    """Load assessment data for a single matter from the database."""
    row = db.query(AssessmentStorage).filter(AssessmentStorage.matter_id == matter_id).first()
    if row and row.data:
        return row.data
    return None


def _db_save_storage(db: Session, matter_id: int, data: Dict[str, Any]):
    """Upsert assessment data for a matter into the database."""
    # Ensure all values are JSON-serialisable (convert datetimes, Decimals, etc.)
    import json as _json
    clean = _json.loads(_json.dumps(data, default=str))

    row = db.query(AssessmentStorage).filter(AssessmentStorage.matter_id == matter_id).first()
    if row:
        row.data = clean
    else:
        row = AssessmentStorage(matter_id=matter_id, data=clean)
        db.add(row)
    db.commit()


def _db_delete_storage(db: Session, matter_id: int):
    """Delete assessment data for a matter from the database."""
    db.query(AssessmentStorage).filter(AssessmentStorage.matter_id == matter_id).delete()
    db.commit()


def _db_load_all_storage(db: Session) -> Dict[int, Dict[str, Any]]:
    """Load assessment data for ALL matters (used by matters.py list endpoint)."""
    rows = db.query(AssessmentStorage).all()
    return {row.matter_id: row.data for row in rows if row.data}


def run_automated_funds_lineage(
    target_transaction: Dict[str, Any],
    bank_statements: List[Dict[str, Any]],
    savings_claim: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Run automated funds lineage analysis for a savings claim.
    Uses RECURSIVE TREE-BUILDING approach - starts from target and traces backward.
    
    ENHANCED: Now detects statement gaps and continues tracing through savings
    transactions even when current account statements don't cover the full period.
    
    Returns:
        Dict with lineage_tree, summary, unresolved_items, external_origins, statement_gaps
    """
    from datetime import datetime as dt
    import re
    
    target_amount = float(target_transaction.get('amount', 0))
    target_date = target_transaction.get('date', '')
    target_account = target_transaction.get('account_id', '') or target_transaction.get('account', '')
    target_desc = target_transaction.get('description', '')
    target_id = target_transaction.get('id') or target_transaction.get('transaction_id') or f"TXN-TARGET-1"
    
    print(f"\n=== AUTOMATED FUNDS LINEAGE (RECURSIVE) ===")
    print(f"  Target: £{target_amount:,.2f} from {target_desc[:50]}...")
    print(f"  Target account: {target_account}")
    print(f"  Target date: {target_date}")
    
    # Group transactions by account
    accounts: Dict[str, List[Dict]] = {}
    for idx, txn in enumerate(bank_statements):
        acc_id = txn.get('account_id', '') or txn.get('account', 'Unknown')
        if acc_id not in accounts:
            accounts[acc_id] = []
        # Ensure each transaction has an ID
        if not txn.get('id') and not txn.get('transaction_id'):
            txn['id'] = f"TXN-{acc_id[:10]}-{idx + 1}"
        accounts[acc_id].append(txn)
    
    print(f"  Found {len(accounts)} account(s): {list(accounts.keys())}")
    
    # Helper to parse dates in multiple formats
    def parse_date_flexible(date_str: str):
        """Parse date string in multiple formats"""
        if not date_str:
            return None
        for fmt in ['%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y']:
            try:
                return dt.strptime(date_str, fmt)
            except:
                continue
        return None
    
    # Calculate date range for each account (for gap detection)
    account_date_ranges: Dict[str, Dict] = {}
    for acc_id, txns in accounts.items():
        dates = []
        for txn in txns:
            d = parse_date_flexible(txn.get('date', ''))
            if d:
                dates.append(d)
        if dates:
            account_date_ranges[acc_id] = {
                'earliest': min(dates),
                'latest': max(dates),
                'earliest_str': min(dates).strftime('%d/%m/%Y'),
                'latest_str': max(dates).strftime('%d/%m/%Y'),
                'transaction_count': len(txns)
            }
    
    print(f"\n  === ACCOUNT STATEMENT COVERAGE ===")
    for acc_id, range_info in account_date_ranges.items():
        print(f"    {acc_id}: {range_info['earliest_str']} to {range_info['latest_str']} ({range_info['transaction_count']} txns)")
    print(f"  ===================================\n")
    
    # Classify accounts (savings vs current vs unknown).
    # P1.4 — surface ambiguity. Previously this function always returned
    # 'savings' or 'current' even when it was guessing from weak signals
    # (e.g. just "more debits than credits"). Now we return a tuple of
    # (classification, confidence) so we can flag accounts that the
    # classifier wasn't sure about. The lineage walker treats 'unknown'
    # as both — i.e. it'll still look for transfers — but the result
    # surfaces them for human review.
    def classify_account(acc_id: str, txns: List[Dict]) -> tuple:
        acc_lower = acc_id.lower()
        # First check: filename / account-name keywords give "high" confidence.
        if any(kw in acc_lower for kw in ['sav', 'isa', 'deposit', 'reserve', 'easy access']):
            return 'savings', 'high'
        if any(kw in acc_lower for kw in ['current', 'checking', 'everyday', 'cheque']):
            return 'current', 'high'
        # Transaction-pattern signal: salary keyword strongly suggests current.
        has_salary = any(
            kw in (t.get('description', '') or '').lower()
            for t in txns
            for kw in ('salary', 'payroll', 'wages')
        )
        if has_salary:
            return 'current', 'medium'
        # Fallback heuristic: more debits than credits => probably current.
        # Flag as low confidence so the result surfaces this guess.
        credits = [t for t in txns if t.get('direction') == 'credit']
        debits = [t for t in txns if t.get('direction') == 'debit']
        if len(debits) > len(credits) * 1.5:
            return 'current', 'low'
        if len(credits) > len(debits) * 1.5:
            return 'savings', 'low'
        # Genuinely ambiguous — equal-ish flow, no keywords.
        return 'unknown', 'none'

    classification_results = {
        acc_id: classify_account(acc_id, txns)
        for acc_id, txns in accounts.items()
    }
    account_types = {acc_id: cls for acc_id, (cls, _conf) in classification_results.items()}
    ambiguous_accounts = [
        {
            'account_id': acc_id,
            'classified_as': cls,
            'confidence': conf,
            'transaction_count': len(accounts.get(acc_id, [])),
        }
        for acc_id, (cls, conf) in classification_results.items()
        if conf in ('low', 'none')
    ]
    savings_accounts = [acc for acc, typ in account_types.items() if typ == 'savings']
    current_accounts = [acc for acc, typ in account_types.items() if typ == 'current']

    print(f"  Account types: {account_types}")
    if ambiguous_accounts:
        print(f"  ⚠ Ambiguous classification: {ambiguous_accounts}")
    
    # Keywords for identifying external origins (legitimate sources)
    external_origin_keywords = [
        'salary', 'wages', 'payroll', 'income', 'employer',
        'interest', 'dividend', 'pension', 'annuity',
        'refund', 'cashback', 'tax refund', 'hmrc', 'rebate',
        'rental', 'rent received', 'tenant', 'bacs'
    ]
    
    # Keywords for suspicious/unverified sources requiring evidence
    suspicious_keywords = [
        'unknown', 'cash deposit', 'anonymous', 'unidentified'
    ]
    
    # Keywords indicating inter-account transfer
    transfer_keywords = [
        'standing order', 'transfer', 'tfr', 'int xfer', 'internal',
        'faster payment', 'fp from', 'fp to', 'bacs', 'chaps',
        'move', 'sweep', 's/o'
    ]
    
    # Counters for summary
    traced_amount = 0
    matched_transfers = 0
    external_origins = []
    unresolved_items = []
    statement_gap_items = []  # NEW: Items where we need more statements
    all_dates = []
    node_counter = [0]  # Use list for mutable counter in nested function
    
    def get_txn_id(txn: Dict, fallback_idx: int = 0) -> str:
        """Get or generate a transaction ID"""
        return txn.get('id') or txn.get('transaction_id') or f"TXN-{fallback_idx + 1}"
    
    def identify_external_origin(txn: Dict) -> tuple:
        """Recognise a credit as a legitimate external origin (not
        suspicious, no further evidence needed).

        Categories covered, ordered to put high-confidence patterns
        first (so e.g. an "insurance settlement" credit isn't pre-
        empted by a vaguer keyword):
          - Salary / wages / payroll / employment income
          - Pension / annuity
          - Dividend / interest
          - HMRC / tax refund / rebate
          - Rental income
          - Inheritance / probate / estate (P1.2)
          - Insurance claim / settlement / payout (P1.2)
          - Asset sale: property completion, vehicle sale, share sale (P1.2)
          - Loan / mortgage disbursement / drawdown (P1.2)
          - Government grant / benefit / DWP (P1.2)
          - Gift / wedding gift / family transfer (P1.2)
        """
        desc = (txn.get('description', '') or '').lower()
        direction = txn.get('direction')

        if any(kw in desc for kw in ['salary', 'wages', 'payroll']):
            return True, 'Salary / Employment Income'
        if any(kw in desc for kw in ['pension', 'annuity']):
            return True, 'Pension Payment'
        if 'dividend' in desc:
            return True, 'Dividend Income'
        if 'interest' in desc and 'transfer' not in desc:
            return True, 'Interest Payment'
        if 'hmrc' in desc or 'tax refund' in desc or 'tax rebate' in desc:
            return True, 'Tax Refund / HMRC'
        if 'rent' in desc and direction == 'credit':
            return True, 'Rental Income'
        if 'bacs' in desc and any(kw in desc for kw in ['employer', 'ltd', 'plc', 'inc']):
            return True, 'Employment Income'

        # ------------------------------------------------------------------
        # P1.2 broadened allow-list — these patterns previously fell
        # through to "requires evidence" even when the credit was
        # clearly from a recognised legitimate source.
        # ------------------------------------------------------------------
        if any(kw in desc for kw in ['inheritance', 'probate', 'estate of', 'executor']):
            return True, 'Inheritance / Probate'
        if any(kw in desc for kw in ['insurance', 'policy payout', 'claim settlement', 'aviva', 'legal & general', 'lv=', 'direct line']):
            # 'insurance' alone is too loose if it's e.g. an insurance
            # premium debit. Require the txn to be a credit.
            if direction == 'credit':
                return True, 'Insurance Settlement'
        if any(kw in desc for kw in ['property sale', 'house sale', 'completion proceeds', 'sale proceeds', 'conveyancing']):
            return True, 'Property Sale Proceeds'
        if any(kw in desc for kw in ['vehicle sale', 'car sale', 'motor sale']):
            return True, 'Vehicle Sale Proceeds'
        if any(kw in desc for kw in ['share sale', 'sale of shares', 'stock sale', 'equity sale']):
            return True, 'Share / Equity Sale'
        if any(kw in desc for kw in ['loan drawdown', 'mortgage drawdown', 'loan advance', 'mortgage advance']):
            return True, 'Loan / Mortgage Disbursement'
        if any(kw in desc for kw in ['dwp', 'universal credit', 'state pension', 'jobseeker', 'esa payment', 'pip ']):
            return True, 'Government Benefit'
        if 'grant' in desc and direction == 'credit':
            return True, 'Government / Body Grant'
        if 'compensation' in desc and direction == 'credit':
            return True, 'Compensation Payment'
        if any(kw in desc for kw in ['gift from', 'birthday gift', 'wedding gift']):
            return True, 'Gift'

        return False, ''
    
    def is_likely_transfer(txn: Dict) -> bool:
        """Check if transaction description suggests it's a transfer"""
        desc = (txn.get('description', '') or '').lower()
        return any(kw in desc for kw in transfer_keywords)
    
    # Rough GBP-anchored FX rates for AML triage when statements span
    # multiple currencies. These are not real-time; they exist so the
    # matcher can spot inter-currency transfers as related rather than
    # marking them all as untraced. Reviewer still has to verify the
    # actual rate on the transaction date.
    _FX_TO_GBP = {
        'GBP': 1.0,    'USD': 0.79,   'EUR': 0.85,   'JPY': 0.0053,
        'AUD': 0.52,   'CAD': 0.58,   'CHF': 0.88,   'INR': 0.0094,
        'CNY': 0.11,   'HKD': 0.10,   'SGD': 0.59,   'NZD': 0.47,
        'AED': 0.21,   'ZAR': 0.041,
    }

    def _to_gbp(amount: float, currency: str) -> Optional[float]:
        rate = _FX_TO_GBP.get((currency or '').upper().strip())
        if rate is None:
            return None
        return amount * rate

    def _share_digit_run(a: str, b: str, min_len: int = 4) -> bool:
        """Reference-number tiebreaker — do both descriptions share
        any digit run of length >= min_len? Useful when two debits sit
        within tolerance and we need to pick the one that actually
        corresponds to the credit."""
        import re as _re
        a_runs = set(_re.findall(rf'\d{{{min_len},}}', a or ''))
        b_runs = set(_re.findall(rf'\d{{{min_len},}}', b or ''))
        return bool(a_runs & b_runs)

    def find_matching_debit(credit_txn: Dict, source_account_txns: List[Dict]) -> Optional[Dict]:
        """Find a debit on another account that matches this credit.

        Matching rules (P1.1, P2.5, P2.6):
          - Amount within max(£1, 0.5% of credit amount) — percentage
            scales for large transfers that lose a few pounds to fees.
          - Date within 3 days for same-currency; 5 days for cross-
            currency (international clearing often slower).
          - When currencies differ, both sides are FX-converted to GBP
            (using the rough table above) and matched at ±2% tolerance.
            Match is annotated with `_fx_converted` so the UI can flag
            it for verification.
          - When multiple candidate debits fit the tolerance window,
            the one whose description shares a digit-run with the
            credit (account / reference number) wins. Otherwise the
            closest by amount, then by date.
        """
        credit_date = parse_date_flexible(credit_txn.get('date', ''))
        if not credit_date:
            return None

        credit_amount = abs(float(credit_txn.get('amount', 0)))
        credit_currency = (credit_txn.get('currency') or 'GBP').upper().strip()
        credit_desc = credit_txn.get('description', '') or ''

        same_cur_tol = max(1.0, credit_amount * 0.005)

        # Pre-compute GBP value for FX path
        credit_gbp = _to_gbp(credit_amount, credit_currency)

        candidates: List[tuple] = []  # (txn, score, fx_converted)

        for txn in source_account_txns:
            if txn.get('direction') != 'debit':
                continue

            txn_amount = abs(float(txn.get('amount', 0)))
            txn_currency = (txn.get('currency') or 'GBP').upper().strip()
            txn_date = parse_date_flexible(txn.get('date', ''))
            if not txn_date:
                continue
            days_diff = abs((txn_date - credit_date).days)

            same_currency = (credit_currency == txn_currency)
            if same_currency:
                if days_diff > 3:
                    continue
                amount_diff = abs(txn_amount - credit_amount)
                if amount_diff > same_cur_tol:
                    continue
                # Score: amount mismatch + day penalty. Lower = better.
                score = amount_diff * 100 + days_diff
                fx_converted = False
            else:
                # Cross-currency path
                if days_diff > 5:
                    continue
                txn_gbp = _to_gbp(txn_amount, txn_currency)
                if credit_gbp is None or txn_gbp is None:
                    continue
                ratio_diff = abs(txn_gbp - credit_gbp) / max(credit_gbp, 1.0)
                if ratio_diff > 0.02:  # 2% — covers normal FX spread
                    continue
                # Heavier base score so a same-currency match beats an
                # FX-converted one when both are available.
                score = 1000 + ratio_diff * 10000 + days_diff
                fx_converted = True

            # Reference-number tiebreaker — if both descriptions share
            # a 4+ digit run (account / payment reference), strongly
            # prefer this candidate.
            if _share_digit_run(credit_desc, txn.get('description', '') or ''):
                score -= 500

            candidates.append((txn, score, fx_converted))

        if not candidates:
            return None

        candidates.sort(key=lambda c: c[1])
        best_txn, _, fx_converted = candidates[0]
        if fx_converted:
            # Annotate so the lineage node can render "FX-matched" badge.
            best_txn['_fx_converted'] = True
        return best_txn
    
    def check_statement_gap(txn_date_str: str, source_acc_id: str) -> Optional[str]:
        """
        Check if a transaction date falls BEFORE another account's statement coverage.
        Returns the account ID that needs more statements, or None if no gap.
        """
        txn_date = parse_date_flexible(txn_date_str)
        if not txn_date:
            return None
        
        # Check if this date is before any other account's earliest date
        for acc_id, range_info in account_date_ranges.items():
            if acc_id == source_acc_id:
                continue
            # If transaction is before this account's coverage, there's a gap
            if txn_date < range_info['earliest']:
                return acc_id
        return None
    
    def find_funding_credits(debit_txn: Dict, account_txns: List[Dict], exclude_id: str = None) -> List[Dict]:
        """Find credits that could have funded a debit.

        P2.7 — Chronological validation. We REQUIRE every candidate
        credit's date to fall on or before the debit's date. A credit
        that arrives AFTER the debit cannot have funded it; including
        such credits in the funding chain is a false-positive that
        previously slipped through whenever the dataset was unsorted.
        The explicit `txn_date <= debit_date` filter rejects them.
        """
        debit_date = parse_date_flexible(debit_txn.get('date', ''))
        if not debit_date:
            return []

        debit_amount = abs(float(debit_txn.get('amount', 0)))

        # Collect candidate credits that PRE-DATE the debit. Each
        # tuple is (txn, parsed_date) so we can sort deterministically.
        prior_credits: List[tuple] = []
        for txn in account_txns:
            if txn.get('direction') != 'credit':
                continue
            txn_id = get_txn_id(txn)
            if exclude_id and txn_id == exclude_id:
                continue
            txn_date = parse_date_flexible(txn.get('date', ''))
            if not txn_date:
                continue
            # *** Chronology guard *** — a credit AFTER the debit cannot
            # have funded it.
            if txn_date > debit_date:
                continue
            days_diff = (debit_date - txn_date).days
            if days_diff <= 365:  # 1 year lookback for funds lineage
                prior_credits.append((txn, txn_date))

        # Sort by date descending (most recent first) — we want the
        # NEWEST funding sources, because those are most likely to be
        # the ones that actually paid for the debit.
        prior_credits.sort(key=lambda x: x[1], reverse=True)

        # Greedily select credits up to the debit amount.
        funding_credits: List[Dict] = []
        running_total = 0.0
        for credit, _ in prior_credits:
            if running_total >= debit_amount:
                break
            funding_credits.append(credit)
            running_total += abs(float(credit.get('amount', 0)))

        return funding_credits
    
    def build_lineage_node(txn: Dict, level: int, visited: set) -> Dict:
        """Recursively build lineage tree from a transaction"""
        nonlocal traced_amount, matched_transfers, external_origins, unresolved_items, statement_gap_items, all_dates
        
        node_counter[0] += 1
        node_id = f"lineage-{node_counter[0]}"
        txn_id = get_txn_id(txn, node_counter[0])
        txn_amount = abs(float(txn.get('amount', 0)))
        txn_date = txn.get('date', '')
        txn_account = txn.get('account_id', '') or txn.get('account', 'Unknown')
        
        # Track dates for accumulation period
        if txn_date:
            parsed_date = parse_date_flexible(txn_date)
            if parsed_date:
                all_dates.append(parsed_date)
        
        # Prevent infinite loops.
        # P1.3 — surface circular references clearly so the analyst
        # can investigate. Round-tripping money through the same
        # accounts is a classic layering technique, so we record it
        # with a distinct severity rather than burying it as a
        # generic unresolved row.
        is_circular = txn_id in visited
        is_too_deep = level > 50
        if is_circular or is_too_deep:
            reason = 'circular_reference' if is_circular else 'max_depth_reached'
            note = (
                'Circular reference — this transaction was already encountered '
                'further up the lineage tree, which can indicate round-tripping '
                'across accounts.'
                if is_circular else
                f'Maximum trace depth ({level}) reached without finding a clear origin.'
            )
            unresolved_items.append({
                'id': txn_id,
                'transaction_id': txn_id,
                'date': txn_date,
                'amount': txn_amount,
                'description': txn.get('description', ''),
                'account': txn_account,
                'reason': reason,
                'severity': 'high' if is_circular else 'medium',
                'message': note,
            })
            return {
                'id': node_id,
                'level': level,
                'transaction': txn,
                'amount': txn_amount,
                'date': txn_date,
                'description': txn.get('description', ''),
                'account': txn_account,
                'source_account': 'Circular Reference' if is_circular else 'Trace Depth Exceeded',
                'destination_account': txn_account,
                'match_type': 'circular' if is_circular else 'requires_evidence',
                'notes': note,
                'children': []
            }
        
        visited.add(txn_id)
        
        # Check if this is an external origin
        is_external, source_type = identify_external_origin(txn)
        if is_external:
            traced_amount += txn_amount
            external_origins.append({
                'date': txn_date,
                'amount': txn_amount,
                'description': txn.get('description', ''),
                'source_type': source_type
            })
            print(f"    {'  ' * level}✓ EXTERNAL: £{txn_amount:,.2f} - {source_type}")
            return {
                'id': node_id,
                'level': level,
                'transaction': txn,
                'amount': txn_amount,
                'date': txn_date,
                'description': txn.get('description', ''),
                'account': txn_account,
                'source_account': source_type,  # External source (salary, interest, etc.)
                'destination_account': txn_account,  # Where it landed
                'match_type': 'external_origin',
                'notes': f'External origin: {source_type}',
                'children': []
            }
        
        # For credits, look for matching debit in other accounts
        if txn.get('direction') == 'credit':
            for other_acc_id, other_acc_txns in accounts.items():
                if other_acc_id == txn_account:
                    continue
                
                matched_debit = find_matching_debit(txn, other_acc_txns)
                if matched_debit:
                    matched_transfers += 1
                    print(f"    {'  ' * level}✓ MATCHED: £{txn_amount:,.2f} from {other_acc_id}")
                    
                    # Find credits that funded this debit
                    funding_credits = find_funding_credits(matched_debit, other_acc_txns, txn_id)
                    
                    # Recursively trace each funding credit
                    children = []
                    for credit in funding_credits:
                        child_node = build_lineage_node(credit, level + 1, visited.copy())
                        children.append(child_node)
                    
                    return {
                        'id': node_id,
                        'level': level,
                        'transaction': txn,
                        'matched_transaction': matched_debit,
                        'amount': txn_amount,
                        'date': txn_date,
                        'description': txn.get('description', ''),
                        'account': txn_account,
                        'source_account': other_acc_id,  # Where the money came from
                        'destination_account': txn_account,  # Where it landed
                        'match_type': 'matched',
                        'notes': f'Verified transfer from {other_acc_id}',
                        'children': children
                    }
            
            # No match found - check if it's due to a statement gap
            if is_likely_transfer(txn):
                gap_account = check_statement_gap(txn_date, txn_account)
                if gap_account:
                    # This is a statement gap - we need more statements
                    gap_range = account_date_ranges.get(gap_account, {})
                    gap_info = {
                        'id': txn_id,
                        'transaction_id': txn_id,
                        'date': txn_date,
                        'amount': txn_amount,
                        'description': txn.get('description', ''),
                        'account': txn_account,
                        'gap_account': gap_account,
                        'gap_account_earliest': gap_range.get('earliest_str', 'Unknown'),
                        'reason': 'statement_gap'
                    }
                    statement_gap_items.append(gap_info)
                    print(f"    {'  ' * level}📋 STATEMENT GAP: £{txn_amount:,.2f} - need {gap_account} statements before {gap_range.get('earliest_str', 'Unknown')}")
                    
                    return {
                        'id': node_id,
                        'level': level,
                        'transaction': txn,
                        'amount': txn_amount,
                        'date': txn_date,
                        'description': txn.get('description', ''),
                        'account': txn_account,
                        'source_account': gap_account,  # The account we need statements for
                        'destination_account': txn_account,  # Where it landed
                        'match_type': 'statement_gap',
                        'notes': f'Statement gap: Need {gap_account} statements from before {gap_range.get("earliest_str", "Unknown")}',
                        'gap_account': gap_account,
                        'gap_earliest': gap_range.get('earliest_str', 'Unknown'),
                        'children': []
                    }
        
        # Check if it's a small regular transaction (assume legitimate)
        desc = (txn.get('description', '') or '').lower()
        is_suspicious = any(kw in desc for kw in suspicious_keywords)
        is_large = txn_amount > 5000
        
        if not is_suspicious and not is_large:
            traced_amount += txn_amount
            print(f"    {'  ' * level}✓ REGULAR: £{txn_amount:,.2f} - small amount")
            return {
                'id': node_id,
                'level': level,
                'transaction': txn,
                'amount': txn_amount,
                'date': txn_date,
                'description': txn.get('description', ''),
                'account': txn_account,
                'source_account': 'Regular Income',  # Generic for small amounts
                'destination_account': txn_account,
                'match_type': 'external_origin',
                'notes': 'Regular transaction (small amount)',
                'children': []
            }
        
        # No match found - requires evidence
        unresolved_items.append({
            'id': txn_id,
            'transaction_id': txn_id,
            'date': txn_date,
            'amount': txn_amount,
            'description': txn.get('description', ''),
            'account': txn_account,
            'reason': 'unverified_source'
        })
        print(f"    {'  ' * level}⚠ UNRESOLVED: £{txn_amount:,.2f} - requires evidence (ID: {txn_id})")
        
        return {
            'id': node_id,
            'level': level,
            'transaction': txn,
            'amount': txn_amount,
            'date': txn_date,
            'description': txn.get('description', ''),
            'account': txn_account,
            'source_account': 'Unknown Source',  # Unknown - needs verification
            'destination_account': txn_account,
            'match_type': 'requires_evidence',
            'notes': 'Source requires verification - evidence needed',
            'children': []
        }
    
    # Build the lineage tree starting from the target transaction
    print(f"\n  Building lineage tree from target...")
    root_node = build_lineage_node(target_transaction, 0, set())
    lineage_tree = [root_node]
    
    # Calculate accumulation period from dates in the lineage tree
    accumulation_days = 0
    if all_dates:
        accumulation_days = (max(all_dates) - min(all_dates)).days
        if accumulation_days == 0 and len(all_dates) > 0:
            accumulation_days = 1
    
    # Calculate traced percentage
    traced_percentage = round((traced_amount / target_amount * 100)) if target_amount > 0 else 0
    
    # Build statement gap summary
    statement_gaps_summary = []
    if statement_gap_items:
        # Group by gap_account
        gap_by_account: Dict[str, List] = {}
        for item in statement_gap_items:
            acc = item.get('gap_account', 'Unknown')
            if acc not in gap_by_account:
                gap_by_account[acc] = []
            gap_by_account[acc].append(item)
        
        for acc, items in gap_by_account.items():
            # Find the earliest date we need statements for
            earliest_needed = min(item['date'] for item in items)
            acc_range = account_date_ranges.get(acc, {})
            
            statement_gaps_summary.append({
                'account': acc,
                'current_coverage_from': acc_range.get('earliest_str', 'Unknown'),
                'current_coverage_to': acc_range.get('latest_str', 'Unknown'),
                'statements_needed_from': earliest_needed,
                'transactions_affected': len(items),
                'total_amount_affected': sum(item['amount'] for item in items)
            })
    
    print(f"\n  === LINEAGE SUMMARY ===")
    print(f"  Total: £{target_amount:,.2f}")
    print(f"  Traced: £{traced_amount:,.2f} ({traced_percentage}%)")
    print(f"  Matched transfers: {matched_transfers}")
    print(f"  External origins: {len(external_origins)}")
    print(f"  Unresolved: {len(unresolved_items)}")
    print(f"  Statement gaps: {len(statement_gap_items)}")
    print(f"  Accumulation period: {accumulation_days} days")
    if statement_gaps_summary:
        print(f"\n  === STATEMENT GAPS DETECTED ===")
        for gap in statement_gaps_summary:
            print(f"    {gap['account']}: Current coverage {gap['current_coverage_from']} to {gap['current_coverage_to']}")
            print(f"      Need statements from: {gap['statements_needed_from']}")
            print(f"      Transactions affected: {gap['transactions_affected']} (£{gap['total_amount_affected']:,.2f})")
    print(f"  ========================\n")
    
    summary = {
        'totalAmount': target_amount,
        'tracedAmount': traced_amount,
        'untracedAmount': target_amount - traced_amount,
        'matchedTransfers': matched_transfers,
        'externalOrigins': len(external_origins),
        'requiresEvidence': len(unresolved_items),
        'statementGapsCount': len(statement_gap_items),
        'accumulationPeriodDays': accumulation_days,
        'traced_percentage': traced_percentage,
        'accountCoverage': account_date_ranges,
        'statementGaps': statement_gaps_summary,
        'ambiguousAccounts': len(ambiguous_accounts),
        'circularReferences': sum(1 for it in unresolved_items if it.get('reason') == 'circular_reference'),
    }

    return {
        'target_transaction': target_transaction,
        'summary': summary,
        'lineage_tree': lineage_tree,
        'unresolved_items': unresolved_items,
        'statement_gap_items': statement_gap_items,
        'external_origins': external_origins,
        'ambiguous_accounts': ambiguous_accounts,
        'traced_percentage': traced_percentage,
        'run_at': datetime.now(timezone.utc).isoformat(),
        'run_by': 'auto'
    }


@router.post("/matters/{matter_id}/sof-assessment/upload")
async def upload_sof_files(
    matter_id: int,
    file: UploadFile = File(...),
    file_category: str = Form(...),  # 'client_info', 'bank_statement', 'supporting_doc'
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Upload and process a file for SoF assessment
    
    File categories:
    - client_info: JSON with client details, purchase info, SoF explanation
    - bank_statement: CSV or PDF bank statement
    - supporting_doc: PDF supporting document
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    print(f"\n📤 Uploading file for matter {matter_id}: {file.filename} ({file_category})")
    
    # Load existing storage for this matter from DB (or create new)
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        print(f"   Creating new storage for matter {matter_id}")
        storage = {
            "client_info": None,
            "bank_statements": [],
            "supporting_docs": [],
            "uploaded_files": [],
            "status": "pending",
            "last_updated": None
        }
    else:
        print(f"   Using existing storage for matter {matter_id}")

    # Determine file type from extension and MIME
    file_ext = file.filename.split('.')[-1].lower() if file.filename else ''
    
    if file_category == 'client_info':
        # Accept JSON (structured) plus free-text formats — PDF, Word,
        # CSV. The file_processor turns the free-text content into the
        # same shape the frontend manual form uses via regex
        # extraction.
        if file_ext == 'json':
            file_type = 'json'
        elif file_ext == 'pdf':
            file_type = 'client_info_document'
        elif file_ext in ('docx', 'doc'):
            file_type = 'client_info_document'
        elif file_ext == 'csv':
            file_type = 'client_info_document'
        elif file_ext == 'txt':
            file_type = 'client_info_document'
        else:
            raise HTTPException(
                status_code=400,
                detail="Client info must be JSON, PDF, Word, CSV or TXT.",
            )
    
    elif file_category == 'bank_statement':
        if file_ext == 'csv':
            file_type = 'csv'
        elif file_ext == 'pdf':
            file_type = 'pdf'
        else:
            raise HTTPException(
                status_code=400,
                detail="Bank statement must be CSV or PDF"
            )
    
    elif file_category == 'supporting_doc':
        if file_ext != 'pdf':
            raise HTTPException(
                status_code=400,
                detail="Supporting documents must be PDF"
            )
        file_type = 'pdf_document'  # Use specific type for documents (NOT bank statements)
    
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file category: {file_category}"
        )
    
    # Read file content for document verification (before process_upload consumes it)
    file_content_for_verification = await file.read()
    await file.seek(0)  # Reset for process_upload

    # Persist file to disk for later retrieval (PDF viewer)
    import os as _os
    _upload_dir = f"/app/uploads/{matter_id}"
    _os.makedirs(_upload_dir, exist_ok=True)
    _safe_filename = _os.path.basename(file.filename or "unknown")
    _dest_path = _os.path.join(_upload_dir, _safe_filename)
    _counter = 1
    _base, _ext = _os.path.splitext(_safe_filename)
    while _os.path.exists(_dest_path):
        _safe_filename = f"{_base}_{_counter}{_ext}"
        _dest_path = _os.path.join(_upload_dir, _safe_filename)
        _counter += 1
    with open(_dest_path, "wb") as _f_out:
        _f_out.write(file_content_for_verification)

    # Process the file
    result = await file_processor.process_upload(file, file_type)

    if not result['success']:
        raise HTTPException(
            status_code=400,
            detail=f"File processing failed: {result['error']}"
        )

    # --- Document Verification (unified) ---
    # PDFs: run structural verification pipeline
    # CSVs: run statement validation pipeline (math check → binary pass/fail)
    import asyncio
    import hashlib as _hashlib

    verification_verdict = None
    verification_score = None
    doc_ver_record = None

    # Module master switch — when Document Verification is disabled on
    # the Configuration page (dv_enabled = false), we keep the upload
    # but skip the forensic pipeline entirely.
    dv_enabled_row = db.query(TransactionConfig).filter(
        TransactionConfig.key == 'dv_enabled'
    ).first()
    dv_enabled = (dv_enabled_row is None) or (str(dv_enabled_row.value).lower() in ('true', '1', 'yes'))

    # ------------------------------------------------------------------
    # Dedupe by file hash. Same file uploaded twice on the same matter
    # MUST produce the same verification result — otherwise the analyst
    # sees a different verdict on re-upload, which is confusing and
    # erodes trust in the pipeline. The pipeline is mostly
    # deterministic, but the OCR / cross-document stages have small
    # variation that can flip borderline flags between runs.
    # Short-circuiting on hash also avoids the cost of re-running the
    # full pipeline on a known file.
    # ------------------------------------------------------------------
    _precomputed_hash = _hashlib.sha256(file_content_for_verification).hexdigest()
    existing_dv = (
        db.query(DocumentVerification)
        .filter(
            DocumentVerification.matter_id == matter_id,
            DocumentVerification.file_hash == _precomputed_hash,
        )
        .first()
    )
    if existing_dv:
        verification_verdict = existing_dv.verdict.value if existing_dv.verdict else None
        verification_score = existing_dv.authenticity_score
        doc_ver_record = existing_dv
        print(
            f"   ↺ Skipping verification — file hash {_precomputed_hash[:12]}…"
            f" already verified as DV #{existing_dv.id} ({verification_verdict})"
        )

    try:
        is_pdf = file_content_for_verification[:4] == b"%PDF"

        if not dv_enabled:
            # Module turned off in the Configuration page. Persist the
            # raw file but skip the entire forensic pipeline; the
            # operator has accepted that this matter will not have
            # document-level checks.
            print(f"   ⚙️  Document Verification module is OFF — skipping pipeline for {file.filename}")
        elif existing_dv:
            # Already verified — skip the rest of the verification work.
            # Skip the entire pipeline branch by raising-to-continue.
            pass
        elif is_pdf:
            # --- PDF: structural verification ---
            loop = asyncio.get_event_loop()
            doc_ver_result: DocVerResult = await loop.run_in_executor(
                None,
                document_verification_pipeline.verify_document,
                file_content_for_verification,
                file.filename or "",
                file_category,
            )
            verification_verdict = doc_ver_result.verdict
            verification_score = doc_ver_result.authenticity_score

            doc_ver_record = DocumentVerification(
                matter_id=matter_id,
                filename=file.filename or "unknown",
                file_hash=doc_ver_result.file_hash_sha256,
                file_category=file_category,
                disk_filename=_safe_filename,
                file_bytes=file_content_for_verification,
                authenticity_score=doc_ver_result.authenticity_score,
                verdict=VerificationVerdict(doc_ver_result.verdict),
                verification_phase="structural_only",
                verification_method="PDF structural analysis",
                structural_pipeline_score=doc_ver_result.authenticity_score,
                metadata_result=doc_ver_result.metadata_result,
                structural_result=doc_ver_result.structural_result,
                font_text_result=doc_ver_result.font_text_result,
                image_result=doc_ver_result.image_result,
                content_consistency_result=doc_ver_result.content_consistency_result,
                signature_result=doc_ver_result.signature_result,
                annotation_form_result=getattr(doc_ver_result, 'annotation_form_result', None),
                hidden_content_result=getattr(doc_ver_result, 'hidden_content_result', None),
                blocked=(doc_ver_result.verdict == "LikelyTampered"),
            )
            db.add(doc_ver_record)
            db.flush()

            for flag in doc_ver_result.flags:
                db.add(DocumentVerificationFlag(
                    verification_id=doc_ver_record.id,
                    pipeline_stage=flag.pipeline_stage,
                    code=flag.code,
                    severity=flag.severity,
                    message=flag.message,
                    details=flag.details,
                ))
            db.commit()
            print(f"   PDF verification: {doc_ver_result.verdict} (score: {doc_ver_result.authenticity_score})")

            # Cross-document corroboration — does this PDF agree with the
            # other documents on the matter (client name, period gaps)?
            try:
                pdf_text = ""
                try:
                    import fitz
                    _doc = fitz.open(stream=file_content_for_verification, filetype="pdf")
                    pdf_text = "\n".join(p.get_text() or "" for p in _doc)
                    _doc.close()
                except Exception:
                    pdf_text = ""
                corr_flags = corroborate_verification(
                    db, matter_id, doc_ver_record.id,
                    new_doc_text=pdf_text,
                    new_period_start=doc_ver_record.period_start,
                    new_period_end=doc_ver_record.period_end,
                )
                for cf in corr_flags:
                    db.add(DocumentVerificationFlag(
                        verification_id=doc_ver_record.id,
                        pipeline_stage=cf.pipeline_stage,
                        code=cf.code,
                        severity=cf.severity,
                        message=cf.message,
                        details=cf.details,
                    ))
                if corr_flags:
                    db.commit()
                    print(f"   Cross-doc corroboration raised {len(corr_flags)} flag(s)")
            except Exception as corr_err:
                print(f"   Cross-document corroboration skipped ({corr_err})")

        elif file_category == "bank_statement" and not is_pdf:
            # --- CSV bank statement: run statement validation pipeline for math check ---
            loop = asyncio.get_event_loop()
            stmt_pipeline = StatementValidationPipeline()
            stmt_result: StmtValResult = await loop.run_in_executor(
                None,
                stmt_pipeline.validate_statement,
                file_content_for_verification,
                None,  # bank_hint
                None,  # period_start
                None,  # period_end
            )

            # Binary pass/fail based on math check
            math_result = stmt_result.math_check_result or {}
            math_passed = math_result.get("checks_passed", 0)
            math_total = math_result.get("checks_total", 0)
            math_pass = math_passed > 0 and math_passed >= math_total

            if math_pass:
                csv_verdict = VerificationVerdict.VERIFIED
                csv_method = "CSV uploaded — Math check: PASS"
                csv_score = stmt_result.authenticity_score
                csv_blocked = False
            else:
                csv_verdict = VerificationVerdict.LIKELY_TAMPERED
                csv_method = "CSV uploaded — Math check: FAIL"
                csv_score = stmt_result.authenticity_score
                csv_blocked = True

            verification_verdict = csv_verdict.value
            verification_score = csv_score

            file_hash = _hashlib.sha256(file_content_for_verification).hexdigest()
            doc_ver_record = DocumentVerification(
                matter_id=matter_id,
                filename=file.filename or "unknown",
                file_hash=file_hash,
                file_category=file_category,
                disk_filename=_safe_filename,
                file_bytes=file_content_for_verification,
                authenticity_score=csv_score,
                verdict=csv_verdict,
                verification_phase="statement_only",
                verification_method=csv_method,
                statement_pipeline_score=stmt_result.authenticity_score,
                file_integrity_result=stmt_result.file_integrity_result,
                template_match_result=stmt_result.template_match_result,
                extraction_result=stmt_result.extraction_result,
                math_check_result=stmt_result.math_check_result,
                anomaly_check_result=stmt_result.anomaly_check_result,
                identified_bank_template=stmt_result.identified_bank_template,
                blocked=csv_blocked,
            )
            db.add(doc_ver_record)
            db.flush()

            # Persist flags
            for flag in stmt_result.flags:
                db.add(DocumentVerificationFlag(
                    verification_id=doc_ver_record.id,
                    pipeline_stage=flag.pipeline_stage,
                    code=flag.code,
                    severity=flag.severity,
                    message=flag.message,
                    details=flag.details if isinstance(flag.details, dict) else None,
                ))

            # Persist extracted transactions
            for et in stmt_result.extracted_transactions:
                db.add(DocumentVerificationTransaction(
                    verification_id=doc_ver_record.id,
                    date=et.date,
                    description=et.description,
                    amount=et.amount,
                    direction=et.direction,
                    balance=et.balance,
                    transaction_type=et.transaction_type,
                    raw_row=et.raw_row,
                ))

            db.commit()
            print(f"   CSV verification: {csv_method} (score: {csv_score})")

            # Cross-document corroboration for CSV uploads. We use the raw
            # bytes decoded as text — good enough for a name-tokens search.
            try:
                try:
                    csv_text = file_content_for_verification.decode("utf-8", errors="ignore")
                except Exception:
                    csv_text = ""
                corr_flags = corroborate_verification(
                    db, matter_id, doc_ver_record.id,
                    new_doc_text=csv_text,
                    new_period_start=doc_ver_record.period_start,
                    new_period_end=doc_ver_record.period_end,
                )
                for cf in corr_flags:
                    db.add(DocumentVerificationFlag(
                        verification_id=doc_ver_record.id,
                        pipeline_stage=cf.pipeline_stage,
                        code=cf.code,
                        severity=cf.severity,
                        message=cf.message,
                        details=cf.details,
                    ))
                if corr_flags:
                    db.commit()
                    print(f"   Cross-doc corroboration raised {len(corr_flags)} flag(s)")
            except Exception as corr_err:
                print(f"   Cross-document corroboration skipped ({corr_err})")

        else:
            # Non-PDF, non-bank-statement files (e.g. client_info.json) — skip verification
            print(f"   Skipping verification for {file_category} file (not a bank statement or PDF)")

    except Exception as ver_err:
        print(f"   Document verification error (non-blocking): {ver_err}")
        import traceback
        traceback.print_exc()
    
    # Validate result based on file category
    if file_category == 'bank_statement':
        if 'bank_statements' not in result.get('data', {}):
            raise HTTPException(
                status_code=400,
                detail="PDF file does not contain valid bank statement data. No transactions could be extracted."
            )
    
    # Store the processed data
    
    if file_category == 'client_info':
        storage['client_info'] = result['data']
    
    elif file_category == 'bank_statement':
        # Merge bank statements
        new_transactions = result['data']['bank_statements']

        # ------------------------------------------------------------------
        # Per-file account attribution
        # ------------------------------------------------------------------
        # The PDF parsers stamp synthetic account_ids like "PDF_P1_T1"
        # based on page + table position. Two different uploaded
        # statements often collide on the same synthetic id, which
        # breaks the funds-lineage logic that groups by account ("Only
        # one account uploaded" even when two were).
        #
        # Override account_id per FILE so each upload is a distinct
        # account, and try to detect the account type from filename
        # / extracted text so the savings → current matcher works.
        if new_transactions:
            import re as _re
            filename_lower = (file.filename or '').lower()

            # ------------------------------------------------------------
            # Extract a real UK bank account number so two statements
            # from the SAME account merge into one logical account, and
            # statements from DIFFERENT accounts stay separate.
            #
            # Order of preference:
            #   1. Full PDF text (account number usually sits in the
            #      header band — most reliable place to find it).
            #   2. CSV/PDF-extracted transaction fields, in case the
            #      parser already pulled an explicit account_number.
            #   3. Any 8-digit run in description / narrative as a
            #      last-ditch search.
            #   4. Fall back to "file:<filename>" if nothing matches.
            #
            # We look for the canonical UK shape: an isolated 8-digit
            # run, optionally preceded by an "account" label. The
            # lookbehind for "account number"/"acct no"/"a/c"/"sort
            # code" rejects 8-digit runs that are clearly NOT the
            # account number (e.g. a date or transaction ref). When the
            # labelled form isn't present we fall back to a free 8-
            # digit run.
            # ------------------------------------------------------------
            account_number = None
            try:
                if is_pdf:
                    import fitz as _fitz
                    _d = _fitz.open(stream=file_content_for_verification, filetype="pdf")
                    pdf_text_full = "\n".join((p.get_text() or "") for p in _d)
                    _d.close()
                else:
                    pdf_text_full = file_content_for_verification.decode('utf-8', errors='ignore')
            except Exception:
                pdf_text_full = ''

            # Prefer labelled matches like "Account number: 12345678" /
            # "A/C 12345678" — these are far more reliable than a bare
            # 8-digit run.
            labelled = _re.search(
                r'(?:account\s*(?:number|no\.?)|a/?c(?:count)?\s*(?:no\.?|number)?|account)\s*[:#]?\s*(\d{8})',
                pdf_text_full,
                _re.IGNORECASE,
            )
            if labelled:
                account_number = labelled.group(1)

            # Fall back to ANY isolated 8-digit run in the header (first
            # 1000 chars) — most banks print the number near the top.
            if not account_number and pdf_text_full:
                m = _re.search(r'\b(\d{8})\b', pdf_text_full[:1000])
                if m:
                    account_number = m.group(1)

            # Last-ditch: scan the parsed transaction fields.
            if not account_number:
                for txn in new_transactions[:50]:
                    for field in ('account_number', 'description', 'narrative'):
                        val = str(txn.get(field, ''))
                        m = _re.search(r'\b(\d{8})\b', val)
                        if m:
                            account_number = m.group(1)
                            break
                    if account_number:
                        break

            # Per-file identifier — falls back to filename when no real
            # account number could be extracted, so different uploads
            # still register as distinct accounts.
            file_account_id = (
                account_number
                or f"file:{file.filename or 'unknown'}"
            )

            # Detect account type from filename hints (only override
            # when the parser left it blank or returned 'Unknown').
            detected_type = None
            if any(k in filename_lower for k in ('savings', 'isa', 'easy access', 'easyaccess', 'reserve')):
                detected_type = 'savings'
            elif any(k in filename_lower for k in ('current', 'cheque', 'cheqing')):
                detected_type = 'current'

            for txn in new_transactions:
                txn['account_id'] = file_account_id
                txn['source_filename'] = file.filename
                if detected_type and not str(txn.get('account_type', '')).lower() in ('current', 'savings'):
                    txn['account_type'] = detected_type

            print(f"   📋 Adding {len(new_transactions)} transactions (file: {file.filename})")
            print(f"      Account ID (file-scoped): {file_account_id}")
            print(f"      Account Type: {new_transactions[0].get('account_type', 'N/A')}")
            print(f"      Bank: {new_transactions[0].get('bank_name', 'N/A')}")
            print(f"      Sort Code: {new_transactions[0].get('sort_code', 'N/A')}")
        
        storage['bank_statements'].extend(new_transactions)
        
        # Show unique accounts after merge
        unique_accounts = set()
        for txn in storage['bank_statements']:
            acc_id = txn.get('account_id', 'Unknown')
            acc_type = txn.get('account_type', '')
            bank = txn.get('bank_name', '')
            unique_accounts.add(f"{bank} {acc_type} ({acc_id})")
        print(f"   🏦 Total accounts now: {len(unique_accounts)}")
        for acc in sorted(unique_accounts):
            print(f"      - {acc}")
    
    elif file_category == 'supporting_doc':
        # Add filename and upload timestamp to the document data for audit trail
        doc_data = result['data'].copy()
        doc_data['filename'] = file.filename
        doc_data['uploaded_at'] = datetime.now(timezone.utc).isoformat()
        storage['supporting_docs'].append(doc_data)
    
    # Track uploaded file
    storage['uploaded_files'].append({
        "filename": file.filename,
        "category": file_category,
        "file_type": result['file_type'],
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "records_count": result['data'].get('transaction_count') or 1
    })
    
    storage['last_updated'] = datetime.now(timezone.utc).isoformat()
    storage['status'] = 'files_uploaded'
    
    # Persist storage to database
    _db_save_storage(db, matter_id, storage)

    response_data = {
        "success": True,
        "matter_id": matter_id,
        "file_category": file_category,
        "filename": file.filename,
        "records_processed": result['data'].get('transaction_count') or 1,
        "message": "File processed successfully",
    }
    if verification_verdict:
        response_data["verification_verdict"] = verification_verdict
        response_data["verification_score"] = verification_score
        response_data["verification_method"] = doc_ver_record.verification_method if doc_ver_record else None

    # Return the parsed client_info payload so the frontend can pre-fill
    # the manual form fields and let the user verify / edit before
    # running the assessment. Only sent for client_info; bank_statement
    # and supporting_doc don't need this.
    if file_category == 'client_info' and isinstance(result.get('data'), dict):
        d = result['data']
        ci = d.get('client_info') if isinstance(d.get('client_info'), dict) else {}
        purchase = d.get('purchase') if isinstance(d.get('purchase'), dict) else {}

        # sof_explanation is sometimes a plain string (regex extractor,
        # simple JSON files) and sometimes a structured object (richer
        # JSON exports with a `sources` array). The frontend renders
        # it as text in a textarea, so flatten objects into plain
        # English here rather than letting React render
        # "[object Object]" or a raw Python repr.
        def _humanise_amount(amount, currency):
            try:
                amt = float(amount)
            except (TypeError, ValueError):
                return None
            cur = (currency or 'GBP').upper().strip()
            sym = {'GBP': '£', 'USD': '$', 'EUR': '€'}.get(cur, '')
            formatted = f"{amt:,.0f}" if amt == int(amt) else f"{amt:,.2f}"
            return f"{sym}{formatted} {cur}" if sym else f"{formatted} {cur}"

        def _humanise_date(d):
            if not d:
                return None
            from datetime import datetime as _dt
            for fmt in ('%Y-%m-%d', '%Y-%m-%dT%H:%M:%S', '%d/%m/%Y', '%d-%m-%Y'):
                try:
                    return _dt.strptime(str(d)[:len(fmt) + 4], fmt).strftime('%d %B %Y')
                except (ValueError, TypeError):
                    continue
            return str(d)

        def _source_to_prose(idx, src):
            """Turn one source dict (property_sale / savings / salary
            / inheritance / etc.) into a numbered prose paragraph."""
            if not isinstance(src, dict):
                return f"{idx}. {src}"
            source_type = (src.get('source_type') or 'Source').replace('_', ' ').title()
            amount_str = _humanise_amount(src.get('amount'), src.get('currency'))
            head = f"{idx}. {source_type}"
            if amount_str:
                head += f" — {amount_str}"
            lines = [head]
            description = (src.get('description') or '').strip()
            if description:
                lines.append(f"   {description}")
            # Optional contextual lines, in a stable order so output
            # reads consistently. Hidden when missing.
            ctx_keys = [
                ('property_address',   'Address'),
                ('title_number',       'Title number'),
                ('completion_date',    'Completion date'),
                ('solicitor_firm',     'Solicitor'),
                ('bank',               'Bank'),
                ('account_number',     'Account'),
                ('sort_code',          'Sort code'),
                ('employer',           'Employer'),
                ('grant_date',         'Granted on'),
                ('grant_authority',    'Granted by'),
                ('payer',              'From'),
                ('reference',          'Reference'),
                ('notes',              'Notes'),
            ]
            for key, label in ctx_keys:
                val = src.get(key)
                if val in (None, ''):
                    continue
                if 'date' in key:
                    val = _humanise_date(val) or val
                lines.append(f"   {label}: {val}")
            return "\n".join(lines)

        def _stringify_sof(v):
            if v is None:
                return ''
            if isinstance(v, str):
                return v
            if isinstance(v, dict):
                # 1) If there's already a narrative-style field, use it.
                for k in ('text', 'summary', 'narrative', 'explanation'):
                    if isinstance(v.get(k), str) and v[k].strip():
                        return v[k]

                # 2) If there's a `sources` array (property sales,
                #    savings, salary, inheritance etc.), render each
                #    as a numbered prose block. This is the AML-rich
                #    shape we see in real client info files.
                sources = v.get('sources')
                if isinstance(sources, (list, tuple)) and sources:
                    blocks: list = []
                    # Lead with a total line when a totals field is present.
                    total = v.get('total_amount') or v.get('total')
                    total_str = _humanise_amount(total, v.get('currency') or 'GBP')
                    if total_str:
                        blocks.append(f"Total funds: {total_str}, from {len(sources)} source(s):")
                    elif len(sources) > 1:
                        blocks.append(f"The funds come from {len(sources)} sources:")
                    for i, s in enumerate(sources, start=1):
                        blocks.append(_source_to_prose(i, s))
                    # Trailing narrative if any.
                    if isinstance(v.get('notes'), str) and v['notes'].strip():
                        blocks.append(f"Notes: {v['notes']}")
                    return "\n\n".join(blocks)

                # 3) Generic fallback — render the dict as readable
                #    "Key: value" lines, prettifying the keys.
                parts = []
                for k, val in v.items():
                    label = k.replace('_', ' ').capitalize()
                    if isinstance(val, (str, int, float)):
                        parts.append(f"{label}: {val}")
                    elif isinstance(val, (list, tuple)):
                        parts.append(f"{label}: {', '.join(str(x) for x in val)}")
                    elif isinstance(val, dict):
                        parts.append(f"{label}: {_stringify_sof(val)}")
                return "\n".join(parts)

            if isinstance(v, (list, tuple)):
                # List of sources without a wrapping dict.
                if v and all(isinstance(x, dict) for x in v):
                    return "\n\n".join(_source_to_prose(i + 1, s) for i, s in enumerate(v))
                return "\n".join(_stringify_sof(x) for x in v if x)
            return str(v)

        response_data["parsed_client_info"] = {
            "client_info": {
                "client_name":         ci.get('client_name'),
                "client_risk_rating":  ci.get('client_risk_rating'),
                "business_sector":     ci.get('business_sector'),
                "pep_status":          bool(ci.get('pep_status') or d.get('flags', {}).get('pep') if isinstance(d.get('flags'), dict) else ci.get('pep_status')),
            },
            "purchase": {
                "amount":                  purchase.get('amount'),
                "currency":                purchase.get('currency') or 'GBP',
                "expected_payment_date":   purchase.get('expected_payment_date'),
                "description":             purchase.get('description'),
            },
            "sof_explanation": _stringify_sof(d.get('sof_explanation')),
        }

    return response_data


@router.get("/matters/{matter_id}/sof-assessment/status")
async def get_sof_assessment_status(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Get current status of SoF assessment for a matter
    """
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Load assessment data from database
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        return {
            "matter_id": matter_id,
            "status": "no_data",
            "uploaded_files": [],
            "ready_for_assessment": False
        }
    
    # Check if ready for assessment
    has_client_info = storage['client_info'] is not None
    has_bank_statements = len(storage['bank_statements']) > 0
    ready = has_client_info and has_bank_statements
    
    return {
        "matter_id": matter_id,
        "status": storage['status'],
        "uploaded_files": storage['uploaded_files'],
        "files_summary": {
            "client_info": "uploaded" if has_client_info else "missing",
            "bank_statements_count": len(storage['bank_statements']),
            "supporting_docs_count": len(storage['supporting_docs'])
        },
        "ready_for_assessment": ready,
        "last_updated": storage['last_updated']
    }


@router.post("/matters/{matter_id}/sof-assessment/run")
async def run_sof_assessment(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Run SoF assessment engine on uploaded data
    Integrates with Transaction Review automatically
    """
    # Load storage from database
    print(f"\n🔄 Run assessment for matter {matter_id}")

    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Get assessment data from DB
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        print(f"   ❌ Matter {matter_id} not found in storage!")
        raise HTTPException(
            status_code=400,
            detail="No files uploaded for this matter"
        )

    print(f"   ✅ Found storage for matter {matter_id}: status={storage.get('status')}, files={len(storage.get('uploaded_files', []))}")
    
    # Validate required data
    if not storage.get('client_info'):
        raise HTTPException(
            status_code=400,
            detail="Client info file required (JSON with client details, purchase, SoF explanation)"
        )
    
    if not storage.get('bank_statements'):
        raise HTTPException(
            status_code=400,
            detail="At least one bank statement file required (CSV or PDF)"
        )
    
    # Extract data with comprehensive error handling
    client_info_data = storage.get('client_info') or {}
    
    # DEFENSIVE: Ensure client_info_data is a dict
    if not isinstance(client_info_data, dict):
        print(f"⚠️ client_info_data is not a dict: {type(client_info_data)}")
        client_info_data = {}
    
    print(f"\n=== CLIENT INFO DEBUG ===")
    print(f"client_info_data type: {type(client_info_data)}")
    print(f"client_info_data keys: {client_info_data.keys() if isinstance(client_info_data, dict) else 'N/A'}")
    print(f"=========================\n")
    
    # Handle different client_info structures
    if isinstance(client_info_data, dict) and 'client_info' in client_info_data:
        client_info = client_info_data.get('client_info') or {}
    else:
        # Flat structure - use the whole thing as client_info
        client_info = {
            'client_name': client_info_data.get('client_name', 'Unknown') if isinstance(client_info_data, dict) else 'Unknown',
            'client_risk_rating': client_info_data.get('client_risk_rating', 'medium') if isinstance(client_info_data, dict) else 'medium',
            'pep_status': client_info_data.get('pep_status', False) if isinstance(client_info_data, dict) else False,
            'business_sector': client_info_data.get('business_sector', 'Unknown') if isinstance(client_info_data, dict) else 'Unknown'
        }
    
    # DEFENSIVE: Ensure client_info is a dict
    if not isinstance(client_info, dict):
        print(f"⚠️ client_info is not a dict: {type(client_info)}")
        client_info = {'client_name': 'Unknown', 'client_risk_rating': 'medium', 'pep_status': False, 'business_sector': 'Unknown'}
    
    if isinstance(client_info_data, dict) and 'purchase' in client_info_data:
        purchase = client_info_data.get('purchase') or {}
    else:
        purchase = {
            'amount': client_info_data.get('purchase_amount', 0) if isinstance(client_info_data, dict) else 0,
            'currency': client_info_data.get('currency', 'GBP') if isinstance(client_info_data, dict) else 'GBP',
            'description': client_info_data.get('purchase_description', 'Property purchase') if isinstance(client_info_data, dict) else 'Property purchase'
        }
    
    # DEFENSIVE: Ensure purchase is a dict
    if not isinstance(purchase, dict):
        print(f"⚠️ purchase is not a dict: {type(purchase)}")
        purchase = {'amount': 0, 'currency': 'GBP', 'description': 'Unknown'}
    
    if isinstance(client_info_data, dict) and 'sof_explanation' in client_info_data:
        sof_explanation = client_info_data.get('sof_explanation') or ''
    else:
        sof_explanation = client_info_data.get('explanation', 'No explanation provided') if isinstance(client_info_data, dict) else 'No explanation provided'
    
    # DEFENSIVE: Ensure sof_explanation is a string or dict
    if sof_explanation is None:
        sof_explanation = ''
    
    bank_statements = storage.get('bank_statements') or []
    
    # Check if client_info JSON has explicit claims array
    # If so, convert to structured format for the assessment engine
    print(f"\n=== CLAIMS CHECK ===")
    print(f"Has 'claims' key: {'claims' in client_info_data}")
    if 'claims' in client_info_data:
        print(f"Claims array: {client_info_data['claims']}")
    
    if 'claims' in client_info_data and client_info_data['claims']:
        # Convert claims array to structured format
        print(f"✅ Converting claims array to structured format")
        sof_explanation = {
            'sources': client_info_data['claims']
        }
        print(f"sof_explanation type: {type(sof_explanation)}")
    print(f"====================\n")
    
    # IMPORTANT: Build fresh known_documents list from current uploads only
    # Do NOT accumulate from previous assessments
    known_documents = []
    supporting_docs_data = storage['supporting_docs']  # Full document data with extracted info
    
    # Add uploaded supporting docs to known documents
    for doc in storage['supporting_docs']:
        doc_type = doc.get('document_type', 'unknown')
        if doc_type != 'unknown':
            known_documents.append(doc_type)
    
    flags = client_info_data.get('flags', {})
    constraints = client_info_data.get('constraints', {})
    
    # DEBUG LOGGING
    print(f"\n=== SoF ASSESSMENT DEBUG ===")
    print(f"Matter ID: {matter_id}")
    print(f"Supporting docs uploaded: {len(supporting_docs_data)}")
    for idx, doc in enumerate(supporting_docs_data):
        print(f"  Doc {idx}: Type={doc.get('document_type')}, Has extracted_data={bool(doc.get('extracted_data'))}")
        if doc.get('extracted_data'):
            print(f"    Extracted fields: {list(doc.get('extracted_data', {}).keys())}")
    print(f"Known documents: {known_documents}")
    print(f"===========================\n")
    
    # ============================================================
    # DOCUMENT VERIFICATION — ASSESSMENT-TIME UPDATE
    # For PDFs that only had structural analysis at upload, now run
    # the statement validation pipeline and combine scores.
    # CSVs already have full results from upload — no re-run needed.
    # ============================================================
    document_verification_summary = {
        "total_documents": 0,
        "verified_count": 0,
        "suspicious_count": 0,
        "likely_tampered_count": 0,
        "blocked_count": 0,
        "overridden_count": 0,
        "average_score": 0.0,
        "verifications": [],
        "all_flags": [],
        "has_blocking_issues": False,
    }

    try:
        pipeline = StatementValidationPipeline()
        uploaded_files_info = storage.get('uploaded_files', [])

        # Group bank statement transactions by their source file
        file_groups = {}
        file_list = [uf for uf in uploaded_files_info if uf.get('category') == 'bank_statement']
        txn_idx = 0
        for fi, uf in enumerate(file_list):
            fname = uf.get('filename', f'statement_{fi}')
            count = uf.get('records_count', 0)
            file_groups[fname] = bank_statements[txn_idx:txn_idx + count]
            txn_idx += count

        if not file_groups and bank_statements:
            file_groups['bank_statement'] = bank_statements

        print(f"\n=== DOCUMENT VERIFICATION — ASSESSMENT UPDATE ===")
        print(f"Processing {len(file_groups)} statement file(s)...")

        score_total = 0.0

        for fname, txns in file_groups.items():
            # Find existing DocumentVerification record for this file
            existing_dv = (
                db.query(DocumentVerification)
                .filter(
                    DocumentVerification.matter_id == matter_id,
                    DocumentVerification.filename == fname,
                )
                .order_by(DocumentVerification.created_at.desc())
                .first()
            )

            # If CSV and already has statement results, skip re-run
            if existing_dv and existing_dv.verification_phase == "statement_only":
                print(f"  📄 {fname}: CSV already verified (score={existing_dv.authenticity_score}, verdict={existing_dv.verdict.value})")
                score_total += existing_dv.authenticity_score
                document_verification_summary["total_documents"] += 1
                if existing_dv.verdict == VerificationVerdict.VERIFIED:
                    document_verification_summary["verified_count"] += 1
                elif existing_dv.verdict == VerificationVerdict.SUSPICIOUS:
                    document_verification_summary["suspicious_count"] += 1
                elif existing_dv.verdict == VerificationVerdict.LIKELY_TAMPERED:
                    document_verification_summary["likely_tampered_count"] += 1
                if existing_dv.blocked:
                    document_verification_summary["blocked_count"] += 1
                if existing_dv.admin_override:
                    document_verification_summary["overridden_count"] += 1
                document_verification_summary["verifications"].append(existing_dv.to_dict())
                for f in existing_dv.flags:
                    document_verification_summary["all_flags"].append(f.to_dict())
                continue

            # Reconstruct CSV from parsed transactions for the pipeline
            import io as _io
            import csv as _csv

            csv_buffer = _io.StringIO()
            writer = _csv.writer(csv_buffer)
            writer.writerow(['date', 'description', 'amount', 'direction', 'balance', 'account_id', 'bank_name'])
            for t in txns:
                amt = t.get('amount', 0)
                direction = t.get('direction', 'debit')
                signed_amt = amt if direction == 'credit' else -amt
                writer.writerow([
                    t.get('date', ''),
                    t.get('description', ''),
                    signed_amt,
                    direction,
                    t.get('balance', ''),
                    t.get('account_id', ''),
                    t.get('bank_name', ''),
                ])

            file_bytes = csv_buffer.getvalue().encode('utf-8')
            bank_hint = txns[0].get('bank_name') if txns else None
            txn_dates = [t.get('date', '') for t in txns if t.get('date')]
            p_start = min(txn_dates) if txn_dates else None
            p_end = max(txn_dates) if txn_dates else None

            vresult = pipeline.validate_statement(
                file_bytes=file_bytes,
                bank_hint=bank_hint,
                period_start=p_start,
                period_end=p_end,
            )

            print(f"  📄 {fname}: stmt_score={vresult.authenticity_score}, status={vresult.status}")

            if existing_dv and existing_dv.verification_phase == "structural_only":
                # PDF: combine structural + statement scores
                structural_score = existing_dv.structural_pipeline_score or existing_dv.authenticity_score or 0
                stmt_score = vresult.authenticity_score
                combined_score = round(structural_score * 0.4 + stmt_score * 0.6, 1)

                existing_dv.statement_pipeline_score = stmt_score
                existing_dv.authenticity_score = combined_score
                existing_dv.verification_phase = "complete"
                existing_dv.verification_method = "PDF structural + statement analysis"
                existing_dv.file_integrity_result = vresult.file_integrity_result
                existing_dv.template_match_result = vresult.template_match_result
                existing_dv.extraction_result = vresult.extraction_result
                existing_dv.math_check_result = vresult.math_check_result
                existing_dv.anomaly_check_result = vresult.anomaly_check_result
                existing_dv.identified_bank_template = vresult.identified_bank_template
                existing_dv.bank_hint = bank_hint
                existing_dv.period_start = p_start
                existing_dv.period_end = p_end

                # Re-evaluate verdict based on combined score
                critical_flags = [f for f in existing_dv.flags if f.severity == "critical"]
                high_flags = [f for f in existing_dv.flags if f.severity == "high"]
                if critical_flags or combined_score < 45:
                    existing_dv.verdict = VerificationVerdict.LIKELY_TAMPERED
                    existing_dv.blocked = True
                elif len(high_flags) >= 2 or combined_score < 75:
                    existing_dv.verdict = VerificationVerdict.SUSPICIOUS
                else:
                    existing_dv.verdict = VerificationVerdict.VERIFIED

                # Add statement pipeline flags
                for f in vresult.flags:
                    db.add(DocumentVerificationFlag(
                        verification_id=existing_dv.id,
                        pipeline_stage=f.pipeline_stage,
                        code=f.code,
                        severity=f.severity,
                        message=f.message,
                        details=f.details if isinstance(f.details, dict) else None,
                    ))

                # Add extracted transactions
                for et in vresult.extracted_transactions:
                    db.add(DocumentVerificationTransaction(
                        verification_id=existing_dv.id,
                        date=et.date,
                        description=et.description,
                        amount=et.amount,
                        direction=et.direction,
                        balance=et.balance,
                        transaction_type=et.transaction_type,
                        raw_row=et.raw_row,
                    ))

                db.commit()
                # Refresh to get updated flags
                db.refresh(existing_dv)
                dv_record = existing_dv
                final_score = combined_score
                print(f"    Combined: structural={structural_score} * 0.4 + stmt={stmt_score} * 0.6 = {combined_score}")
            else:
                # No existing record — create new (shouldn't happen normally, but handle gracefully)
                import hashlib as _hl
                dv_record = DocumentVerification(
                    matter_id=matter_id,
                    filename=fname,
                    file_hash=_hl.sha256(file_bytes).hexdigest(),
                    file_category="bank_statement",
                    authenticity_score=vresult.authenticity_score,
                    verdict=VerificationVerdict.VERIFIED if vresult.status == "Trusted"
                        else VerificationVerdict.SUSPICIOUS if vresult.status == "Review"
                        else VerificationVerdict.LIKELY_TAMPERED,
                    verification_phase="statement_only",
                    verification_method="Statement analysis",
                    statement_pipeline_score=vresult.authenticity_score,
                    file_integrity_result=vresult.file_integrity_result,
                    template_match_result=vresult.template_match_result,
                    extraction_result=vresult.extraction_result,
                    math_check_result=vresult.math_check_result,
                    anomaly_check_result=vresult.anomaly_check_result,
                    identified_bank_template=vresult.identified_bank_template,
                    bank_hint=bank_hint,
                    period_start=p_start,
                    period_end=p_end,
                    blocked=(vresult.status == "HighRisk"),
                )
                db.add(dv_record)
                db.flush()
                for f in vresult.flags:
                    db.add(DocumentVerificationFlag(
                        verification_id=dv_record.id,
                        pipeline_stage=f.pipeline_stage,
                        code=f.code,
                        severity=f.severity,
                        message=f.message,
                        details=f.details if isinstance(f.details, dict) else None,
                    ))
                for et in vresult.extracted_transactions:
                    db.add(DocumentVerificationTransaction(
                        verification_id=dv_record.id,
                        date=et.date,
                        description=et.description,
                        amount=et.amount,
                        direction=et.direction,
                        balance=et.balance,
                        transaction_type=et.transaction_type,
                        raw_row=et.raw_row,
                    ))
                db.commit()
                final_score = vresult.authenticity_score

            # Build summary
            document_verification_summary["total_documents"] += 1
            score_total += dv_record.authenticity_score
            if dv_record.verdict == VerificationVerdict.VERIFIED:
                document_verification_summary["verified_count"] += 1
            elif dv_record.verdict == VerificationVerdict.SUSPICIOUS:
                document_verification_summary["suspicious_count"] += 1
            elif dv_record.verdict == VerificationVerdict.LIKELY_TAMPERED:
                document_verification_summary["likely_tampered_count"] += 1
            if dv_record.blocked and not dv_record.admin_override:
                document_verification_summary["blocked_count"] += 1
            if dv_record.admin_override:
                document_verification_summary["overridden_count"] += 1
            document_verification_summary["verifications"].append(dv_record.to_dict())
            for f in (dv_record.flags or []):
                document_verification_summary["all_flags"].append(f.to_dict())

        if document_verification_summary["total_documents"] > 0:
            document_verification_summary["average_score"] = round(
                score_total / document_verification_summary["total_documents"], 1
            )
        document_verification_summary["has_blocking_issues"] = any(
            v.get("blocked") and not v.get("admin_override")
            for v in document_verification_summary["verifications"]
        )

        print(f"=== VERIFICATION COMPLETE: avg={document_verification_summary['average_score']}, "
              f"verified={document_verification_summary['verified_count']}, "
              f"suspicious={document_verification_summary['suspicious_count']}, "
              f"tampered={document_verification_summary['likely_tampered_count']} ===\n")

    except Exception as val_err:
        import traceback
        print(f"⚠️ Document verification update error (non-fatal): {val_err}")
        print(traceback.format_exc())

    # Store both keys for backward compatibility
    storage['document_verification_summary'] = document_verification_summary
    storage['statement_validation_summary'] = document_verification_summary
    
    # Run assessment engine
    engine = SoFAssessmentEngine(matter_id=matter_id, db=db)
    
    # PRESERVE MANUAL ACCEPTANCES: Store any existing manual acceptances before re-running
    previous_manual_acceptances = {}
    previous_assessment = storage.get('assessment_result')
    if previous_assessment:
        prev_claims = previous_assessment.get('claims', [])
        prev_evidence = previous_assessment.get('evidence_matches', [])
        
        for idx, evidence in enumerate(prev_evidence):
            doc_ver = evidence.get('document_verification', {})
            if doc_ver.get('manual_review_status') == 'accepted':
                # Store by claim source_type and amount for matching
                claim = prev_claims[idx] if idx < len(prev_claims) else {}
                key = f"{claim.get('source_type', '')}_{claim.get('expected_amount', 0)}"
                
                # Store ALL relevant data needed to fully restore the accepted state
                previous_manual_acceptances[key] = {
                    'manual_review_status': 'accepted',
                    'manually_accepted_by': doc_ver.get('manually_accepted_by'),
                    'manually_accepted_at': doc_ver.get('manually_accepted_at'),
                    'acceptance_reason': doc_ver.get('acceptance_reason', doc_ver.get('manual_acceptance_reason', '')),
                    'verification_note': doc_ver.get('verification_note', 'Marked as verified post user acceptance of differences'),
                }
                print(f"  📋 Preserved manual acceptance for {key}")
    
    try:
        assessment_result = engine.assess(
            client_info=client_info,
            purchase=purchase,
            sof_explanation=sof_explanation,
            bank_statements=bank_statements,
            known_documents=known_documents,
            supporting_docs_data=supporting_docs_data,  # Pass full document data
            constraints=constraints,
            flags=flags
        )
        
        # Store result (will be modified by funds lineage below if applicable)
        storage['assessment_result'] = assessment_result
        storage['status'] = 'completed'
        storage['last_updated'] = datetime.now(timezone.utc).isoformat()
        
        # AUTO-RUN FUNDS LINEAGE FOR SAVINGS CLAIMS
        # Check if there's a savings claim and automatically run lineage analysis.
        # Skipped entirely when the Funds Lineage module is turned off in
        # the Configuration page (fl_enabled = false).
        fl_enabled_row = db.query(TransactionConfig).filter(
            TransactionConfig.key == 'fl_enabled'
        ).first()
        fl_enabled = (fl_enabled_row is None) or (str(fl_enabled_row.value).lower() in ('true', '1', 'yes'))

        savings_claim = None
        savings_claim_idx = None
        claims = assessment_result.get('claims', []) if fl_enabled else []
        evidence_matches = assessment_result.get('evidence_matches', []) if fl_enabled else []
        
        for idx, claim in enumerate(claims):
            source_type = (claim.get('source_type', '') or '').lower()
            if 'saving' in source_type or 'accumul' in source_type:
                savings_claim = claim
                savings_claim_idx = idx
                break
        
        if savings_claim:
            print(f"\n=== AUTO FUNDS LINEAGE FOR SAVINGS CLAIM ===")
            # Use expected_amount field from claim, or from evidence_matches if available
            savings_amount = savings_claim.get('expected_amount', 0) or savings_claim.get('amount', 0)
            
            # If still 0, try to get from evidence_matches
            if savings_amount == 0 and savings_claim_idx is not None and savings_claim_idx < len(evidence_matches):
                savings_amount = evidence_matches[savings_claim_idx].get('expected_amount', 0)
            
            print(f"Savings claim amount: £{savings_amount:,.2f}")
            
            # Find EXACT matching transaction for the savings claim amount
            # Only auto-run if we find an exact match - otherwise user must select manually
            target_transaction = None
            
            print(f"  Looking for exact credit of £{savings_amount:,.2f} in bank statements...")
            for txn in bank_statements:
                if txn.get('direction') == 'credit':
                    txn_amount = float(txn.get('amount', 0))
                    # Exact match within £1 tolerance for rounding
                    if abs(txn_amount - savings_amount) < 1.0:
                        target_transaction = txn
                        print(f"    ✓ EXACT MATCH FOUND: £{txn_amount:,.2f} - {txn.get('description', '')[:50]}")
                        break
            
            if target_transaction:
                print(f"  Auto-running lineage for exact match: £{target_transaction.get('amount'):,.2f}")
                
                # Run automated funds lineage analysis
                lineage_result = run_automated_funds_lineage(
                    target_transaction=target_transaction,
                    bank_statements=bank_statements,
                    savings_claim=savings_claim
                )
                
                # Store lineage results
                storage['funds_lineage'] = lineage_result
                
                # Update the savings claim evidence with lineage info
                print(f"\n  === ADDING LINEAGE RESULTS TO EVIDENCE ===")
                print(f"  savings_claim_idx: {savings_claim_idx}")
                print(f"  len(evidence_matches): {len(evidence_matches)}")
                
                if savings_claim_idx is not None and savings_claim_idx < len(evidence_matches):
                    evidence = evidence_matches[savings_claim_idx]
                    evidence['funds_lineage_verified'] = True
                    evidence['lineage_summary'] = lineage_result.get('summary', {})
                    
                    # Get or create document_verification
                    if 'document_verification' not in evidence:
                        evidence['document_verification'] = {
                            'claim_id': savings_claim_idx,
                            'verified': False,
                            'confidence': 0.4,
                            'differences': [],
                            'issues': []
                        }
                    
                    doc_ver = evidence['document_verification']
                    
                    # Clear any existing differences (we're replacing with lineage results)
                    doc_ver['differences'] = []
                    
                    # Get lineage summary
                    summary = lineage_result.get('summary', {})
                    traced_pct = summary.get('traced_percentage', 0)
                    traced_amt = summary.get('tracedAmount', 0)
                    total_amt = summary.get('totalAmount', 0)
                    external_origins = summary.get('externalOrigins', 0)
                    unresolved = lineage_result.get('unresolved_items', [])
                    statement_gaps = lineage_result.get('statement_gap_items', [])
                    statement_gaps_summary = summary.get('statementGaps', [])
                    
                    # Add lineage completion status as first difference
                    severity = 'verified' if traced_pct >= 80 else 'review_required' if traced_pct >= 50 else 'evidence_required'
                    doc_ver['differences'].append({
                        'field': 'funds_lineage_analysis',
                        'severity': severity,
                        'issue': f'Funds Lineage Complete: {traced_pct}% of £{total_amt:,.2f} traced to origin',
                        'expected': f'£{total_amt:,.2f} traced to legitimate sources',
                        'found': f'£{traced_amt:,.2f} traced ({external_origins} external income sources identified)'
                    })
                    
                    # Add statement coverage info if available
                    account_coverage = summary.get('accountCoverage', {})
                    if account_coverage:
                        coverage_lines = []
                        for acc_id, range_info in account_coverage.items():
                            coverage_lines.append(f"{acc_id}: {range_info.get('earliest_str', 'N/A')} to {range_info.get('latest_str', 'N/A')}")
                        if coverage_lines:
                            doc_ver['differences'].append({
                                'field': 'statement_coverage',
                                'severity': 'info',
                                'issue': 'Statement Coverage',
                                'expected': 'Full coverage of accumulation period',
                                'found': '; '.join(coverage_lines)
                            })
                    
                    # Add statement gap items (where we need more statements)
                    for gap in statement_gaps:
                        doc_ver['differences'].append({
                            'field': 'statement_gap',
                            'severity': 'documents_required',
                            'issue': f"Statement Gap: £{gap.get('amount', 0):,.2f} on {gap.get('date', 'unknown')}",
                            'expected': f"Need {gap.get('gap_account', 'account')} statements from before {gap.get('gap_account_earliest', 'unknown')}",
                            'found': gap.get('description', 'Transfer to savings - source account statements needed'),
                            'transaction_id': gap.get('id', gap.get('transaction_id', 'N/A')),
                            'amount': gap.get('amount', 0),
                            'date': gap.get('date', 'unknown'),
                            'gap_account': gap.get('gap_account', ''),
                            'gap_account_earliest': gap.get('gap_account_earliest', '')
                        })
                    
                    # Add each unresolved item as a difference with transaction ID
                    for item in unresolved:
                        doc_ver['differences'].append({
                            'field': 'untraced_funds',
                            'severity': 'evidence_required',
                            'issue': f"Untraced: £{item.get('amount', 0):,.2f} on {item.get('date', 'unknown')}",
                            'expected': 'Source documentation to verify origin',
                            'found': item.get('description', 'Source unknown'),
                            'transaction_id': item.get('id', item.get('transaction_id', 'N/A')),
                            'amount': item.get('amount', 0),
                            'date': item.get('date', 'unknown')
                        })
                    
                    # CHECK FOR DISCREPANCY: traced + identified untraced vs claimed amount
                    unresolved_total = sum(item.get('amount', 0) for item in unresolved)
                    gap_total = sum(item.get('amount', 0) for item in statement_gaps)
                    identified_untraced = unresolved_total + gap_total
                    accounted_for = traced_amt + identified_untraced
                    discrepancy = total_amt - accounted_for
                    
                    print(f"  📊 DISCREPANCY CHECK:")
                    print(f"     Total claimed: £{total_amt:,.2f}")
                    print(f"     Traced amount: £{traced_amt:,.2f}")
                    print(f"     Unresolved total: £{unresolved_total:,.2f} ({len(unresolved)} items)")
                    print(f"     Gap total: £{gap_total:,.2f} ({len(statement_gaps)} items)")
                    print(f"     Identified untraced: £{identified_untraced:,.2f}")
                    print(f"     Accounted for: £{accounted_for:,.2f}")
                    print(f"     Discrepancy: £{discrepancy:,.2f}")
                    
                    # Only add discrepancy if significant (more than £100)
                    if discrepancy > 100:
                        doc_ver['differences'].append({
                            'field': 'funds_discrepancy',
                            'severity': 'discrepancy',
                            'issue': f"Funds accounted for: Traced (£{traced_amt:,.2f}) + Untraced identified (£{identified_untraced:,.2f}) = £{accounted_for:,.2f}",
                            'expected': f"Total claimed: £{total_amt:,.2f}",
                            'found': f"Shortfall of £{discrepancy:,.2f} requires evidence (£{total_amt:,.2f} - £{accounted_for:,.2f} = £{discrepancy:,.2f})",
                            'traced_amount': traced_amt,
                            'untraced_amount': identified_untraced,
                            'accounted_for': accounted_for,
                            'claimed_amount': total_amt,
                            'discrepancy_amount': discrepancy
                        })
                        print(f"  ✅ DISCREPANCY ADDED TO DIFFERENCES")
                    else:
                        print(f"  ℹ️ No significant discrepancy (< £100)")
                    
                    # Log final differences count
                    print(f"  📋 Total differences added: {len(doc_ver['differences'])}")
                    for i, d in enumerate(doc_ver['differences']):
                        print(f"     {i+1}. {d.get('field')}: {d.get('severity')}")
                    
                    # Update confidence based on traced percentage
                    if traced_pct >= 95:
                        doc_ver['confidence'] = 0.9
                        doc_ver['verified'] = True
                    elif traced_pct >= 80:
                        doc_ver['confidence'] = 0.7
                    elif traced_pct >= 50:
                        doc_ver['confidence'] = 0.5
                    else:
                        doc_ver['confidence'] = 0.3
                    
                    # Set manual_review_status for savings claims - pending if not fully traced
                    if traced_pct < 95:
                        doc_ver['manual_review_status'] = 'pending'
                    
                    # Update issues
                    doc_ver['issues'] = [f'Funds lineage analysis complete: {traced_pct}% traced']
                    if statement_gaps:
                        doc_ver['issues'].append(f'{len(statement_gaps)} transactions need additional account statements')
                    if unresolved:
                        doc_ver['issues'].append(f'{len(unresolved)} items require additional evidence')
                    
                    # ADD QUESTION FOR MISSING BANK STATEMENTS if there are statement gaps
                    if statement_gaps and 'next_actions' in assessment_result:
                        # Build a question about the missing statements
                        gap_accounts = set()
                        total_gap_amount = 0
                        earliest_gap_date = None
                        
                        for gap in statement_gaps:
                            gap_accounts.add(gap.get('gap_account', 'source account'))
                            total_gap_amount += gap.get('amount', 0)
                            gap_date = gap.get('date', '')
                            if gap_date and (earliest_gap_date is None or gap_date < earliest_gap_date):
                                earliest_gap_date = gap_date
                        
                        # Format the question
                        accounts_str = ', '.join(gap_accounts) if gap_accounts else 'source account'
                        gap_question = (
                            f"ADDITIONAL BANK STATEMENTS REQUIRED: We have identified {len(statement_gaps)} "
                            f"transaction(s) totalling £{total_gap_amount:,.2f} that cannot be fully traced due to "
                            f"incomplete bank statement coverage. Please provide earlier statements for your "
                            f"{accounts_str} to cover the period before {earliest_gap_date or 'the earliest transaction'}."
                        )
                        
                        # Add to questions if not already present
                        if 'questions' not in assessment_result['next_actions']:
                            assessment_result['next_actions']['questions'] = []
                        
                        # Check if a similar question already exists
                        existing_questions = assessment_result['next_actions']['questions']
                        if not any('ADDITIONAL BANK STATEMENTS' in q for q in existing_questions):
                            # Insert at the beginning as it's important
                            assessment_result['next_actions']['questions'].insert(0, gap_question)
                            print(f"  📋 Added question for {len(statement_gaps)} statement gap(s)")
                        
                        # Also add to documents list
                        if 'documents' not in assessment_result['next_actions']:
                            assessment_result['next_actions']['documents'] = []
                        
                        doc_request = f"Bank statements for {accounts_str} covering period before {earliest_gap_date or 'earliest transaction'}"
                        existing_docs = assessment_result['next_actions']['documents']
                        if not any('Bank statements' in d and accounts_str in d for d in existing_docs):
                            assessment_result['next_actions']['documents'].insert(0, doc_request)
                    
                    print(f"  Added {len(doc_ver['differences'])} differences from lineage")
                    print(f"  Statement gaps: {len(statement_gaps)}")
                    print(f"  Confidence: {doc_ver['confidence']}")
                    print(f"  Verified: {doc_ver.get('verified', False)}")
                
                print(f"Funds lineage analysis complete: {lineage_result.get('summary', {}).get('traced_percentage', 0)}% traced")
                
                # IMPORTANT: Update the assessment_result in storage with modified evidence
                storage['assessment_result'] = assessment_result
            else:
                # No exact match found - user must select target transaction manually via Funds Lineage page
                print(f"  ⚠️ No exact match for £{savings_amount:,.2f} found in bank statements")
                print(f"  User must select target transaction manually via Funds Lineage tab")
                
                # Mark that lineage is pending manual selection
                if savings_claim_idx is not None and savings_claim_idx < len(evidence_matches):
                    evidence = evidence_matches[savings_claim_idx]
                    evidence['funds_lineage_pending'] = True
                    
                    # Get or create document_verification
                    if 'document_verification' not in evidence:
                        evidence['document_verification'] = {
                            'claim_id': savings_claim_idx,
                            'verified': False,
                            'confidence': 0.3,
                            'differences': [],
                            'issues': []
                        }
                    
                    doc_ver = evidence['document_verification']
                    doc_ver['differences'] = [{
                        'field': 'funds_lineage_required',
                        'severity': 'action_required',
                        'issue': f'Funds Lineage Required: No exact match for £{savings_amount:,.2f} found',
                        'expected': f'Select the transaction representing £{savings_amount:,.2f} savings in Funds Lineage tab',
                        'found': 'Manual selection required - go to Funds Lineage tab to select target transaction'
                    }]
                    doc_ver['issues'] = ['Funds lineage analysis required - select target transaction manually']
                    doc_ver['confidence'] = 0.3
                    
                    # Update storage
                    storage['assessment_result'] = assessment_result
        
        # ============================================================
        # FINAL STEP: RESTORE MANUAL ACCEPTANCES (after ALL processing)
        # This ensures manual acceptances are preserved even if other
        # code modifies the evidence during processing
        # ============================================================
        if previous_manual_acceptances:
            print(f"\n  🔄 RESTORING MANUAL ACCEPTANCES (final step)...")
            new_claims = assessment_result.get('claims', [])
            new_evidence = assessment_result.get('evidence_matches', [])
            
            for idx, claim in enumerate(new_claims):
                key = f"{claim.get('source_type', '')}_{claim.get('expected_amount', 0)}"
                if key in previous_manual_acceptances:
                    if idx < len(new_evidence):
                        preserved = previous_manual_acceptances[key]
                        evidence = new_evidence[idx]
                        
                        # Get or create document_verification
                        if 'document_verification' not in evidence:
                            evidence['document_verification'] = {}
                        doc_ver = evidence['document_verification']
                        
                        # SET ALL FIELDS NEEDED FOR "FULLY VERIFIED" DISPLAY:
                        # 1. Document verification fields
                        doc_ver['manual_review_status'] = 'accepted'
                        doc_ver['manually_accepted_by'] = preserved['manually_accepted_by']
                        doc_ver['manually_accepted_at'] = preserved['manually_accepted_at']
                        doc_ver['acceptance_reason'] = preserved.get('acceptance_reason', '')
                        doc_ver['verification_note'] = preserved.get('verification_note', 'Marked as verified post user acceptance of differences')
                        doc_ver['confidence'] = 1.0  # CRITICAL: >= 0.999 for fullyVerified
                        doc_ver['verified'] = True
                        doc_ver['requires_review'] = False
                        
                        # Mark all differences as accepted
                        if 'differences' in doc_ver:
                            for diff in doc_ver['differences']:
                                diff['accepted'] = True
                                diff['accepted_by'] = preserved['manually_accepted_by']
                                diff['accepted_at'] = preserved['manually_accepted_at']
                        
                        # 2. Evidence-level fields (CRITICAL for frontend display)
                        evidence['verified'] = True  # hasBank check
                        evidence['document_verified'] = True  # hasDocs check
                        
                        print(f"     ✅ Restored: {key}")
                        print(f"        - evidence.verified = {evidence['verified']}")
                        print(f"        - evidence.document_verified = {evidence['document_verified']}")
                        print(f"        - doc_ver.confidence = {doc_ver['confidence']}")
                        print(f"        - doc_ver.manual_review_status = {doc_ver['manual_review_status']}")
            
            # Update storage with restored evidence
            storage['assessment_result'] = assessment_result
        
        # Persist storage to database (includes lineage if run)
        _db_save_storage(db, matter_id, storage)

        return {
            "success": True,
            "matter_id": matter_id,
            "assessment": assessment_result,
            "funds_lineage_run": savings_claim is not None and storage.get('funds_lineage') is not None,
            "funds_lineage_pending": savings_claim is not None and storage.get('funds_lineage') is None,
            "document_verification_summary": storage.get('document_verification_summary', {}),
            "statement_validation_summary": storage.get('statement_validation_summary', {}),
        }

    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        print(f"\n=== ASSESSMENT ERROR ===")
        print(f"Error: {str(e)}")
        print(f"Traceback:\n{error_traceback}")
        print(f"========================\n")
        
        storage['status'] = 'error'
        storage['error'] = str(e)
        storage['error_traceback'] = error_traceback
        
        # Include traceback in response for debugging
        raise HTTPException(
            status_code=500,
            detail={
                "error": f"Assessment engine error: {str(e)}",
                "traceback": error_traceback
            }
        )


@router.get("/matters/{matter_id}/sof-assessment/results")
async def get_sof_assessment_results(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Get full SoF assessment results
    """
    print(f"\n📥 Getting assessment for matter {matter_id}")

    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Load from database
    storage = _db_load_storage(db, matter_id)

    if storage is None:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )
    
    if storage.get('status') != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage.get('status')})"
        )
    
    # Per-claim reviewer/compliance actions travel with the assessment
    # so the frontend can render claim status + the action buttons.
    assessment = dict(storage['assessment_result'])
    assessment['claim_actions'] = storage.get('claim_actions', {})
    assessment['matter_compliance_status'] = getattr(matter, 'compliance_status', None) or 'none'

    return {
        "matter_id": matter_id,
        "assessment": assessment,
        "document_verification_summary": storage.get('document_verification_summary', {}),
        "statement_validation_summary": storage.get('statement_validation_summary', {}),
        "metadata": {
            "uploaded_files": storage['uploaded_files'],
            "bank_statements_count": len(storage['bank_statements']),
            "supporting_docs_count": len(storage['supporting_docs']),
            "completed_at": storage['last_updated']
        }
    }


@router.get("/matters/{matter_id}/sof-assessment/file-note")
async def download_file_note(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Download audit-ready file note as Word document (.docx)
    """

    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Get assessment data from DB
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )

    if storage['status'] != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage['status']})"
        )
    
    # Get file note text
    file_note = storage['assessment_result']['file_note_summary']
    
    # Generate Word document
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from fastapi.responses import StreamingResponse
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Source of Funds Assessment File Note', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add matter reference
    matter_ref = doc.add_paragraph()
    matter_ref.add_run(f'Matter Reference: {matter.reference_number or f"MAT-{matter.id}"}').bold = True
    matter_ref.add_run(f'\nClient: {matter.client_name}')
    matter_ref.add_run(f'\nDate: {datetime.now().strftime("%d %B %Y")}')
    matter_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse and format the file note sections
    sections_text = file_note.split('===')
    
    for section in sections_text:
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        section_title = lines[0].strip()
        section_content = '\n'.join(lines[1:]).strip()
        
        # Add section heading
        if section_title:
            heading = doc.add_heading(section_title, level=1)
            # Style the heading
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add section content
        if section_content:
            # Split content into paragraphs
            paragraphs = section_content.split('\n\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                    
                para = doc.add_paragraph()
                
                # Handle bullet points and formatting
                if para_text.strip().startswith('•') or para_text.strip().startswith('-'):
                    para.style = 'List Bullet'
                    para_text = para_text.strip()[1:].strip()
                
                # Add text with basic formatting
                for line in para_text.split('\n'):
                    if not line.strip():
                        continue
                    
                    # Bold text between ** **
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            para.add_run(part)
                        else:
                            para.add_run(part).bold = True
                    
                    para.add_run('\n')
                
                # Set paragraph spacing
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('_' * 80)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph()
    footer_text.add_run('\nDocument generated on: ' + datetime.now().strftime("%d %B %Y at %H:%M"))
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Return as downloadable Word document
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=SoF_File_Note_Matter_{matter.reference_number or matter.id}.docx"
        }
    )


@router.post("/matters/{matter_id}/sof-assessment/accept-differences")
async def accept_claim_differences(
    matter_id: int,
    request: dict,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Manually accept differences for a specific claim after manual review
    
    Body:
        {
            "claim_index": int,
            "accepted_by": str,
            "reason": str (optional)
        }
    """
    from datetime import datetime

    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Get assessment data from DB
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        raise HTTPException(
            status_code=404,
            detail="No assessment data found for this matter"
        )

    if storage['status'] != 'completed':
        raise HTTPException(
            status_code=400,
            detail=f"Assessment not completed (status: {storage['status']})"
        )
    
    claim_index = request.get('claim_index')
    accepted_by = request.get('accepted_by', 'User')
    reason = request.get('reason', 'Manual review completed - differences accepted')
    
    if claim_index is None:
        raise HTTPException(status_code=400, detail="claim_index is required")
    
    # Get the evidence match for this claim
    evidence_matches = storage['assessment_result']['evidence_matches']
    
    if claim_index < 0 or claim_index >= len(evidence_matches):
        raise HTTPException(status_code=400, detail="Invalid claim_index")
    
    evidence = evidence_matches[claim_index]
    
    # Check if this claim requires review
    if not evidence.get('document_verification', {}).get('requires_review', False):
        raise HTTPException(
            status_code=400,
            detail="This claim does not require review (already at 100% confidence)"
        )
    
    # Update the manual review status
    doc_verification = evidence['document_verification']
    doc_verification['manual_review_status'] = 'accepted'
    doc_verification['manually_accepted_by'] = accepted_by
    doc_verification['manually_accepted_at'] = datetime.now(timezone.utc).isoformat()
    doc_verification['acceptance_reason'] = reason
    
    # MARK THE CLAIM AS VERIFIED after user acceptance
    evidence['verified'] = True
    evidence['document_verified'] = True
    doc_verification['verified'] = True
    doc_verification['confidence'] = 1.0  # 100% confidence after manual acceptance
    doc_verification['requires_review'] = False  # No longer requires review
    doc_verification['verification_note'] = 'Marked as verified post user acceptance of differences'
    
    # Mark all differences as accepted
    if 'differences' in doc_verification:
        for diff in doc_verification['differences']:
            diff['accepted'] = True
            diff['accepted_by'] = accepted_by
            diff['accepted_at'] = datetime.now(timezone.utc).isoformat()
    
    # Update the overall assessment outcome if all claims are now verified
    all_verified = all(
        em.get('verified', False) and em.get('document_verified', False)
        for em in evidence_matches
    )
    
    if all_verified and 'outcome' in storage['assessment_result']:
        storage['assessment_result']['outcome']['status'] = 'approved'
        storage['assessment_result']['outcome']['recommendation'] = 'PROCEED - All claims verified (including manual acceptances)'
    
    # Persist the updated storage to database
    _db_save_storage(db, matter_id, storage)

    return {
        "success": True,
        "message": f"Differences accepted for claim {claim_index + 1} - Claim now marked as verified",
        "claim_index": claim_index,
        "accepted_by": accepted_by,
        "accepted_at": doc_verification['manually_accepted_at'],
        "reason": reason,
        "verification_note": doc_verification['verification_note'],
        "updated_status": {
            "verified": True,
            "document_verified": True,
            "requires_review": False,
            "manual_review_status": 'accepted',
            "confidence": 1.0
        }
    }


@router.post("/matters/{matter_id}/sof-assessment/funds-lineage")
async def save_funds_lineage(
    matter_id: int,
    lineage_data: dict,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Save funds lineage analysis results for a matter.
    This integrates with the SoF assessment to verify savings claims.
    Updates the Assessment Summary with lineage results.
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Load storage from database
    storage = _db_load_storage(db, matter_id)
    if storage is None:
        storage = {
            "client_info": {},
            "bank_statements": [],
            "supporting_docs": [],
            "uploaded_files": [],
            "status": "pending",
            "last_updated": None
        }
    
    # Save lineage results
    storage['funds_lineage'] = {
        'target_transaction': lineage_data.get('target_transaction'),
        'summary': lineage_data.get('summary'),
        'lineage_tree': lineage_data.get('lineage_tree'),
        'unresolved_items': lineage_data.get('unresolved_items', []),
        'external_origins': lineage_data.get('external_origins', []),
        'traced_percentage': lineage_data.get('traced_percentage', 0),
        'run_at': datetime.now(timezone.utc).isoformat(),
        'run_by': 'user'
    }
    
    print(f"\n=== SAVING MANUAL FUNDS LINEAGE ===")
    print(f"  Matter: {matter_id}")
    print(f"  Target: £{lineage_data.get('target_transaction', {}).get('amount', 0):,.2f}")
    
    # Update assessment result if it exists - integrate lineage into document_verification
    if storage.get('assessment_result') and storage['funds_lineage'].get('summary'):
        lineage_summary = storage['funds_lineage']['summary']
        unresolved_items = storage['funds_lineage'].get('unresolved_items', [])
        
        traced_pct = lineage_summary.get('traced_percentage', 0)
        traced_amt = lineage_summary.get('tracedAmount', 0)
        total_amt = lineage_summary.get('totalAmount', 0)
        external_origins = lineage_summary.get('externalOrigins', 0)
        
        print(f"  Traced: {traced_pct}% (£{traced_amt:,.2f} of £{total_amt:,.2f})")
        print(f"  Unresolved items: {len(unresolved_items)}")
        
        # Find savings claim in evidence matches and update document_verification
        evidence_matches = storage['assessment_result'].get('evidence_matches', [])
        for idx, match in enumerate(evidence_matches):
            source_type = (match.get('claim_source', '') or '').lower()
            if 'saving' in source_type or 'accumul' in source_type:
                print(f"  Updating evidence match {idx} (savings claim)")
                
                # Update with lineage information
                match['funds_lineage_verified'] = True
                match['funds_lineage_pending'] = False
                match['lineage_summary'] = {
                    'traced_amount': traced_amt,
                    'untraced_amount': lineage_summary.get('untracedAmount', 0),
                    'traced_percentage': traced_pct,
                    'external_origins_found': external_origins,
                    'requires_evidence_count': lineage_summary.get('requiresEvidence', 0),
                    'accumulation_period_days': lineage_summary.get('accumulationPeriodDays', 0),
                    'totalAmount': total_amt,
                    'tracedAmount': traced_amt,
                    'externalOrigins': external_origins
                }
                
                # Get or create document_verification
                if 'document_verification' not in match:
                    match['document_verification'] = {
                        'claim_id': idx,
                        'verified': False,
                        'confidence': 0.3,
                        'differences': [],
                        'issues': []
                    }
                
                doc_ver = match['document_verification']
                
                # Clear old differences and replace with lineage results
                doc_ver['differences'] = []
                
                # Add lineage completion status as first difference
                severity = 'verified' if traced_pct >= 80 else 'review_required' if traced_pct >= 50 else 'evidence_required'
                doc_ver['differences'].append({
                    'field': 'funds_lineage_analysis',
                    'severity': severity,
                    'issue': f'Funds Lineage Complete: {traced_pct}% of £{total_amt:,.2f} traced to origin',
                    'expected': f'£{total_amt:,.2f} traced to legitimate sources',
                    'found': f'£{traced_amt:,.2f} traced ({external_origins} external income sources identified)'
                })
                
                # Add each unresolved item as a difference with transaction ID
                for item in unresolved_items:
                    doc_ver['differences'].append({
                        'field': 'untraced_funds',
                        'severity': 'evidence_required',
                        'issue': f"Untraced: £{item.get('amount', 0):,.2f} on {item.get('date', 'unknown')}",
                        'expected': 'Source documentation to verify origin',
                        'found': item.get('description', 'Source unknown'),
                        'transaction_id': item.get('id', item.get('transaction_id', 'N/A')),
                        'amount': item.get('amount', 0),
                        'date': item.get('date', 'unknown')
                    })
                
                # Update confidence based on traced percentage
                if traced_pct >= 95:
                    doc_ver['confidence'] = 0.9
                    doc_ver['verified'] = True
                elif traced_pct >= 80:
                    doc_ver['confidence'] = 0.7
                elif traced_pct >= 50:
                    doc_ver['confidence'] = 0.5
                else:
                    doc_ver['confidence'] = 0.3
                
                # Set manual_review_status for savings claims - pending if not fully traced
                if traced_pct < 95:
                    doc_ver['manual_review_status'] = 'pending'
                
                # Also update the match-level confidence
                match['confidence'] = doc_ver['confidence']
                
                # Update issues
                doc_ver['issues'] = [f'Funds lineage analysis complete: {traced_pct}% traced']
                if unresolved_items:
                    doc_ver['issues'].append(f'{len(unresolved_items)} items require additional evidence')
                
                print(f"  Updated document_verification with {len(doc_ver['differences'])} differences")
                print(f"  Confidence: {doc_ver['confidence']}")
                break
    
    storage['last_updated'] = datetime.now(timezone.utc).isoformat()
    
    # Persist storage to database
    _db_save_storage(db, matter_id, storage)

    return {
        "success": True,
        "matter_id": matter_id,
        "message": "Funds lineage saved and assessment updated",
        "lineage_summary": storage['funds_lineage'].get('summary')
    }


@router.get("/matters/{matter_id}/sof-assessment/funds-lineage")
async def get_funds_lineage(
    matter_id: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db)
):
    """
    Retrieve saved funds lineage analysis results for a matter.
    """
    
    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    
    # Load from database
    storage = _db_load_storage(db, matter_id) or {}
    
    if not storage.get('funds_lineage'):
        return {
            "exists": False,
            "message": "No funds lineage analysis has been run for this matter"
        }
    
    return {
        "exists": True,
        "funds_lineage": storage['funds_lineage']
    }


@router.delete("/matters/{matter_id}/sof-assessment/reset")
async def reset_sof_assessment(
    matter_id: int,
    current_user: User = Depends(require_admin),
    db: Session = Depends(get_sync_db)
):
    """
    Reset/clear SoF assessment data for a matter
    """

    # Verify matter exists
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # Clear storage from database
    _db_delete_storage(db, matter_id)

    return {
        "success": True,
        "matter_id": matter_id,
        "message": "SoF assessment data cleared"
    }


@router.delete("/matters/{matter_id}/sof-assessment/uploaded-file")
async def delete_uploaded_file(
    matter_id: int,
    filename: str,
    category: str,  # 'client_info' | 'bank_statement' | 'supporting_doc'
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db),
):
    """
    Remove a single uploaded file from the SoF assessment workspace.

    Cleans up four places: the entry in storage['uploaded_files'], the
    category-specific buckets (client_info / bank_statements transactions
    tagged with source_filename / supporting_docs), every
    DocumentVerification row (with its cascaded flags + transactions),
    and the file on disk. After a successful delete the user can
    re-upload the same file from scratch.
    """
    if category not in ("client_info", "bank_statement", "supporting_doc"):
        raise HTTPException(status_code=400, detail=f"Unknown category: {category}")

    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    storage = _db_load_storage(db, matter_id)
    if not storage:
        raise HTTPException(status_code=404, detail="No uploaded files for this matter")

    uploaded_files = storage.get("uploaded_files", [])
    match_idx = next(
        (i for i, uf in enumerate(uploaded_files)
         if uf.get("filename") == filename and uf.get("category") == category),
        None,
    )
    if match_idx is None:
        raise HTTPException(status_code=404, detail="File not found in this matter")

    del uploaded_files[match_idx]
    storage["uploaded_files"] = uploaded_files

    # Strip out the parsed contents for this file from the category bucket.
    if category == "client_info":
        # The client_info slot only holds one document at a time.
        storage["client_info"] = None
    elif category == "bank_statement":
        storage["bank_statements"] = [
            txn for txn in storage.get("bank_statements", [])
            if txn.get("source_filename") != filename
        ]
    elif category == "supporting_doc":
        storage["supporting_docs"] = [
            doc for doc in storage.get("supporting_docs", [])
            if doc.get("filename") != filename
        ]

    # Clear any cached assessment outputs — they're stale once the
    # underlying files change. The user will re-run after re-upload.
    storage.pop("results", None)
    storage.pop("rationale", None)
    storage.pop("funds_lineage", None)
    storage["last_updated"] = datetime.now(timezone.utc).isoformat()
    if not storage["uploaded_files"]:
        storage["status"] = "pending"

    _db_save_storage(db, matter_id, storage)

    # Drop DocumentVerification rows (cascades to flags + transactions
    # via ondelete=CASCADE on their FKs).
    dv_rows = db.query(DocumentVerification).filter(
        DocumentVerification.matter_id == matter_id,
        DocumentVerification.filename == filename,
    ).all()

    disk_targets: set[str] = {filename}
    for dv in dv_rows:
        if dv.disk_filename:
            disk_targets.add(dv.disk_filename)
        db.delete(dv)
    db.commit()

    upload_dir = f"/app/uploads/{matter_id}"
    for disk_name in disk_targets:
        path = os.path.join(upload_dir, os.path.basename(disk_name))
        if os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass

    return {
        "success": True,
        "matter_id": matter_id,
        "filename": filename,
        "category": category,
        "message": "File removed",
    }


# ---------------------------------------------------------------------------
# Per-claim reviewer / compliance actions
# ---------------------------------------------------------------------------
# Claims are objects inside the assessment_result JSON, not DB rows, so
# the reviewer actions on them are stored in storage['claim_actions'],
# a dict keyed by claim index:
#   { "0": {"sufficient": {...}, "compliance": {...}}, ... }

def _claim_action_entry(storage: dict, claim_index: int) -> dict:
    actions = storage.setdefault('claim_actions', {})
    return actions.setdefault(str(claim_index), {})


def _any_claim_in_review(storage: dict) -> bool:
    for entry in (storage.get('claim_actions') or {}).values():
        comp = entry.get('compliance') or {}
        if comp.get('state') == 'in_review':
            return True
    return False


@router.post("/matters/{matter_id}/sof-assessment/claims/{claim_index}/sufficient-evidence")
async def mark_claim_sufficient(
    matter_id: int,
    claim_index: int,
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db),
):
    """Mark a claim's evidence as sufficient — the reviewer confirms
    nothing further is needed and the claim is treated as verified."""
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    storage = _db_load_storage(db, matter_id)
    if not storage:
        raise HTTPException(status_code=404, detail="No assessment data for this matter")

    who = current_user.full_name or current_user.email
    when = datetime.now(timezone.utc).isoformat()
    entry = _claim_action_entry(storage, claim_index)
    entry['sufficient'] = {'by': who, 'at': when}
    _db_save_storage(db, matter_id, storage)

    claims = (storage.get('assessment_result') or {}).get('claims') or []
    claim_label = ''
    if 0 <= claim_index < len(claims):
        claim_label = str(claims[claim_index].get('source_type', '')).replace('_', ' ')

    db.add(AuditLog(
        matter_id=matter_id,
        user_id=current_user.id,
        action=AuditLogAction.APPROVED,
        entity_type="sof_claim",
        entity_id=matter_id,
        description=(
            f"Claim {claim_index + 1} ({claim_label}) marked as having sufficient "
            f"evidence by {who}."
        ),
        details={"claim_index": claim_index, "reviewer": who},
    ))
    db.commit()
    return {"success": True, "claim_actions": storage.get('claim_actions', {})}


@router.post("/matters/{matter_id}/sof-assessment/claims/{claim_index}/send-to-compliance")
async def send_claim_to_compliance(
    matter_id: int,
    claim_index: int,
    request: Dict[str, Any],
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db),
):
    """Send a single claim to the compliance team for review. Puts the
    claim — and the matter — into 'in_review' and notifies admins. A
    reason is required so compliance can see why it was referred."""
    reason = (request.get('reason') or '').strip()
    if len(reason) < 10:
        raise HTTPException(
            status_code=400,
            detail="A reason (at least 10 characters) is required so compliance can see why the claim was referred.",
        )
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    storage = _db_load_storage(db, matter_id)
    if not storage:
        raise HTTPException(status_code=404, detail="No assessment data for this matter")

    who = current_user.full_name or current_user.email
    now = datetime.now(timezone.utc)
    entry = _claim_action_entry(storage, claim_index)
    entry['compliance'] = {
        'state': 'in_review',
        'sent_by': who,
        'sent_at': now.isoformat(),
        'reason': reason,
    }
    _db_save_storage(db, matter_id, storage)

    matter.compliance_status = 'in_review'
    matter.compliance_submitted_at = now
    matter.compliance_submitted_by = who
    matter.compliance_reason = reason

    claims = (storage.get('assessment_result') or {}).get('claims') or []
    claim_label = ''
    if 0 <= claim_index < len(claims):
        claim_label = str(claims[claim_index].get('source_type', '')).replace('_', ' ')

    db.add(AuditLog(
        matter_id=matter_id,
        user_id=current_user.id,
        action=AuditLogAction.UPDATED,
        entity_type="sof_claim",
        entity_id=matter_id,
        description=(
            f"Claim {claim_index + 1} ({claim_label}) on matter "
            f"{matter.reference_number} sent to compliance review by {who}. "
            f"Reason: {reason}"
        ),
        details={"claim_index": claim_index, "sent_by": who, "reason": reason},
    ))
    for admin in db.query(User).filter(User.role == UserRole.ADMIN).all():
        try:
            db.add(Notification(
                user_id=admin.id,
                matter_id=matter_id,
                type="compliance_review",
                title="Claim sent for compliance review",
                message=(
                    f"{who} sent a Source of Funds claim on matter "
                    f"{matter.reference_number} ({matter.client_name}) for "
                    f"compliance review. Reason: {reason}"
                ),
            ))
        except Exception:
            pass
    db.commit()
    return {
        "success": True,
        "claim_actions": storage.get('claim_actions', {}),
        "compliance_status": matter.compliance_status,
    }


@router.post("/matters/{matter_id}/sof-assessment/claims/{claim_index}/cancel-compliance")
async def cancel_claim_compliance(
    matter_id: int,
    claim_index: int,
    request: Dict[str, Any],
    current_user: User = Depends(require_analyst),
    db: Session = Depends(get_sync_db),
):
    """Cancel a claim's compliance review — the reviewer has decided it
    no longer needs compliance sign-off. A rationale is required."""
    rationale = (request.get('rationale') or '').strip()
    if len(rationale) < 10:
        raise HTTPException(
            status_code=400,
            detail="A rationale (at least 10 characters) is required to cancel a compliance review.",
        )
    matter = db.query(Matter).filter(Matter.id == matter_id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    storage = _db_load_storage(db, matter_id)
    if not storage:
        raise HTTPException(status_code=404, detail="No assessment data for this matter")

    who = current_user.full_name or current_user.email
    now = datetime.now(timezone.utc)
    entry = _claim_action_entry(storage, claim_index)
    comp = entry.setdefault('compliance', {})
    comp['state'] = 'cancelled'
    comp['cancelled_by'] = who
    comp['cancelled_at'] = now.isoformat()
    comp['cancel_rationale'] = rationale
    _db_save_storage(db, matter_id, storage)

    # If no claim is still under review, lift the matter's status.
    if not _any_claim_in_review(storage):
        matter.compliance_status = 'none'

    db.add(AuditLog(
        matter_id=matter_id,
        user_id=current_user.id,
        action=AuditLogAction.UPDATED,
        entity_type="sof_claim",
        entity_id=matter_id,
        description=(
            f"Compliance review cancelled for claim {claim_index + 1} on matter "
            f"{matter.reference_number} by {who}. Rationale: {rationale}"
        ),
        details={"claim_index": claim_index, "cancelled_by": who, "rationale": rationale},
    ))
    db.commit()
    return {
        "success": True,
        "claim_actions": storage.get('claim_actions', {}),
        "compliance_status": matter.compliance_status,
    }
